"""Aida Web UI. Flask app with split chat/results layout."""

from __future__ import annotations

import json
import logging
import os
import secrets
import sys
import threading
import time
from functools import wraps
from pathlib import Path
from urllib.error import HTTPError
from urllib.parse import urlencode
from urllib.request import Request, urlopen

from flask import Flask, Response, jsonify, redirect, render_template_string, request, session, url_for

# Add src to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..'))

import anthropic

from aida.agents.aggregate import compute_aggregate
from aida.agents.alternatives import find_alternatives
from aida.agents.baseline import calculate_baseline
from aida.agents.intake import run_intake
from aida.agents.report import generate_report_markdown
from aida.errors import UserFacingError
from aida.llm_json import ModelOutputError
from aida.models import Baseline, Project, Selections

logger = logging.getLogger(__name__)

# Timeout errors to catch from Anthropic SDK
_TIMEOUT_ERRORS = (anthropic.APITimeoutError, TimeoutError)


def step_failed(exc: Exception, step: str):
    """Turn an unexpected step failure into something a building manager can act
    on, and keep the detail in the server log.

    Raw exception text used to go straight to the chat window. Sara's June test
    run died on "Fel: Expecting value: line 1 column 4 (char 3)", which says
    nothing about what to do next. Anything unrecognised now reads as a plain
    Swedish sentence; the traceback (and, for ModelOutputError, the raw model
    output) lands in the log instead.
    """
    if isinstance(exc, UserFacingError):
        logger.warning("%s: %s", step, exc)
        return jsonify({'error': str(exc)}), exc.status_code
    if isinstance(exc, ModelOutputError):
        logger.error("%s: unreadable model output: %s | raw=%r", step, exc, exc.raw[:2000])
        return jsonify({
            'error': f'Aida kunde inte tolka svaret från modellen under {step}. '
                     'Försök igen, det brukar gå igenom vid nästa försök.'
        }), 502
    logger.exception("%s failed", step)
    return jsonify({
        'error': f'Något gick fel under {step}. Försök igen. '
                 'Står felet kvar, hör av dig så tittar vi i loggarna.'
    }), 500

app = Flask(__name__)

AIDA_PASSWORD = os.environ.get('AIDA_PASSWORD', '')
SUPABASE_URL = os.environ.get('SUPABASE_URL', '').strip()

# Session cookie signing key. When legacy password auth is the active mechanism
# (no Supabase JWT), a per-process random key breaks sessions across serverless
# instances (login loops). Require a stable key in that case. With Supabase JWT
# the Flask session isn't used for auth, so a random fallback is harmless.
_secret = os.environ.get('AIDA_SECRET_KEY', '')
if not _secret:
    if AIDA_PASSWORD and not SUPABASE_URL:
        raise RuntimeError(
            "AIDA_SECRET_KEY must be set when AIDA_PASSWORD is used — a random "
            "per-instance key breaks login across serverless instances."
        )
    _secret = secrets.token_hex(32)
app.secret_key = _secret
SUPABASE_ANON_KEY = os.environ.get('SUPABASE_ANON_KEY', '').strip()
SUPABASE_JWT_SECRET = os.environ.get('SUPABASE_JWT_SECRET', '').strip()

try:
    import jwt as pyjwt
    from jwt import PyJWKClient
except ImportError:
    pyjwt = None
    PyJWKClient = None

# JWKS client for ES256 token verification (cached, Supabase default since 2026)
_jwks_client = None


def _get_jwks_client():
    global _jwks_client
    if _jwks_client is None and PyJWKClient and SUPABASE_URL:
        _jwks_client = PyJWKClient(f"{SUPABASE_URL}/auth/v1/.well-known/jwks.json")
    return _jwks_client


def get_user_claims():
    """Verify the Supabase JWT and return its claims, or None.

    Returns a dict with at least 'sub' and, when the token carries it,
    'email'. The email is what the allowlist gate needs — the user id of a
    self-registered account tells you nothing about whether that account
    should exist.
    """
    if not pyjwt or not SUPABASE_URL:
        return None
    auth_header = request.headers.get('Authorization', '')
    if not auth_header.startswith('Bearer '):
        return None
    token = auth_header[7:]

    # Try ES256 via JWKS first (Supabase default since 2024)
    jwks = _get_jwks_client()
    if jwks:
        try:
            signing_key = jwks.get_signing_key_from_jwt(token)
            payload = pyjwt.decode(
                token, signing_key.key,
                algorithms=['ES256'], audience='authenticated'
            )
            if payload.get('sub'):
                return payload
        except Exception as e:
            app.logger.debug("ES256 JWKS validation failed: %s", e)

    # Fallback: HS256 with local secret
    if SUPABASE_JWT_SECRET:
        try:
            payload = pyjwt.decode(
                token, SUPABASE_JWT_SECRET,
                algorithms=['HS256'], audience='authenticated'
            )
            if payload.get('sub'):
                return payload
        except Exception as e:
            app.logger.debug("HS256 validation failed: %s", e)

    # Last resort: verify token via Supabase auth API (handles any algorithm)
    try:
        resp = urlopen(
            Request(
                f"{SUPABASE_URL}/auth/v1/user",
                headers={
                    'apikey': SUPABASE_ANON_KEY,
                    'Authorization': f'Bearer {token}',
                },
            ),
            timeout=5,
        )
        user_data = json.loads(resp.read().decode())
        uid = user_data.get('id')
        if uid:
            app.logger.info("Token validated via Supabase /auth/v1/user fallback")
            return {'sub': uid, 'email': user_data.get('email')}
    except Exception as e:
        app.logger.debug("Supabase /auth/v1/user fallback failed: %s", e)

    return None


def get_user_from_token():
    """Extract user_id from Supabase JWT in Authorization header."""
    claims = get_user_claims()
    return claims.get('sub') if claims else None


def _parse_allowlist(raw):
    return {e.strip().lower() for e in raw.split(',') if e.strip()}


ALLOWED_EMAILS = _parse_allowlist(os.environ.get('AIDA_ALLOWED_EMAILS', ''))


def email_is_allowed(claims):
    """Whether these token claims may use the app.

    Supabase self-signup is a project-level setting outside this codebase,
    so the app cannot assume the account set is curated. This gate makes
    that moot: an account nobody invited gets a valid token and still
    reaches nothing. Hiding the "Skapa konto" link does NOT do this — the
    anon key is public in the page, so /auth/v1/signup stays reachable
    regardless of what the UI offers.

    Unset AIDA_ALLOWED_EMAILS keeps the previous behaviour (any
    authenticated user), so deploying this change alone locks nobody out.
    """
    if not ALLOWED_EMAILS:
        return True
    email = (claims.get('email') or '').strip().lower()
    return bool(email) and email in ALLOWED_EMAILS


def supabase_request(method, path, data=None, token=None, params=None, prefer=None):
    """Make a request to Supabase REST API (PostgREST).

    `prefer` replaces the default Prefer header. Used for upserts, which need
    `resolution=merge-duplicates` on top of the representation the callers here
    already expect back.
    """
    url = f"{SUPABASE_URL}/rest/v1/{path}"
    if params:
        url += '?' + urlencode(params)
    headers = {
        'apikey': SUPABASE_ANON_KEY,
        'Content-Type': 'application/json',
        'Prefer': prefer or 'return=representation',
    }
    if token:
        headers['Authorization'] = f'Bearer {token}'
    body = json.dumps(data).encode() if data else None
    req = Request(url, data=body, headers=headers, method=method)
    try:
        with urlopen(req, timeout=30) as resp:
            resp_data = resp.read().decode()
            return json.loads(resp_data) if resp_data else None
    except HTTPError as e:
        error_body = e.read().decode()
        # Log the full Supabase error server-side, but don't leak schema/RLS
        # details to the client (routes surface str(e) on 500).
        logger.warning("Supabase error %s: %s", e.code, error_body)
        raise Exception(f"Supabase error {e.code}")


def require_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        # Supabase JWT auth
        if SUPABASE_URL:
            claims = get_user_claims()
            if not claims:
                return jsonify({'error': 'Ej inloggad'}), 401
            if not email_is_allowed(claims):
                app.logger.warning(
                    "Blocked non-allowlisted user: %s", claims.get('email')
                )
                return jsonify({'error': 'Kontot saknar behörighet'}), 403
            request.user_id = claims['sub']
            return f(*args, **kwargs)
        # Legacy password auth
        if not AIDA_PASSWORD:
            # Neither Supabase nor a password is configured. In a serverless
            # (production) deploy that means the LLM-cost endpoints would be
            # public — fail closed. Locally, keep the no-auth convenience.
            if os.environ.get('VERCEL'):
                return jsonify({'error': 'Autentisering ej konfigurerad'}), 503
            return f(*args, **kwargs)
        if session.get('authenticated'):
            return f(*args, **kwargs)
        if request.is_json:
            return jsonify({'error': 'Ej inloggad'}), 401
        return redirect(url_for('login'))
    return decorated


RATE_LIMIT_PER_MIN = int(os.environ.get('AIDA_RATE_LIMIT_PER_MIN', '15'))
RATE_LIMIT_PER_DAY = int(os.environ.get('AIDA_RATE_LIMIT_PER_DAY', '150'))

# {caller_key: [monotonic timestamps]}, trimmed on each check.
_rate_hits: dict[str, list[float]] = {}
_rate_lock = threading.Lock()


def _rate_limit_key():
    """Who to count against. User id when known, else client IP."""
    uid = getattr(request, 'user_id', None)
    if uid:
        return f"u:{uid}"
    fwd = request.headers.get('X-Forwarded-For', '')
    ip = fwd.split(',')[0].strip() or request.remote_addr or 'unknown'
    return f"ip:{ip}"


def _check_rate_limit(key):
    """Sliding-window check. Returns retry-after seconds, or None if allowed.

    Deliberately in-process. On Vercel each warm instance keeps its own
    counters, so this is a *cost guard*, not a security boundary: a caller
    spread across N instances gets up to N times the quota. It still caps
    the runaway case this exists for — one client looping a chat endpoint —
    and it costs no external dependency. A hard per-user quota needs
    durable storage (a Supabase table keyed on user id); worth doing if the
    app ever opens beyond a known set of users.
    """
    now = time.monotonic()
    with _rate_lock:
        hits = [t for t in _rate_hits.get(key, []) if now - t < 86400]
        recent = [t for t in hits if now - t < 60]
        if len(recent) >= RATE_LIMIT_PER_MIN:
            _rate_hits[key] = hits
            return max(1, int(60 - (now - recent[0])))
        if len(hits) >= RATE_LIMIT_PER_DAY:
            _rate_hits[key] = hits
            return max(1, int(86400 - (now - hits[0])))
        hits.append(now)
        _rate_hits[key] = hits

        # Opportunistic cleanup so idle callers do not accumulate forever.
        if len(_rate_hits) > 1000:
            for k in [k for k, v in _rate_hits.items()
                      if not v or now - v[-1] > 86400]:
                del _rate_hits[k]
    return None


def rate_limited(f):
    """Cap calls to endpoints that spend money on LLM tokens.

    Per caller, counted in requests, in process. Deliberately NOT a global
    spend cap: one was built and removed the same day (2026-08-31). A ceiling
    on total daily cost cannot tell "a stranger is hammering the app" from
    "four colleagues are testing it", and the second is what Aida exists for.
    Blocking the success case is worse than the risk it removes, especially
    since the OpenRouter key carries a hard monthly limit that stops runaway
    spending for real, and scripts/aida_budget_watch.py in the generalassistant
    repo warns while there is still month left to react in.

    Applied *inside* require_auth so the counter keys on the authenticated
    user where possible. Endpoints that only read or write rows are not
    wrapped — they are cheap, and throttling them would break the UI's
    normal save/load traffic.
    """
    @wraps(f)
    def decorated(*args, **kwargs):
        retry_after = _check_rate_limit(_rate_limit_key())
        if retry_after is not None:
            resp = jsonify({
                'error': 'För många anrop. Vänta en stund och försök igen.'
            })
            resp.status_code = 429
            resp.headers['Retry-After'] = str(retry_after)
            return resp
        return f(*args, **kwargs)
    return decorated


def require_supabase_auth(f):
    """Like require_auth but only allows Supabase JWT (for CRUD endpoints)."""
    @wraps(f)
    def decorated(*args, **kwargs):
        if not SUPABASE_URL:
            return jsonify({'error': 'Supabase ej konfigurerat'}), 501
        claims = get_user_claims()
        if not claims:
            return jsonify({'error': 'Ej inloggad'}), 401
        if not email_is_allowed(claims):
            app.logger.warning(
                "Blocked non-allowlisted user: %s", claims.get('email')
            )
            return jsonify({'error': 'Kontot saknar behörighet'}), 403
        request.user_id = claims['sub']
        return f(*args, **kwargs)
    return decorated


LOGIN_TEMPLATE = r"""<!DOCTYPE html>
<html lang="sv">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Aida | Logga in</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Roboto:wght@400;500;700&display=swap" rel="stylesheet">
<style>
:root { --kk-gold: #FFCC01; --kk-dark-red: #B5201F; --kk-burgundy: #890200; --kk-charcoal: #444; --kk-cream: #FFF9DE; --kk-warm-bg: #FAF9F6; --kk-gray-200: #e6e4e0; --kk-gray-400: #8a8883; --kk-gray-500: #6a6864; --kk-gold-light: #FFF1B6; }
* { margin: 0; padding: 0; box-sizing: border-box; }
body { font-family: 'Roboto', sans-serif; height: 100vh; display: flex; align-items: center; justify-content: center; background: var(--kk-warm-bg); }
.login-box { background: white; border-radius: 12px; padding: 40px; width: 360px; box-shadow: 0 4px 24px rgba(0,0,0,0.08); border-top: 3px solid var(--kk-gold-light); }
.login-box h1 { font-size: 24px; color: var(--kk-charcoal); margin-bottom: 8px; }
.login-box p { font-size: 13px; color: var(--kk-gray-500); margin-bottom: 24px; }
.login-box input { width: 100%; padding: 12px 16px; border: 1px solid var(--kk-gray-200); border-radius: 8px; font-size: 14px; font-family: inherit; outline: none; }
.login-box input:focus { border-color: var(--kk-dark-red); box-shadow: 0 0 0 2px rgba(181,32,31,0.15); }
.login-box button { width: 100%; padding: 12px; background: var(--kk-charcoal); color: white; border: none; border-radius: 8px; font-size: 14px; font-weight: 600; cursor: pointer; margin-top: 12px; font-family: inherit; }
.login-box button:hover { background: var(--kk-dark-red); }
.error { color: var(--kk-dark-red); font-size: 12px; margin-top: 8px; }
.footer { position: fixed; bottom: 16px; font-size: 11px; color: var(--kk-gray-500); }
</style>
</head>
<body>
<div class="login-box">
  <h1>Aida</h1>
  <p>Klimatkalkyl och beslutsstöd för ombyggnationer</p>
  <form method="POST">
    <input type="password" name="password" placeholder="Lösenord" autofocus>
    {% if error %}<div class="error">{{ error }}</div>{% endif %}
    <button type="submit">Logga in</button>
  </form>
</div>
<div class="footer"></div>
</body>
</html>"""


@app.route('/login', methods=['GET', 'POST'])
def login():
    if not AIDA_PASSWORD:
        return redirect(url_for('index'))
    error = None
    if request.method == 'POST':
        if request.form.get('password') == AIDA_PASSWORD:
            session['authenticated'] = True
            return redirect(url_for('index'))
        error = 'Fel lösenord'
    return render_template_string(LOGIN_TEMPLATE, error=error)


@app.route('/')
def index():
    if SUPABASE_URL:
        return render_template_string(HTML_TEMPLATE,
            supabase_url=SUPABASE_URL,
            supabase_anon_key=SUPABASE_ANON_KEY,
            has_supabase=True)
    if AIDA_PASSWORD and not session.get('authenticated'):
        return redirect(url_for('login'))
    return render_template_string(HTML_TEMPLATE,
        supabase_url='', supabase_anon_key='', has_supabase=False)


@app.route('/docs/<path:filename>')
def serve_docs(filename):
    """Serve static docs files."""
    # Resolve relative to this file: src/aida/web/app.py -> project_root/docs/
    docs_dir = os.path.realpath(os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..', '..', 'docs'))
    filepath = os.path.realpath(os.path.join(docs_dir, filename))
    # commonpath (not startswith) so "/docs-private" can't pass the prefix check,
    # and realpath so symlinks can't escape the docs dir.
    if os.path.commonpath([docs_dir, filepath]) != docs_dir:
        return 'Forbidden', 403
    try:
        with open(filepath) as f:
            return f.read(), 200, {'Content-Type': 'text/html; charset=utf-8'}
    except FileNotFoundError:
        return 'Not found', 404


@app.route('/api/intake', methods=['POST'])
@require_auth
@rate_limited
def api_intake():
    data = request.json or {}
    description = data.get('description', '')
    if not description:
        return jsonify({'error': 'Beskrivning saknas'}), 400

    try:
        result = run_intake(description)
        return jsonify(result)
    except _TIMEOUT_ERRORS:
        return jsonify({'error': 'Analysen tog för lång tid. Försök igen.'}), 504
    except Exception as e:
        return step_failed(e, 'analysen av projektbeskrivningen')


@app.route('/api/baseline', methods=['POST'])
@require_auth
@rate_limited
def api_baseline():
    data = request.json or {}
    project_data = data.get('project')
    component_ids = data.get('component_ids')

    if not project_data:
        return jsonify({'error': 'Projekt saknas'}), 400

    try:
        project = Project.from_dict(project_data)

        # Partial rerun: filter the project to a subset of components before LLM matching.
        # An empty/missing component_ids means "all components" (full rerun, original behavior).
        if isinstance(component_ids, list) and component_ids:
            requested = set(component_ids)
            project.components = [c for c in project.components if c.id in requested]
            if not project.components:
                return jsonify({'error': 'Inga komponenter matchade angivna component_ids'}), 400

        baseline = calculate_baseline(project)
        return jsonify(baseline.to_dict())
    except _TIMEOUT_ERRORS:
        return jsonify({'error': 'Baslinjeberäkningen tog för lång tid. Försök igen.'}), 504
    except Exception as e:
        return step_failed(e, 'baslinjeberäkningen')


@app.route('/api/alternatives', methods=['POST'])
@require_auth
@rate_limited
def api_alternatives():
    data = request.json or {}
    project_data = data.get('project')
    baseline_data = data.get('baseline')
    component_ids = data.get('component_ids')

    if not project_data or not baseline_data:
        return jsonify({'error': 'Projekt eller baslinje saknas'}), 400

    try:
        project = Project.from_dict(project_data)
        baseline = Baseline.from_dict(baseline_data)

        # Partial rerun: filter both project and baseline to the requested subset.
        # find_alternatives iterates over baseline.components, so both must agree on which
        # components are in scope. Empty/missing list = all components (full rerun).
        if isinstance(component_ids, list) and component_ids:
            requested = set(component_ids)
            project_ids_before = {c.id for c in project.components}
            baseline_ids_before = {c.component_id for c in baseline.components}
            project.components = [c for c in project.components if c.id in requested]
            baseline.components = [c for c in baseline.components if c.component_id in requested]
            if not project.components:
                missing = sorted(requested - project_ids_before)
                return jsonify({'error': f'Komponent saknas i projektet: {missing}'}), 400
            if not baseline.components:
                missing = sorted(requested - baseline_ids_before)
                return jsonify({
                    'error': f'Komponent saknas i baslinjen: {missing}. Kör om baslinjen först.'
                }), 400

        user_feedback = data.get('user_feedback')
        result = find_alternatives(project, baseline, user_feedback=user_feedback)
        return jsonify(result.to_dict())
    except _TIMEOUT_ERRORS:
        return jsonify({'error': 'Alternativanalysen tog för lång tid. Försök igen, eller minska antal komponenter.'}), 504
    except Exception as e:
        return step_failed(e, 'alternativsökningen')


@app.route('/api/aggregate', methods=['POST'])
@require_auth
def api_aggregate():
    from aida.overrides import apply_to_selections_payload

    data = request.json or {}
    try:
        project = Project.from_dict(data.get('project', {}))
        # Laid on before the model classes are built, because Selections.from_dict
        # drops keys it does not know and an override that vanished here would
        # produce a total that quietly disagrees with the sheet on screen.
        payload = apply_to_selections_payload(data.get('selections', {}), data.get('overrides'))
        selections = Selections.from_dict(payload)
        result = compute_aggregate(project, selections)
        return jsonify(result.to_dict())
    except Exception as e:
        return step_failed(e, 'sammanställningen')


@app.route('/api/report', methods=['POST'])
@require_auth
def api_report():
    from aida.overrides import apply_to_selections_payload

    data = request.json or {}
    try:
        project = Project.from_dict(data.get('project', {}))
        overrides = data.get('overrides') or {}
        payload = apply_to_selections_payload(data.get('selections', {}), overrides)
        selections = Selections.from_dict(payload)
        if not selections.components:
            return jsonify({'error': 'Inga komponenter valda'}), 400
        # The overrides go in twice on purpose: once folded into the figures, so
        # every total is the one on screen, and once as a list, so the report can
        # say which numbers are not Aida's own. Word gets nothing but the
        # markdown, so a marking that is not in the text does not exist.
        markdown = generate_report_markdown(
            project, selections, overrides=data.get('overrides'),
        )
        return jsonify({'markdown': markdown})
    except _TIMEOUT_ERRORS:
        return jsonify({'error': 'Rapportgenereringen tog för lång tid. Försök igen.'}), 504
    except Exception as e:
        return step_failed(e, 'rapportgenereringen')


@app.route('/api/report/docx', methods=['POST'])
@require_auth
def api_report_docx():
    """Convert markdown report to .docx and return as download."""
    data = request.json or {}
    markdown = data.get('markdown', '')
    if not markdown:
        return jsonify({'error': 'Markdown saknas'}), 400

    try:
        import io
        import re
        from datetime import date

        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt, RGBColor

        BRAND_BLUE = RGBColor(0x4A, 0x90, 0xD9)
        GRAY_66 = RGBColor(0x66, 0x66, 0x66)
        GRAY_99 = RGBColor(0x99, 0x99, 0x99)
        HEADER_BG = "4A90D9"

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Base style
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)

        # Heading styles
        for level in range(1, 4):
            h_style = doc.styles[f'Heading {level}']
            h_style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        # Aida branding header
        header_p = doc.add_paragraph()
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header_p.add_run('Aida')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BRAND_BLUE
        run = header_p.add_run('  |  Klimatberäkning av ombyggnad')
        run.font.size = Pt(10)
        run.font.color.rgb = GRAY_66

        active_table = None
        table_is_first_row = False

        def _add_rich_runs(paragraph, text):
            """Parse inline markdown (bold, italic) into runs on a paragraph."""
            # Split on bold (**text**) and italic (*text*) markers
            parts = re.split(r'(\*\*[^*]+?\*\*|\*[^*]+?\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = paragraph.add_run(part[2:-2])
                    r.bold = True
                elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                    r = paragraph.add_run(part[1:-1])
                    r.italic = True
                else:
                    paragraph.add_run(part)

        def _add_rich_paragraph(text, style_name=None):
            """Add paragraph with bold/italic markdown spans preserved."""
            p = doc.add_paragraph(style=style_name)
            _add_rich_runs(p, text)
            return p

        def _style_header_cell(cell):
            """Apply white-on-blue header styling to a table cell."""
            from docx.oxml.ns import qn
            shading = cell._element.find(qn('w:tcPr'))
            if shading is None:
                tc_pr = cell._element.makeelement(qn('w:tcPr'), {})
                cell._element.insert(0, tc_pr)
            else:
                tc_pr = shading
            shading_el = tc_pr.makeelement(qn('w:shd'), {
                qn('w:val'): 'clear',
                qn('w:color'): 'auto',
                qn('w:fill'): HEADER_BG,
            })
            tc_pr.append(shading_el)
            for paragraph in cell.paragraphs:
                for r in paragraph.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    r.font.size = Pt(10)

        for line in markdown.split('\n'):
            stripped = line.strip()

            if stripped.startswith('#### '):
                active_table = None
                p = doc.add_paragraph()
                run = p.add_run(stripped[5:])
                run.bold = True
                run.font.size = Pt(11)
            elif stripped.startswith('### '):
                active_table = None
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith('## '):
                active_table = None
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('# '):
                active_table = None
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                active_table = None
                _add_rich_paragraph(stripped[2:], style_name='List Bullet')
            elif re.match(r'^\d+\.\s', stripped):
                active_table = None
                text = re.sub(r'^\d+\.\s', '', stripped)
                _add_rich_paragraph(text, style_name='List Number')
            elif stripped.startswith('|') and '|' in stripped[1:]:
                # Split on unescaped pipes only. A "|" inside a cell is written
                # as "\|" by report.cell(), because a row that splits into the
                # wrong number of cells used to be dropped without a word - and
                # the rows most likely to contain a pipe are the appendix rows
                # that disclose a figure is not Aida's.
                cells = [
                    c.strip().replace('\\|', '|')
                    for c in re.split(r'(?<!\\)\|', stripped)[1:-1]
                ]
                # Skip separator rows (|---|---|)
                if cells and all(set(c) <= {'-', ':', ' '} for c in cells):
                    continue
                if cells:
                    if active_table is None:
                        active_table = doc.add_table(rows=0, cols=len(cells))
                        active_table.style = 'Table Grid'
                        active_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        table_is_first_row = True
                    # Fit rather than drop. A row we cannot place is still a row
                    # the reader was meant to see, so a mismatch loses at most
                    # the layout, never the disclosure.
                    ncols = len(active_table.columns)
                    if len(cells) > ncols:
                        cells = cells[:ncols - 1] + [' '.join(cells[ncols - 1:])]
                    elif len(cells) < ncols:
                        cells = cells + [''] * (ncols - len(cells))
                    if len(cells) == ncols:
                        row = active_table.add_row()
                        for i, cell_text in enumerate(cells):
                            cell_text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
                            row.cells[i].text = cell_text
                            # Smaller font in table cells
                            for paragraph in row.cells[i].paragraphs:
                                for r in paragraph.runs:
                                    r.font.size = Pt(10)
                            if table_is_first_row:
                                _style_header_cell(row.cells[i])
                        table_is_first_row = False
            elif stripped == '':
                active_table = None
            elif stripped.startswith('---') or stripped.startswith('***'):
                # Horizontal rule -- skip
                active_table = None
            else:
                active_table = None
                _add_rich_paragraph(stripped)

        # Footer with disclaimer
        doc.add_paragraph()
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_p.add_run('Genererad av Aida | AI-stödd klimatanalys för ombyggnadsprojekt')
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY_99

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        today = date.today().strftime('%Y-%m-%d')
        filename = f'Aida_rapport_{today}.docx'

        return Response(
            buf.getvalue(),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename={filename}'},
        )
    except ImportError:
        return jsonify({'error': 'python-docx är inte installerat på servern'}), 500
    except Exception as e:
        return step_failed(e, 'Word-exporten')


@app.route('/api/route', methods=['POST'])
@require_auth
@rate_limited
def api_route():
    """Intent router (orchestration increment 1).

    Classifies the message and, when it is an advisory question, answers it in the
    same call without touching analysis state. The frontend calls this BEFORE its
    own flow switch: an advisory reply is rendered and stops; new_project and
    flow_action fall through to the existing flow.

    Request:  {message, history?, project?, baseline?, alternatives?, selections?}
    Response: {intent: 'advisory_question', reply, ...} | {intent: 'new_project'|'flow_action'}
    """
    from aida.agents.orchestrator import route

    data = request.json or {}
    message = (data.get('message') or '').strip()
    if not message:
        return jsonify({'error': 'Meddelande saknas'}), 400
    try:
        result = route(
            message=message,
            history=data.get('history', []),
            project=data.get('project'),
            baseline=data.get('baseline'),
            alternatives=data.get('alternatives'),
            selections=data.get('selections'),
        )
        return jsonify(result)
    except _TIMEOUT_ERRORS:
        return jsonify({'error': 'Routern svarade inte i tid. Försök igen.'}), 504
    except Exception as e:
        app.logger.exception("route failed")
        return step_failed(e, 'tolkningen av ditt meddelande')


@app.route('/api/chat', methods=['POST'])
@require_auth
@rate_limited
def api_chat():
    """Conversational endpoint with tool-use.

    Request:
      {
        message: str,
        history: [{role, content}],
        project?: Project,
        baseline?: Baseline,
        alternatives?: AlternativesResult,
        selections?: {component_id: Selection}
      }

    Response:
      {
        reply: str,
        state_updates?: {project?, baseline?, alternatives?, selections?},
        tool_calls?: [...]
      }
    """
    from aida.agents.chat_agent import run_chat_agent

    data = request.json or {}
    try:
        result = run_chat_agent(
            message=data.get('message', ''),
            history=data.get('history', []),
            project=data.get('project'),
            baseline=data.get('baseline'),
            alternatives=data.get('alternatives'),
            selections=data.get('selections'),
            overrides=data.get('overrides'),
            as_built=data.get('as_built'),
            epd_resolver=resolve_epd,
        )
        return jsonify(result)
    except _TIMEOUT_ERRORS:
        return jsonify({'error': 'Chatten svarade inte i tid. Försök igen.'}), 504
    except Exception as e:
        app.logger.exception("chat_agent failed")
        return step_failed(e, 'chatt-svaret')


# Tools a cell may reach. Deliberately not the whole HANDLERS table: rerun_baseline
# and rerun_alternatives exist to let the *model* queue work, and exposing them
# here would let a crafted request ask for a full recompute of everything, which
# is the one operation that costs every other component its choices.
_CELL_TOOLS = ('update_component', 'add_component', 'remove_component', 'select_alternative',
               'set_override', 'clear_override', 'set_as_built', 'bind_epd')


# One client for the whole process. The index is 5.8 MB of JSON and the client
# memoises it per instance, so building a new one per request would re-parse it
# every time somebody types a letter in the match field.
_epd_client = None
_epd_client_lock = threading.Lock()


def _get_epd_client():
    global _epd_client
    if _epd_client is None:
        with _epd_client_lock:
            if _epd_client is None:
                from aida.data.environdec_client import EnvirondecClient
                _epd_client = EnvirondecClient()
    return _epd_client


# Declarations from registers Environdec does not carry. Forbo publishes through
# EPD Norge and UL Environment, JELD-WEN through EPD Hub, so a search of the
# Environdec index alone returns nothing for Marmoleum or Swedoor -- the first
# two things a Karlstad byggledare types. `external_epds.json` is the file where
# those are entered by hand; adding a product is a row there and nothing else.
_EXTERNAL_EPD_PATH = Path(__file__).resolve().parent.parent / 'data' / 'external_epds.json'
_external_epds_cache = None
_external_epds_lock = threading.Lock()


def _external_epds() -> list[dict]:
    """The hand-entered declarations, read once per process.

    Failing to read the file is not an error worth refusing a search over: the
    Environdec hits are still worth showing. It is worth a log line, because the
    silent version of this is a search that quietly stops finding Marmoleum
    again and looks exactly like the day before the fix.
    """
    global _external_epds_cache
    if _external_epds_cache is None:
        with _external_epds_lock:
            if _external_epds_cache is None:
                try:
                    with open(_EXTERNAL_EPD_PATH, encoding='utf-8') as f:
                        _external_epds_cache = [r for r in json.load(f) if r.get('uuid')]
                except (OSError, json.JSONDecodeError, TypeError) as e:
                    logger.error("Kunde inte läsa externa EPD:er: %s", e)
                    _external_epds_cache = []
    return _external_epds_cache


def _search_external_epds(query: str, limit: int) -> list[dict]:
    """Substring match over name and owner, which is all six rows need.

    No ranking and no fuzziness on purpose. The list is tens of rows, entered by
    us, and every term a user would reach for ("Marmoleum", "Forbo", "Swedoor")
    is literally in the name or the owner. A scorer here would be machinery
    guarding a decision that has not been reached yet.
    """
    q = query.strip().lower()
    if not q:
        return []
    return [r for r in _external_epds()
            if q in (r.get('name') or '').lower()
            or q in (r.get('owner') or '').lower()][:limit]


def _external_epd_by_id(epd_id: str) -> dict | None:
    for r in _external_epds():
        if r.get('uuid') == epd_id:
            return r
    return None


def resolve_epd(epd_id: str, version: str = '') -> dict | None:
    """Look up one declaration and return it in the shape `bind_epd` stores.

    The GWP has to come from the register rather than from whoever is asking:
    this is the number that ends up in a klimatredovisning, and a figure that
    travelled through a form field or a language model is not a declaration, it
    is a claim about one.

    GWP-fossil A1-A3 only, and None when the declaration has no fossil figure.
    Falling back to GWP-total would silently mix in a biogenic credit and make
    the row incomparable to the Boverket baseline the rest of the tool uses.

    The external rows are the one place the sentence above is not literally
    true: their figure was transcribed from a published declaration by us, not
    fetched from a register API. That is still a declaration behind the number
    and not a guess, but it is a weaker chain of custody, so every row carries
    `gwp_source` and the two rows whose figure was inferred rather than read say
    so all the way into the klimatredovisning.
    """
    if not epd_id:
        return None
    ext = _external_epd_by_id(epd_id)
    if ext is not None:
        # gwp_a1a3 is the fossil figure in this file, by the same rule as above.
        gwp = ext.get('gwp_a1a3')
        return {
            'id': ext['uuid'],
            'name': ext.get('name', ''),
            'gwp_per_unit': gwp if isinstance(gwp, (int, float)) else None,
            'unit': ext.get('unit', ''),
            'gwp_basis': 'fossil' if isinstance(gwp, (int, float)) else '',
            'reg_no': ext.get('reg_no', ''),
            'gwp_source': ext.get('gwp_source', 'estimated'),
        }
    detail = _get_epd_client().fetch_epd_detail(epd_id, version=version)
    if detail is None:
        return None
    gwp = detail.gwp_fossil_a1a3
    basis = 'fossil'
    if gwp is None and detail.gwp_ghg_a1a3 is not None:
        # Same marking PR #550 established: usable, but never as if it were the
        # same indicator. The badge follows the number everywhere it goes.
        gwp = detail.gwp_ghg_a1a3
        basis = 'ghg'
    return {
        'id': detail.uuid,
        'name': detail.name,
        'gwp_per_unit': gwp,
        'unit': detail.declared_unit,
        'gwp_basis': basis if gwp is not None else '',
        'reg_no': detail.reg_no or _reg_no_from_index(detail.uuid),
        # Every outcome carries the field, so a later source cannot be added
        # without someone deciding what its chain of custody is.
        'gwp_source': 'declared',
    }


def _reg_no_from_index(uuid: str) -> str:
    """The registration number, from the index when the detail record lacks it.

    The two endpoints disagree about where this lives. The index carries `regNo`
    on every row; the detail document is supposed to carry it under
    `descriptionAndOwnership.other`, and for a good share of records it simply is
    not there (Golvabia Maxwear, S-P-06406, is one).

    Worth reaching for rather than shrugging at, because this is the string a
    reader of a klimatredovisning uses to look the declaration up and check the
    figure against its source. Without it the document says a number came from an
    EPD and gives no way to find which one. The search results show it, so it is
    on screen at the moment of binding and then thrown away, which is the worst
    of the three options.
    """
    if not uuid:
        return ''
    try:
        for entry in _get_epd_client().fetch_index():
            if entry.uuid == uuid:
                return entry.reg_no or ''
    except Exception:
        # A missing registration number is a weaker document, not a failed bind.
        logger.warning("Kunde inte slå upp regNo för %s", uuid[:8])
    return ''


@app.route('/api/match', methods=['POST'])
@require_auth
@rate_limited
def api_match():
    """Candidates for "which declaration is the thing that actually got installed".

    Searches the whole Environdec index (18 849 rows), not `epd_alternatives.json`
    (1 428). That file is pruned to the best candidates per category, which is
    right when proposing what to build and wrong here: the product a contractor
    actually delivered is usually not one of the best in its category, and a
    search that cannot find it would push the user towards a typvärde for no
    reason other than our own curation.

    Plus `external_epds.json`, and that is not a second helping of the same
    thing. Environdec is one register among several, and the ones it does not
    carry are not obscure: Forbo publishes Marmoleum through EPD Norge and UL
    Environment, JELD-WEN publishes Swedoor through EPD Hub. Searching the index
    alone returned nothing for either, so the tool answered "no declaration
    exists" about products whose figure it was already shipping in the
    alternatives catalog. Those rows go first in the list, because a hand-entered
    row only exists at all because somebody decided the product mattered.

    The GWP is deliberately not fetched for every candidate. That is one API call
    per row, and the list exists to be looked at, not summed. It arrives when a
    row is bound, which is also when it starts to matter.

    Request:  {query, geo?, component_hint?, limit?}
    Response: {candidates: [{id, name, owner, geo, reg_no, valid_until}]}
    """
    data = request.json or {}
    query = (data.get('query') or '').strip()
    if not query:
        return jsonify({'error': 'Sökningen behöver ett produktnamn.'}), 400
    try:
        limit = min(int(data.get('limit') or 12), 40)
    except (TypeError, ValueError):
        limit = 12

    external = _search_external_epds(query, limit)
    remaining = limit - len(external)

    hits = []
    if remaining > 0:
        try:
            hits = _get_epd_client().search_index(
                query,
                geo_filter=(data.get('geo') or '').strip(),
                component_hint=(data.get('component_hint') or '').strip(),
                max_results=remaining,
            )
        except Exception as e:
            # Only when the index was the whole answer. A reachable Marmoleum row
            # should not be withheld because Environdec happened to be down.
            if not external:
                return step_failed(e, 'EPD-sökningen')
            logger.warning("Environdec svarade inte, visar bara externa träffar: %s", e)

    candidates = [{
        'id': r['uuid'],
        'version': '',
        'name': r.get('name', ''),
        'owner': r.get('owner', ''),
        'geo': r.get('geo', ''),
        'reg_no': r.get('reg_no', ''),
        'valid_until': '',
        # The badge the search results show. It travels with the row into the
        # bind and out into the klimatredovisning, so what a user was told at
        # the moment of choosing is what the document says afterwards.
        'gwp_source': r.get('gwp_source', 'estimated'),
    } for r in external]

    candidates += [{
        'id': h.uuid,
        'version': h.version,
        'name': h.name,
        'owner': h.owner,
        'geo': h.geo,
        'reg_no': h.reg_no,
        'valid_until': h.valid_until,
        'gwp_source': 'declared',
    } for h in hits]

    # The cap is this endpoint's promise, not the index client's. Passing
    # max_results and trusting it makes the response size a property of a
    # dependency, and `limit` exists to bound what crosses the wire.
    return jsonify({'candidates': candidates[:limit]})


@app.route('/api/followup', methods=['POST'])
@require_auth
def api_followup():
    """The outcome table: what was installed, against baseline and against plan.

    Computed here and nowhere else. The override side needed a JS twin because
    `effectiveState` runs on every render; this does not, because every input to
    it (as_built, baseline, selections, overrides) only ever changes through a
    request that already goes to the server. So the client asks once per change
    and renders what it is told, and there is one implementation of the arithmetic
    instead of two that have to be kept honest about each other.

    Request:  {project, baseline?, selections?, as_built?, overrides?}
    Response: {rows, totals, uncertainties}
    """
    from aida import followup as followup_mod

    data = request.json or {}
    if not isinstance(data.get('project'), dict):
        return jsonify({'error': 'project måste vara ett objekt.'}), 400
    for bag in ('baseline', 'selections', 'as_built', 'overrides'):
        value = data.get(bag)
        if value is not None and not isinstance(value, dict):
            return jsonify({'error': f'{bag} måste vara ett objekt eller saknas.'}), 400

    try:
        return jsonify(followup_mod.compute(
            data.get('project'),
            data.get('baseline'),
            data.get('selections') or {},
            data.get('as_built') or {},
            overrides=data.get('overrides') or {},
        ))
    except Exception as e:
        return step_failed(e, 'uppföljningen')


@app.route('/api/followup/report', methods=['POST'])
@require_auth
def api_followup_report():
    """The klimatredovisning, rendered from the outcome table.

    Its own endpoint rather than a template flag on /api/report, because the two
    have different preconditions and the shared one would have to lose its.
    /api/report refuses an empty `selections.components`, and rightly so: a plan
    with nothing chosen is not a report. A follow-up trips that check while being
    perfectly valid, because following up a job Aida never planned is the case
    §12.6 was written for.

    Request:  {project, baseline?, selections?, as_built?, overrides?, property_ref?}
    Response: {markdown}
    """
    from aida import followup as followup_mod
    from aida.agents.report import render_followup_report

    data = request.json or {}
    if not isinstance(data.get('project'), dict):
        return jsonify({'error': 'project måste vara ett objekt.'}), 400
    for bag in ('baseline', 'selections', 'as_built', 'overrides'):
        value = data.get(bag)
        if value is not None and not isinstance(value, dict):
            return jsonify({'error': f'{bag} måste vara ett objekt eller saknas.'}), 400

    try:
        result = followup_mod.compute(
            data.get('project'),
            data.get('baseline'),
            data.get('selections') or {},
            data.get('as_built') or {},
            overrides=data.get('overrides') or {},
        )
        markdown = render_followup_report(
            data.get('project'), result,
            overrides=data.get('overrides') or {},
            property_ref=str(data.get('property_ref') or '').strip()[:200],
        )
        return jsonify({'markdown': markdown})
    except Exception as e:
        return step_failed(e, 'klimatredovisningen')


@app.route('/api/followup/facts', methods=['POST'])
@require_supabase_auth
def api_followup_facts():
    """Record what was estimated against what it turned out to be.

    Collected, not fed back. §12.6 is explicit that none of this loops into
    pricing yet: a correction drawn from a handful of projects is a rumour with a
    number on it. The point of writing it down now is that in two years there
    will be enough of it to decide with.

    The facts are computed here rather than accepted from the caller. A row in
    this table is a claim about how far Aida's estimate was off, and one that
    arrived ready-made from a browser would be a claim nobody checked.

    Request:  {analysis_id, project, baseline?, selections?, as_built?, overrides?}
    Response: {written: n}
    """
    from aida import followup as followup_mod

    token = request.headers.get('Authorization', '').replace('Bearer ', '')
    data = request.json or {}
    analysis_id = str(data.get('analysis_id') or '').strip()
    if not analysis_id:
        return jsonify({'error': 'analysis_id saknas.'}), 400
    if not isinstance(data.get('project'), dict):
        return jsonify({'error': 'project måste vara ett objekt.'}), 400
    for bag in ('baseline', 'selections', 'as_built', 'overrides'):
        value = data.get(bag)
        if value is not None and not isinstance(value, dict):
            return jsonify({'error': f'{bag} måste vara ett objekt eller saknas.'}), 400

    try:
        result = followup_mod.compute(
            data.get('project'),
            data.get('baseline'),
            data.get('selections') or {},
            data.get('as_built') or {},
            overrides=data.get('overrides') or {},
        )
        rows = followup_mod.facts(analysis_id, result['rows'])
        if not rows:
            return jsonify({'written': 0})
        for row in rows:
            row['user_id'] = request.user_id
        # Upsert on (analysis_id, component_id, field): a second follow-up of the
        # same project corrects its own earlier rows rather than adding a second
        # opinion beside them.
        supabase_request(
            'POST', 'follow_up_facts', data=rows, token=token,
            prefer='return=minimal,resolution=merge-duplicates',
        )
        return jsonify({'written': len(rows)})
    except Exception as e:
        return step_failed(e, 'uppföljningsfakta')


@app.route('/api/mutate', methods=['POST'])
@require_auth
@rate_limited
def api_mutate():
    """One state mutation, from a cell instead of from the chat.

    The point of this endpoint is that it does NOT contain any rules. It hands the
    edit to the same functions the chat agent's tools call, so an edit made by
    typing in a cell and the same edit made by asking Aida in the chat cannot
    produce different state. No model runs here, so it is fast and free.

    Request:  {tool, input, project?, baseline?, alternatives?, selections?}
    Response: {ok, message, state_updates}
    """
    from aida.mutations import apply_mutation

    data = request.json or {}
    tool = data.get('tool')
    if tool not in _CELL_TOOLS:
        return jsonify({'error': f'Otillåtet verktyg: {tool}'}), 400
    if not isinstance(data.get('input'), dict):
        return jsonify({'error': 'input måste vara ett objekt.'}), 400
    # The handlers reach straight into these bags. A missing or wrongly typed one
    # raises deep inside and surfaces as a 500 logged as an unexpected server
    # error, when what actually happened is an ordinary malformed request. A
    # project is required because every handler starts from its components; the
    # rest may legitimately be absent on an analysis that has not got that far.
    if not isinstance(data.get('project'), dict):
        return jsonify({'error': 'project måste vara ett objekt.'}), 400
    for bag, kind in (('baseline', dict), ('alternatives', dict), ('selections', dict),
                      ('overrides', dict), ('as_built', dict)):
        value = data.get(bag)
        if value is not None and not isinstance(value, kind):
            return jsonify({'error': f'{bag} måste vara ett objekt eller saknas.'}), 400

    inp = data['input']
    if tool == 'bind_epd' and 'epd' not in inp:
        # Same resolution the chat goes through, in the same place in the flow:
        # the client sends an id off the candidate list and the GWP is fetched
        # here, from the register. A client-supplied figure would be a number
        # that never passed a declaration on its way into the report.
        inp = dict(inp)
        epd_id = inp.pop('epd_id', None)
        if epd_id:
            try:
                found = resolve_epd(epd_id, version=inp.pop('version', '') or '')
            except Exception as e:
                return step_failed(e, 'EPD-uppslaget')
            if not found:
                return jsonify({'error': 'Hittade ingen deklaration med det id:t.'}), 404
            inp['epd'] = found
        else:
            inp['epd'] = None

    try:
        result = apply_mutation(
            tool=tool,
            inp=inp,
            project=data.get('project'),
            baseline=data.get('baseline'),
            alternatives=data.get('alternatives'),
            selections=data.get('selections') or {},
            # A cell has no model to follow the prompt's mandatory rerun pattern,
            # so the module applies it instead.
            auto_rerun=True,
            overrides=data.get('overrides') or {},
            as_built=data.get('as_built') or {},
        )
        return jsonify(result)
    except Exception as e:
        app.logger.exception("mutate failed: %s", tool)
        return step_failed(e, 'ändringen')


# === Analyses CRUD (Supabase) ===


def _nullable_text(value):
    """Empty form fields arrive as '', which Postgres rejects for a DATE and
    stores as a meaningless '' for TEXT. Both should read as "not answered".

    Applies to property_ref and planned_start: the browser sends '' when the
    user clears the field, and an empty string is not the same answer as a
    missing key. Without this, clearing the month field fails the whole PATCH
    with a date parse error and the analysis silently stops autosaving.
    """
    if value is None:
        return None
    if isinstance(value, str) and not value.strip():
        return None
    return value


@app.route('/api/analyses', methods=['POST'])
@require_supabase_auth
def create_analysis():
    token = request.headers.get('Authorization', '').replace('Bearer ', '')
    data = request.json or {}
    row = {
        'user_id': request.user_id,
        'name': data.get('name', 'Nytt projekt'),
        'status': data.get('status', 'intake'),
        'project_data': data.get('project_data'),
        'baseline_data': data.get('baseline_data'),
        'alternatives_data': data.get('alternatives_data'),
        'selections_data': data.get('selections_data'),
        'report_markdown': data.get('report_markdown'),
        # Orchestration increment 4: the chat is part of the analysis, not of
        # the browser that happened to run it.
        'conversation_data': data.get('conversation_data'),
        # What was actually installed (§12.6). Its own column rather than a ride
        # in project_data, for the same reason conversation_data got one: an
        # analysis started straight in follow-up mode has no project_data to
        # ride in until intake has run, and following up a project that was
        # never calculated in Aida is a normal case, not an exception.
        'as_built_data': data.get('as_built_data'),
        # Which building this analysis is about, and roughly when the work is
        # planned. Both optional. They exist so analyses stop being isolated
        # events: two analyses on the same school can be related, and the set
        # can be read as a pipeline over time. Asked for by Anna Florqvist
        # 2026-09-07 (matching material against the ten-year plan), and the
        # cheapest way to keep that option open while the tool is being tested.
        'property_ref': _nullable_text(data.get('property_ref')),
        'planned_start': _nullable_text(data.get('planned_start')),
    }
    result = supabase_request('POST', 'analyses', data=row, token=token)
    if isinstance(result, list):
        if not result:
            return jsonify({'error': 'Kunde inte skapa analys'}), 500
        return jsonify(result[0])
    return jsonify(result)


@app.route('/api/analyses', methods=['GET'])
@require_supabase_auth
def list_analyses():
    token = request.headers.get('Authorization', '').replace('Bearer ', '')
    params = {
        'select': 'id,name,status,created_at,updated_at,property_ref,planned_start',
        'user_id': f'eq.{request.user_id}',
        'order': 'updated_at.desc',
        'limit': '20',
    }
    result = supabase_request('GET', 'analyses', token=token, params=params)
    return jsonify(result or [])


@app.route('/api/analyses/<analysis_id>', methods=['GET'])
@require_supabase_auth
def get_analysis(analysis_id):
    token = request.headers.get('Authorization', '').replace('Bearer ', '')
    params = {
        'id': f'eq.{analysis_id}',
        'user_id': f'eq.{request.user_id}',
    }
    result = supabase_request('GET', 'analyses', token=token, params=params)
    if not result:
        return jsonify({'error': 'Ej hittad'}), 404
    return jsonify(result[0])


@app.route('/api/analyses/<analysis_id>', methods=['PUT'])
@require_supabase_auth
def update_analysis(analysis_id):
    token = request.headers.get('Authorization', '').replace('Bearer ', '')
    data = request.json or {}
    update = {}
    for key in ('name', 'status', 'project_data', 'baseline_data',
                'alternatives_data', 'selections_data', 'report_markdown',
                'conversation_data', 'as_built_data'):
        if key in data:
            update[key] = data[key]
    # Kept apart from the loop above because '' has to become NULL rather than
    # be written through: an empty month string is a date parse error, and it
    # would fail the whole PATCH, not just this field.
    for key in ('property_ref', 'planned_start'):
        if key in data:
            update[key] = _nullable_text(data[key])
    params = {
        'id': f'eq.{analysis_id}',
        'user_id': f'eq.{request.user_id}',
    }
    result = supabase_request('PATCH', 'analyses', data=update, token=token, params=params)
    if not result:
        return jsonify({'error': 'Ej hittad'}), 404
    return jsonify(result[0] if isinstance(result, list) else result)


@app.route('/api/analyses/<analysis_id>', methods=['DELETE'])
@require_supabase_auth
def delete_analysis(analysis_id):
    token = request.headers.get('Authorization', '').replace('Bearer ', '')
    params = {
        'id': f'eq.{analysis_id}',
        'user_id': f'eq.{request.user_id}',
    }
    result = supabase_request('DELETE', 'analyses', token=token, params=params)
    if not result:
        return jsonify({'error': 'Ej hittad'}), 404
    return jsonify({'ok': True})


HTML_TEMPLATE = r"""<!DOCTYPE html>
<html lang="sv">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Aida | Klimatkalkyl för ombyggnationer</title>
<link rel="icon" href="data:image/svg+xml,<svg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24' fill='none' stroke='%23E84E0D' stroke-width='1.5' stroke-linecap='round'><circle cx='12' cy='12' r='5'/><path d='M12 1v3M12 20v3M1 12h3M20 12h3M4.2 4.2l2.1 2.1M17.7 17.7l2.1 2.1M4.2 19.8l2.1-2.1M17.7 6.3l2.1-2.1'/></svg>">
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Roboto:wght@300;400;500;700&display=swap" rel="stylesheet">
<script src="https://cdn.jsdelivr.net/npm/marked/marked.min.js"></script>
<script src="https://cdn.jsdelivr.net/npm/dompurify@3/dist/purify.min.js"></script>
{% if has_supabase %}<script src="https://cdn.jsdelivr.net/npm/@supabase/supabase-js@2"></script>{% endif %}
<style>
/* === Karlstads kommun färgpalett (karlstad.se-manér) === */
:root {
  --kk-gold: #FFCC01;
  --kk-gold-light: #FFF1B6;
  --kk-orange: #EF7D00;
  --kk-red-orange: #E84E0D;
  --kk-red: #D41318;
  --kk-dark-red: #B5201F;
  --kk-burgundy: #890200;
  --kk-cream: #FFF9DE;
  --kk-warm-bg: #FAF9F6;
  --kk-charcoal: #444444;
  --kk-text: #2f2e2c;
  /* Warm ramp, same one Zaid uses. The old ramp was Tailwind's neutral greys,
     which are dead flat next to gold and orange - the two panels read as two
     kits because one was warm and the other was not.
     The two darkest steps swapped jobs. --kk-gray-400 measured 2.5:1 on white
     and carried 24 pieces of secondary text, so it never passed AA for any of
     them; it is now the non-text grey (dots, hairlines, disabled) at 3.2:1 on
     --kk-gray-100, and that text moved to --kk-gray-500, which measures 5.5 on
     white and 5.3 on --kk-warm-bg. */
  --kk-gray-50: #fafaf9;
  --kk-gray-100: #f4f3f1;
  --kk-gray-200: #e6e4e0;
  --kk-gray-300: #cfccc7;
  --kk-gray-400: #8a8883;
  --kk-gray-500: #6a6864;
  --kk-gray-600: #555350;
  --green-saving: #4a7c59;
  /* State tints for the pipeline rail. The pale pair are backgrounds, the ink
     pair are the only colours small text sits in on them. Both inks are the
     brand hue darkened until 11.5px clears 4.5:1 on its own tint: #E84E0D
     measures 3.5 there and #4a7c59 measures 4.4, so neither could be used as
     written. --kk-step-idle is the same story for grey on --kk-gray-100, where
     --kk-gray-500 lands at 4.35. */
  --kk-orange-pale: #FEF4EF;
  --kk-orange-soft: #FDE7DC;
  --kk-orange-ink: #A33A08;
  --kk-green-pale: #EEF5F0;
  --kk-green-ink: #3D6A4A;
  --kk-step-idle: #666666;
}

* { margin: 0; padding: 0; box-sizing: border-box; }
body { font-family: 'Roboto', -apple-system, BlinkMacSystemFont, sans-serif; height: 100vh; display: flex; flex-direction: column; background: var(--kk-warm-bg); color: var(--kk-text); }

/* === Top bar (karlstad.se: white with warm accent line) === */
.topbar { background: white; color: var(--kk-charcoal); height: 56px; display: flex; align-items: center; justify-content: space-between; padding: 0 24px; flex-shrink: 0; border-bottom: 3px solid var(--kk-gold-light); }
/* The two sides share the free space equally, so the project name sits in the
   middle of the bar and not in the middle of whatever is left over. With plain
   space-between it drifted left every time a button was added to the right.
   flex-shrink is 0 on purpose: below roughly 900px the sides keep their content
   width and the middle gives way, which is a squeeze rather than a break. */
.topbar-logo { flex: 1 0 0; display: flex; align-items: center; gap: 10px; }
.topbar-logo svg { width: 28px; height: 28px; color: var(--kk-red-orange); }
.topbar-logo span { font-size: 16px; font-weight: 700; letter-spacing: 0.5px; color: var(--kk-charcoal); }
.topbar-center { font-size: 14px; color: var(--kk-gray-500); }
/* One control family across the whole chrome: everything you can press is a
   pill. Filled means it starts something (Nytt projekt), outlined means it is a
   setting or a side door, and the mode switch is the same outline joined into a
   group. The two top-bar items were bare underlined-blue-adjacent links wearing
   inline styles, which read as prose rather than as controls. */
.btn-pill { display: inline-flex; align-items: center; gap: 6px; height: 28px; padding: 0 12px; background: none; border: 1px solid var(--kk-gray-200); border-radius: 100px; font-family: inherit; font-size: 11.5px; font-weight: 500; color: var(--kk-gray-500); text-decoration: none; white-space: nowrap; cursor: pointer; transition: background 0.15s, border-color 0.15s, color 0.15s; }
.btn-pill:hover { background: var(--kk-gray-100); border-color: var(--kk-gray-300); color: var(--kk-charcoal); }
.btn-pill:focus-visible { outline: 2px solid var(--kk-dark-red); outline-offset: 2px; }
.topbar-right { flex: 1 0 0; justify-content: flex-end; font-size: 12px; color: var(--kk-gray-500); }

/* === Pipeline rail ===
   Was six numbered circles on a connecting line, spread edge to edge. That
   spread is why nothing could share the row, and it cost 75px. Pills pack left,
   cost about 45, and leave the right end to the mode switch — which belongs on
   this row rather than in the results panel, because the mode decides what
   shape the process takes.

   Active and done are now different hues. The old rail painted both charcoal
   and told them apart by scale, so "where am I" and "what is finished" were the
   same colour. The numbers are gone with the circles: the labels name the steps
   and the chevrons carry the order, so a number was a third way of saying the
   same thing.

   The two ink colours are the brand orange and green darkened until 11.5px text
   on their own pale tint clears 4.5:1. Zaid's #E84E0D on #FEF4EF measures 3.5,
   which is fine for his larger type and not for this. */
.progress-bar { display: flex; align-items: center; gap: 10px; padding: 9px 24px; flex-shrink: 0; background: white; border-bottom: 1px solid var(--kk-gray-200); }
.progress-bar.steps-hidden .progress-track { display: none; }
.progress-track { display: flex; align-items: center; gap: 3px; min-width: 0; overflow-x: auto; scrollbar-width: none; }
.progress-track::-webkit-scrollbar { display: none; }
.step-item { display: flex; align-items: center; gap: 6px; padding: 4px 11px; border: none; border-radius: 100px; font-family: inherit; font-size: 11.5px; font-weight: 500; white-space: nowrap; color: var(--kk-step-idle); background: var(--kk-gray-100); transition: background 0.2s, color 0.2s; }
.step-item:disabled { cursor: default; }
.step-dot { width: 7px; height: 7px; border-radius: 50%; background: var(--kk-gray-400); flex-shrink: 0; }
.step-item.active { color: var(--kk-orange-ink); background: var(--kk-orange-pale); box-shadow: inset 0 0 0 1px var(--kk-orange-soft); font-weight: 600; }
.step-item.active .step-dot { background: var(--kk-red-orange); }
.step-item.done { color: var(--kk-green-ink); background: var(--kk-green-pale); }
.step-item.done .step-dot { background: var(--kk-green-ink); }
.step-sep { color: var(--kk-gray-400); font-size: 12px; flex-shrink: 0; user-select: none; }

/* === Main layout === */
.main { display: flex; flex: 1; overflow: hidden; padding: 16px 24px 0; gap: 24px; }

/* === Chat panel (mockup: rounded, warm bg) === */
.chat-panel { width: 40%; display: flex; flex-direction: column; flex-shrink: 0; }
/* The chat used to be gold: cream header, cream input row, pale yellow body.
   The results panel next to it was neutral grey, so the two halves of the same
   screen were two kits. Both are now the same warm off-white with white cards
   on top, and the gold is back to being an accent - the top bar's line, the
   user's own bubbles, the system notes. */
.chat-container { flex: 1; display: flex; flex-direction: column; background: var(--kk-gray-50); border-radius: 12px; border: 1px solid var(--kk-gray-200); overflow: hidden; min-height: 0; }
.chat-header { padding: 10px 16px; border-bottom: 1px solid var(--kk-gray-200); background: white; display: flex; justify-content: space-between; align-items: center; }
.chat-header h2 { font-size: 15px; font-weight: 600; color: var(--kk-charcoal); }
.messages { flex: 1; overflow-y: auto; padding: 16px; display: flex; flex-direction: column; gap: 10px; }
.msg { padding: 10px 14px; border-radius: 16px; max-width: 85%; font-size: 13px; line-height: 1.5; }
.msg.user { background: var(--kk-cream); color: var(--kk-text); align-self: flex-end; border-bottom-right-radius: 4px; }
.msg.bot { background: white; color: var(--kk-text); align-self: flex-start; border-bottom-left-radius: 4px; border: 1px solid var(--kk-gray-200); }
@keyframes msgIn { from { opacity: 0; transform: translateY(8px); } to { opacity: 1; transform: translateY(0); } }
.msg { animation: msgIn 0.25s ease-out; }
.msg.system { background: var(--kk-cream); font-size: 12px; text-align: center; align-self: center; max-width: 100%; color: var(--kk-gray-500); border: 1px solid var(--kk-gray-200); }
.msg p { margin: 0 0 8px; }
.msg p:last-child { margin-bottom: 0; }
.msg ol, .msg ul { margin: 6px 0; padding-left: 20px; }
.msg li { margin-bottom: 4px; }
.msg h1, .msg h2, .msg h3, .msg h4 { margin: 8px 0 4px; line-height: 1.3; }
.msg h1 { font-size: 16px; } .msg h2 { font-size: 15px; } .msg h3 { font-size: 14px; } .msg h4 { font-size: 13px; }
.msg code { background: rgba(0,0,0,0.06); padding: 1px 4px; border-radius: 3px; font-size: 12px; }
.msg pre { background: rgba(0,0,0,0.06); padding: 8px 10px; border-radius: 6px; overflow-x: auto; margin: 6px 0; }
.msg pre code { background: none; padding: 0; }
.msg table { border-collapse: collapse; width: 100%; margin: 6px 0; font-size: 12px; }
.msg table th, .msg table td { padding: 4px 8px; border: 1px solid var(--kk-gray-200); text-align: left; }
.msg table th { background: var(--kk-gray-50); font-weight: 600; }
.msg blockquote { border-left: 3px solid var(--kk-gray-300); margin: 6px 0; padding: 2px 10px; color: var(--kk-gray-500); }
.chat-input { padding: 12px 16px; border-top: 1px solid var(--kk-gray-200); background: white; display: flex; align-items: center; gap: 8px; }
.chat-input input { flex: 1; padding: 10px 16px; border: 1px solid var(--kk-gray-200); border-radius: 24px; font-size: 13px; font-family: inherit; background: white; outline: none; }
.chat-input input:focus { border-color: var(--kk-dark-red); box-shadow: 0 0 0 2px rgba(181,32,31,0.15); }
.chat-input button { width: 40px; height: 40px; border-radius: 50%; background: var(--kk-charcoal); color: white; border: none; cursor: pointer; display: flex; align-items: center; justify-content: center; transition: background 0.2s; flex-shrink: 0; }
.chat-input button:hover:not(:disabled) { background: var(--kk-dark-red); }
.chat-input button:disabled { opacity: 0.4; cursor: not-allowed; }
.chat-disclaimer { text-align: center; font-size: 11px; color: var(--kk-gray-500); padding: 6px 0 12px; }

/* === Results panel (mockup: tabs + white bg) === */
/* Same card as the chat: one border, one radius, the same warm surface. It used
   to be a bare grey column, which is what made the two panels look like they
   came from different products. */
.results-panel { width: 60%; display: flex; flex-direction: column; overflow: hidden; min-height: 0; background: var(--kk-gray-50); border: 1px solid var(--kk-gray-200); border-radius: 12px; }
.results-content { flex: 1; overflow-y: auto; padding: 20px 8px; background: var(--kk-gray-50); }

/* === Component cards (mockup style) === */
.comp-card { background: white; border: 1px solid var(--kk-gray-200); border-radius: 8px; overflow: hidden; margin-bottom: 16px; }
.comp-card-header { padding: 12px 16px; background: var(--kk-gray-50); border-bottom: 1px solid var(--kk-gray-200); }
.comp-card-header h3 { font-size: 14px; font-weight: 600; color: var(--kk-charcoal); }
.comp-table { width: 100%; border-collapse: collapse; font-size: 13px; }
.comp-table th { padding: 8px 12px; text-align: left; font-weight: 500; color: var(--kk-gray-500); font-size: 12px; border-bottom: 1px solid var(--kk-gray-200); }
.comp-table td { padding: 10px 12px; border-bottom: 1px solid var(--kk-gray-100); }
.comp-table tr:last-child td { border-bottom: none; }
.alt-row { cursor: pointer; transition: background 0.15s; }
.alt-row:hover { background: var(--kk-gray-50); }
.alt-row.selected { background: var(--kk-gold-light) !important; }
.alt-row input[type=radio] { accent-color: var(--kk-charcoal); }
/* Per-component usage_context — subtle callout under component name in tables */
.usage-context { margin-top: 6px; padding: 8px 12px 8px 14px; background: var(--kk-gray-50); border-left: 2px solid var(--kk-gray-300); border-radius: 0 3px 3px 0; font-size: 12.5px; line-height: 1.5; color: var(--kk-gray-500); font-style: normal; }
.usage-context-label { display: block; font-size: 9.5px; font-weight: 700; letter-spacing: 1.3px; text-transform: uppercase; color: var(--kk-gray-500); margin: 0 0 3px; }

/* === Needs analysis card (editorial pairing) === */
.needs-card { background: white; border: 1px solid var(--kk-gray-200); border-radius: 6px; overflow: hidden; margin-bottom: 16px; }
.needs-card-head { padding: 14px 20px 12px; border-bottom: 1px solid var(--kk-gray-100); display: flex; align-items: baseline; justify-content: space-between; gap: 16px; flex-wrap: wrap; }
.needs-card-title { font-size: 15px; font-weight: 500; color: var(--kk-charcoal); }
.needs-card-sub { font-size: 11.5px; color: var(--kk-gray-500); font-style: italic; }
.needs-body { padding: 20px 24px 22px; }
.needs-empty { padding: 16px 20px; font-size: 12px; color: var(--kk-gray-500); }

/* Voice blocks */
.voice-block { position: relative; padding: 4px 0 4px 22px; }
.voice-block + .voice-block { margin-top: 0; }
.voice-block::before { content: ''; position: absolute; left: 0; top: 6px; bottom: 6px; width: 3px; border-radius: 2px; }
.voice-user::before { background: var(--kk-gray-300); }
.voice-aida::before { background: var(--kk-red); }
.voice-label { font-size: 10.5px; font-weight: 700; letter-spacing: 1.5px; text-transform: uppercase; margin-bottom: 6px; display: flex; align-items: center; gap: 8px; }
.voice-user .voice-label { color: var(--kk-gray-500); }
.voice-aida .voice-label { color: var(--kk-red); }
.voice-label .dot { display: inline-block; width: 6px; height: 6px; border-radius: 50%; background: currentColor; opacity: 0.7; }
.voice-text { font-size: 14.5px; line-height: 1.6; color: var(--kk-charcoal); }
.voice-user .voice-text { color: #5a5854; }
.voice-text em.empty { color: var(--kk-gray-500); }

/* Transition between user and aida */
.voice-transition { margin: 12px 0 12px 22px; font-size: 11.5px; color: var(--kk-gray-500); display: flex; align-items: center; gap: 8px; letter-spacing: 0.3px; }
.voice-transition::before { content: ''; height: 16px; border-left: 1.5px dashed var(--kk-gray-300); margin-left: -22px; width: 22px; }

/* Inferens edit affordance */
.voice-aida { position: relative; }
.voice-aida-actions { position: absolute; top: 0; right: 0; }
.voice-aida-edit { background: none; border: 1px solid var(--kk-gray-200); border-radius: 100px; padding: 4px 11px 4px 9px; font-size: 11px; color: var(--kk-gray-500); cursor: pointer; display: inline-flex; align-items: center; gap: 6px; font-family: inherit; transition: all 0.15s; }
.voice-aida-edit:hover { background: #FDF7F7; border-color: var(--kk-red); color: var(--kk-red); }
.voice-aida-edit svg { width: 11px; height: 11px; }
.voice-aida.is-editing .voice-text { display: none; }
.voice-aida.is-editing .voice-aida-edit { display: none; }
.voice-aida-textarea { display: none; width: 100%; min-height: 140px; border: 1.5px solid var(--kk-red); border-radius: 4px; padding: 12px 14px; font-family: inherit; font-size: 14.5px; line-height: 1.6; color: var(--kk-charcoal); background: white; resize: vertical; box-sizing: border-box; }
.voice-aida.is-editing .voice-aida-textarea { display: block; }
.voice-aida-textarea:focus { outline: none; box-shadow: 0 0 0 3px rgba(181, 32, 31, 0.15); }
.voice-aida-edit-actions { display: none; gap: 8px; margin-top: 10px; justify-content: flex-end; }
.voice-aida.is-editing .voice-aida-edit-actions { display: flex; }
.btn-na-cancel { background: none; border: 1px solid var(--kk-gray-300); color: var(--kk-gray-500); padding: 5px 14px; font-size: 11.5px; border-radius: 3px; cursor: pointer; font-family: inherit; }
.btn-na-save { background: var(--kk-charcoal); border: 1px solid var(--kk-charcoal); color: white; padding: 5px 14px; font-size: 11.5px; border-radius: 3px; cursor: pointer; font-family: inherit; }
.btn-na-save:hover { background: #2a2a2a; }

/* Meta blocks (assumptions + would_clarify) */
.needs-meta-row { margin-top: 22px; padding-top: 16px; border-top: 1px solid var(--kk-gray-100); display: grid; grid-template-columns: 1fr 1fr; gap: 24px; }
.needs-meta-label { font-size: 10.5px; font-weight: 700; letter-spacing: 1.3px; text-transform: uppercase; color: var(--kk-gray-500); margin-bottom: 6px; }
.needs-meta-list { list-style: none; margin: 0; padding: 0; font-size: 12.5px; line-height: 1.55; color: var(--kk-gray-500); }
.needs-meta-list li { position: relative; padding: 3px 0 3px 18px; }
.needs-meta-list li::before { position: absolute; left: 0; top: 3px; }
.needs-meta-assumptions li::before { content: '·'; font-size: 18px; line-height: 1; color: var(--kk-gray-300); }
.needs-meta-clarify li::before { content: '?'; font-style: italic; color: var(--kk-red); opacity: 0.55; }
@media (max-width: 640px) {
  .needs-meta-row { grid-template-columns: 1fr; gap: 16px; }
  .needs-card-head { flex-direction: column; align-items: flex-start; gap: 4px; }
  .voice-aida-actions { position: static; margin-bottom: 6px; }
}
.type-badge { display: inline-block; padding: 2px 8px; border-radius: 12px; font-size: 11px; font-weight: 600; }
.type-baseline { background: var(--kk-gray-100); color: var(--kk-charcoal); }
.type-reuse { background: var(--kk-gold-light); color: #7A6000; }
.type-optimized { background: #FDE8D0; color: var(--kk-red-orange); }

/* === Summary cards === */
.summary { display: grid; grid-template-columns: repeat(3, 1fr); gap: 12px; margin: 16px 0; }
.card { background: var(--kk-gray-50); border: 1px solid var(--kk-gray-200); border-radius: 8px; padding: 16px; }
.card .card-title { font-size: 11px; font-weight: 600; color: var(--kk-gray-500); text-transform: uppercase; letter-spacing: 0.5px; }
.card .value { font-size: 24px; font-weight: 700; color: var(--kk-charcoal); margin-top: 4px; }
.card .sublabel { font-size: 12px; color: var(--kk-gray-500); }
.card.saving .value { color: var(--green-saving); }

/* === Source badges === */
.source-badge { display: inline-block; padding: 2px 6px; border-radius: 4px; font-size: 10px; font-weight: 600; margin-right: 3px; }
.source-verified { background: #F0E0E0; color: var(--kk-burgundy); }
.source-aggregate { background: #FFE9D6; color: #7A4810; }
.source-estimate { background: var(--kk-gold-light); color: #8B6914; }
/* The only outlined badge on the page. Every source badge is a filled warm pill
   because every source is one of ours; this number is not, and the difference is
   a category rather than a shade. §12.5. */
.source-manual { background: transparent; color: var(--kk-charcoal); border: 1px solid var(--kk-charcoal); }
/* The figure itself is the affordance, so a dense numeric column does not gain
   an icon per row. */
.override-open { border-bottom: 1px dotted var(--kk-gray-400); cursor: pointer; }
.override-open:hover { border-bottom-color: var(--kk-dark-red); }
.override-row td { background: var(--kk-gray-50); border-top: none; }
.override-form { display: flex; flex-wrap: wrap; gap: 8px; align-items: flex-end; padding: 4px 0 6px; text-align: left; }
.override-form label { font-size: 10px; font-weight: 500; color: var(--kk-gray-500); letter-spacing: 0.6px; text-transform: uppercase; display: block; margin-bottom: 3px; }
.override-form input { font: inherit; font-size: 13px; padding: 6px 8px; border: 1px solid var(--kk-gray-300); border-radius: 4px; background: white; }
.override-form input:focus { outline: none; border-color: var(--kk-dark-red); box-shadow: 0 0 0 2px rgba(181,32,31,0.12); }
.override-form .ov-note { flex: 1 1 280px; min-width: 200px; }
.override-form .ov-value { width: 130px; }
.override-form .ov-actions { display: flex; gap: 8px; align-items: center; flex: 0 0 auto; }
/* Lifting an override is neither the primary action nor a sibling of Avbryt, so
   it reads as a text link rather than a third chip. */
.override-form .ov-clear { border-color: transparent; text-decoration: underline; padding-left: 4px; }
.override-hint { flex-basis: 100%; font-size: 11px; color: var(--kk-gray-500); }
.override-hint.bad { color: var(--kk-red-orange); }
/* EPD-sökträffar. Capped and scrolled rather than allowed to grow: twelve hits
   would push the rest of the outcome table off screen, and the row being matched
   has to stay visible while the choice is made. */
.match-results { flex-basis: 100%; display: flex; flex-direction: column; gap: 2px; margin-top: 6px; max-height: 260px; overflow-y: auto; border: 1px solid var(--kk-gray-200); border-radius: 6px; background: white; }
/* Full-width row, not a chip: the product name is the thing being read, and a
   name truncated to fit a pill is a match nobody can verify. */
.match-hit { display: block; width: 100%; text-align: left; font: inherit; padding: 7px 10px; background: none; border: none; border-bottom: 1px solid var(--kk-gray-100); cursor: pointer; }
.match-hit:last-child { border-bottom: none; }
.match-hit:hover { background: var(--kk-gray-50); }
.match-hit:focus-visible { outline: 2px solid var(--kk-dark-red); outline-offset: -2px; }
.match-hit-name { display: block; font-size: 13px; color: var(--kk-charcoal); }
/* Owner, geography and registration number. Which of two similarly named
   products this is depends on exactly this line, so it is quiet but present. */
.match-hit-meta { display: block; font-size: 11px; color: var(--kk-gray-500); margin-top: 1px; }
.source-legend { display: flex; gap: 16px; margin: 4px 0 12px; font-size: 12px; color: var(--kk-gray-500); }
.method-label { margin: 4px 0 8px; font-size: 11px; color: var(--kk-gray-500); font-style: italic; }

/* === Buttons === */
.btn { padding: 10px 20px; background: var(--kk-dark-red); color: white; border: none; border-radius: 8px; cursor: pointer; font-size: 13px; font-weight: 600; font-family: inherit; margin-top: 12px; transition: background 0.2s; }
.btn:hover { background: var(--kk-burgundy); }
.btn:disabled { opacity: 0.4; cursor: not-allowed; }
.btn-secondary { background: var(--kk-gray-100); color: var(--kk-charcoal); }
.btn-secondary:hover { background: var(--kk-gray-200); }

.section-title { font-size: 15px; font-weight: 600; margin: 16px 0 6px; color: var(--kk-charcoal); }
.report-area { background: white; border: 1px solid var(--kk-gray-200); border-radius: 8px; padding: 20px; margin-top: 16px; font-size: 13px; line-height: 1.6; max-height: 500px; overflow-y: auto; }
.report-area h1 { font-size: 20px; font-weight: 700; margin: 0 0 12px; color: var(--kk-charcoal); border-bottom: 2px solid var(--kk-gray-200); padding-bottom: 6px; }
.report-area h2 { font-size: 16px; font-weight: 600; margin: 16px 0 8px; color: var(--kk-charcoal); }
.report-area h3 { font-size: 14px; font-weight: 600; margin: 12px 0 6px; color: var(--kk-charcoal); }
.report-area p { margin: 0 0 10px; }
.report-area ul, .report-area ol { margin: 6px 0 10px; padding-left: 24px; }
.report-area li { margin-bottom: 4px; }
.report-area table { border-collapse: collapse; width: 100%; margin: 10px 0; font-size: 12px; }
.report-area table th, .report-area table td { padding: 6px 10px; border: 1px solid var(--kk-gray-200); text-align: left; }
.report-area table th { background: var(--kk-gray-50); font-weight: 600; font-size: 11px; color: var(--kk-gray-500); }
.report-area strong { font-weight: 600; }
.report-area blockquote { border-left: 3px solid var(--kk-gold); margin: 8px 0; padding: 4px 12px; background: var(--kk-cream); color: var(--kk-gray-500); font-style: italic; }
.report-area hr { border: none; border-top: 1px solid var(--kk-gray-200); margin: 16px 0; }

/* === Footer (karlstad.se: warm cream) === */
/* Was a full-width gold band. That made the least important sentence on the
   page the loudest thing on it, and it was the last surface still shouting after
   the panels went warm-neutral. Gold stays where it marks something: the top
   bar's line, the confirm bar, callouts. */
.footer { background: white; color: var(--kk-gray-500); height: 36px; display: flex; align-items: center; justify-content: center; font-size: 11px; flex-shrink: 0; border-top: 1px solid var(--kk-gray-200); }

/* === Scrollbar === */
::-webkit-scrollbar { width: 8px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background: var(--kk-gray-300); border-radius: 8px; }
::-webkit-scrollbar-thumb:hover { background: var(--kk-gray-400); }
html { scrollbar-width: thin; scrollbar-color: #d4d4d4 transparent; }

.empty-state { color: var(--kk-gray-500); text-align: center; margin-top: 80px; }
.empty-state p { font-size: 14px; }

/* === Auth overlay === */
#authOverlay { display: flex; align-items: center; justify-content: center; flex: 1; background: var(--kk-warm-bg); }
#authOverlay .login-box { background: white; border-radius: 12px; padding: 40px; width: 360px; box-shadow: 0 4px 24px rgba(0,0,0,0.08); border-top: 3px solid var(--kk-gold-light); }
#authOverlay .login-box h1 { font-size: 24px; color: var(--kk-charcoal); margin-bottom: 8px; }
#authOverlay .login-box p { font-size: 13px; color: var(--kk-gray-500); margin-bottom: 24px; }
#authOverlay .login-box input { width: 100%; padding: 12px 16px; border: 1px solid var(--kk-gray-200); border-radius: 8px; font-size: 14px; font-family: inherit; outline: none; margin-bottom: 8px; }
#authOverlay .login-box input:focus { border-color: var(--kk-dark-red); box-shadow: 0 0 0 2px rgba(181,32,31,0.15); }
#authOverlay .login-box button { width: 100%; padding: 12px; background: var(--kk-charcoal); color: white; border: none; border-radius: 8px; font-size: 14px; font-weight: 600; cursor: pointer; margin-top: 4px; font-family: inherit; }
#authOverlay .login-box button:hover { background: var(--kk-dark-red); }
#authOverlay .login-box button:disabled { opacity: 0.4; cursor: not-allowed; }
#authOverlay .error { color: var(--kk-dark-red); font-size: 12px; margin: 4px 0; }
#appContainer { display: flex; flex-direction: column; flex: 1; min-height: 0; }

/* === Dropdown menus === */
.project-btn { background: none; border: none; color: var(--kk-gray-500); font-size: 14px; cursor: pointer; display: flex; align-items: center; gap: 6px; padding: 6px 12px; border-radius: 6px; font-family: inherit; }
.project-btn:hover { background: var(--kk-gray-100); color: var(--kk-charcoal); }
.user-btn { background: none; border: none; color: var(--kk-gray-500); cursor: pointer; padding: 6px; border-radius: 50%; display: flex; align-items: center; }
.user-btn:hover { background: var(--kk-gray-100); color: var(--kk-charcoal); }
/* Starting a new run is the one action the top bar exists for, so it is the
   only filled control up there. It wears the same dark red as every primary
   button inside the app (.btn), which makes red this interface's "act" colour
   rather than a warning. Charcoal is spoken for: step circles, mode switch and
   the send button all use it, so a charcoal pill up here read as one more piece
   of progress chrome instead of a thing to press. White on #B5201F is 6.6:1,
   and the fill itself clears 3:1 against the white bar, so the pill's edge is
   visible without a border. */
.topbar-new-btn { display: inline-flex; align-items: center; gap: 6px; height: 32px; padding: 0 14px 0 10px; background: var(--kk-dark-red); color: white; border: none; border-radius: 100px; font-size: 13px; font-weight: 600; font-family: inherit; cursor: pointer; white-space: nowrap; transition: background 0.15s; }
.topbar-new-btn:hover { background: var(--kk-burgundy); }
.topbar-new-btn:focus-visible { outline: 2px solid var(--kk-charcoal); outline-offset: 2px; }
.topbar-new-btn svg { width: 14px; height: 14px; flex-shrink: 0; }
.dropdown-menu { position: absolute; top: calc(100% + 4px); background: white; border: 1px solid var(--kk-gray-200); border-radius: 8px; box-shadow: 0 4px 16px rgba(0,0,0,0.12); min-width: 220px; z-index: 100; padding: 4px 0; }
.dropdown-right { right: 0; }
.dropdown-header { padding: 8px 16px; font-size: 11px; font-weight: 600; color: var(--kk-gray-500); text-transform: uppercase; }
.dropdown-divider { border-top: 1px solid var(--kk-gray-200); margin: 4px 0; }
.dropdown-item { display: flex; align-items: center; gap: 8px; width: 100%; padding: 8px 16px; border: none; background: none; font-size: 13px; color: var(--kk-charcoal); cursor: pointer; font-family: inherit; text-align: left; }
.dropdown-item:hover { background: var(--kk-gray-50); }
.dropdown-item.active { background: var(--kk-gold-light); }

/* === Results tabs ===
   Holds tabs and nothing else again. The mode switch spent a few hours here
   between its own row and the pipeline rail; the rail is the right home,
   because the mode is a fact about the process rather than about this panel. */
.results-tabs { display: flex; border-bottom: 2px solid var(--kk-gray-200); flex-shrink: 0; background: white; border-radius: 12px 12px 0 0; }
.tab { padding: 10px 20px; background: none; border: none; font-size: 13px; font-weight: 500; color: var(--kk-gray-500); cursor: pointer; border-bottom: 2px solid transparent; margin-bottom: -2px; font-family: inherit; transition: all 0.2s; }
.tab:hover:not(:disabled) { color: var(--kk-charcoal); }
.tab.active { color: var(--kk-charcoal); border-bottom-color: var(--kk-dark-red); font-weight: 600; }
.tab:disabled { opacity: 0.35; cursor: not-allowed; }

/* === Mode switch (orchestration-redesign §12) ===
   Deliberately NOT styled like .tab. A tab answers "which part am I looking at";
   the mode answers "what shape is the whole thing". Giving them the same
   underline language would claim they are the same kind of choice.
   It lives at the right end of the pipeline rail. That row already answers
   "where is this process", and the mode answers "what shape does it take", so
   the two belong together; the steps pack left, so the right end was free.
   Same-row adjacency does not claim equivalence the way shared styling would.
   It owned a 38px row of its own until 2026-09-09, which was 7% of the results
   panel spent on a preference. */
.mode-switch { display: flex; gap: 2px; margin-left: auto; flex-shrink: 0; }
.mode-btn { padding: 4px 12px; font-size: 11.5px; font-weight: 500; font-family: inherit; color: var(--kk-gray-500); background: var(--kk-gray-100); border: 1px solid var(--kk-gray-200); cursor: pointer; transition: all 0.15s; }
.mode-btn:first-child { border-radius: 100px 0 0 100px; }
.mode-btn:last-child { border-radius: 0 100px 100px 0; }
.mode-btn:hover:not(.active) { color: var(--kk-charcoal); background: var(--kk-gray-200); }
.mode-btn.active { background: var(--kk-charcoal); border-color: var(--kk-charcoal); color: white; font-weight: 600; }
.mode-btn:focus-visible { outline: 2px solid var(--kk-dark-red); outline-offset: 2px; }

/* === Sheet (Arbetsblad) ===
   One scrolling column. Sections are separated by space and a hairline, not by
   boxes: the comp-cards inside are already boxes, and nesting them would read as
   a card kit rather than a document. No numbering anywhere, because the sheet is
   explicitly not a sequence. That is the whole difference from Stegvis. */
.sheet { padding: 0 12px 40px; }
.sheet-section + .sheet-section { margin-top: 32px; padding-top: 28px; border-top: 1px solid var(--kk-gray-200); }
.sheet-head { display: flex; align-items: baseline; gap: 10px; margin-bottom: 12px; }
.sheet-head h2 { font-size: 13px; font-weight: 600; color: var(--kk-charcoal); letter-spacing: 0.01em; }
.sheet-pending { font-size: 11px; color: var(--kk-gray-500); font-weight: 400; }
.sheet-empty { border: 1px dashed var(--kk-gray-300); border-radius: 8px; padding: 18px 20px; background: white; }
.sheet-empty p { font-size: 12.5px; color: var(--kk-gray-500); line-height: 1.5; max-width: 60ch; }
.sheet-empty .btn { margin-top: 14px; }

/* === Editable cells (orchestration-redesign §12.4) ===
   A cell reads as text until you touch it. Drawing every editable value as an
   input box would turn a results table into a form, and most of the time the
   user is reading, not typing. The affordance appears on hover and focus, which
   is the spreadsheet convention the name Arbetsblad is borrowing from. */
.cell-input { width: 100%; font: inherit; color: inherit; background: transparent; border: 1px solid transparent; border-radius: 4px; padding: 3px 6px; margin: -3px -6px; font-family: inherit; }
.cell-input:hover:not(:disabled) { border-color: var(--kk-gray-200); background: white; }
.cell-input:focus { outline: none; border-color: var(--kk-dark-red); background: white; box-shadow: 0 0 0 2px rgba(181,32,31,0.12); }
.cell-input:disabled { color: var(--kk-gray-500); cursor: not-allowed; }
.cell-input.saving { border-color: var(--kk-gold); background: var(--kk-cream); }
td.cell-num .cell-input { text-align: right; }
select.cell-input { cursor: pointer; }
.cell-remove { background: none; border: none; color: var(--kk-gray-500); cursor: pointer; font-size: 15px; line-height: 1; padding: 4px 6px; border-radius: 4px; font-family: inherit; }
.cell-remove:hover:not(:disabled) { color: var(--kk-dark-red); background: var(--kk-gray-100); }
.cell-remove:disabled { opacity: 0.3; cursor: not-allowed; }
.cell-add { padding: 10px 16px; border-top: 1px solid var(--kk-gray-200); }
.cell-add button { background: none; border: 1px dashed var(--kk-gray-300); color: var(--kk-gray-500); border-radius: 6px; padding: 6px 12px; font-size: 12px; cursor: pointer; font-family: inherit; }
.cell-add button:hover:not(:disabled) { border-color: var(--kk-dark-red); color: var(--kk-dark-red); }
.cell-add button:disabled { opacity: 0.4; cursor: not-allowed; }

/* === Confirm actions in chat (legacy, kept for history rendering) === */
.confirm-actions { display: none; }
.confirm-hint { display: none; }

/* === Sticky confirm bar above chat input === */
.confirm-bar { display: none; padding: 10px 16px; background: var(--kk-cream); border-top: 1px solid var(--kk-gray-200); align-items: center; gap: 10px; }
.confirm-bar.visible { display: flex; }
.confirm-bar .confirm-bar-text { flex: 1; font-size: 12px; color: var(--kk-gray-500); }
.confirm-bar .btn-confirm-sticky { padding: 8px 20px; background: var(--kk-charcoal); color: white; border: none; border-radius: 20px; font-size: 12px; font-weight: 600; cursor: pointer; font-family: inherit; transition: background 0.2s; }
.confirm-bar .btn-confirm-sticky:hover:not(:disabled) { background: var(--kk-dark-red); }
.confirm-bar .btn-confirm-sticky:disabled { opacity: 0.4; cursor: not-allowed; }

/* === Typing indicator (Feature 1) === */
.typing-indicator { display: flex; align-items: center; gap: 5px; padding: 10px 14px; }
.typing-dot { width: 7px; height: 7px; border-radius: 50%; background: var(--kk-gray-400); animation: typingBounce 1.2s ease-in-out infinite; }
.typing-dot:nth-child(2) { animation-delay: 0.2s; }
.typing-dot:nth-child(3) { animation-delay: 0.4s; }
@keyframes typingBounce { 0%, 60%, 100% { transform: translateY(0); opacity: 0.4; } 30% { transform: translateY(-6px); opacity: 1; } }
.elapsed-time { font-size: 11px; color: var(--kk-gray-500); margin-left: 4px; }
.typing-text { font-size: 12px; color: var(--kk-gray-500); margin-left: 8px; font-style: italic; }
.action-btn { padding: 6px 14px; background: var(--kk-charcoal); color: white; border: none; border-radius: 16px; font-size: 12px; font-weight: 600; cursor: pointer; font-family: inherit; transition: background 0.2s; }
.action-btn:hover:not(:disabled) { background: var(--kk-dark-red); }
.action-btn:disabled { opacity: 0.5; cursor: not-allowed; }
.project-rename-input { background: transparent; border: 1px solid var(--kk-gray-300); border-radius: 4px; padding: 2px 6px; font-size: inherit; font-family: inherit; color: inherit; outline: none; min-width: 120px; }
.project-rename-input:focus { border-color: var(--kk-charcoal); }
.meta-label { display: block; font-size: 13px; font-weight: 600; color: var(--kk-charcoal); margin-bottom: 5px; }
.meta-input { width: 100%; padding: 10px 12px; border: 1px solid var(--kk-gray-200); border-radius: 8px; font-size: 14px; font-family: inherit; color: var(--kk-charcoal); outline: none; }
.meta-input:focus { border-color: var(--kk-dark-red); box-shadow: 0 0 0 2px rgba(181,32,31,0.15); }
.meta-hint { font-size: 12px; color: var(--kk-gray-500); margin: 6px 0 18px; }
.meta-actions { display: flex; justify-content: flex-end; gap: 8px; }
.project-meta-line { font-size: 11px; color: var(--kk-gray-500); overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }

/* === Reasoning expander (Feature 2) === */
.reasoning-toggle { background: none; border: none; color: var(--kk-gray-500); font-size: 11px; cursor: pointer; padding: 0; font-family: inherit; text-decoration: underline; white-space: nowrap; }
.reasoning-toggle:hover { color: var(--kk-charcoal); }
.reasoning-row td { padding: 4px 12px 8px 44px; font-size: 12px; color: var(--kk-gray-500); line-height: 1.5; background: var(--kk-gray-50); border-bottom: 1px solid var(--kk-gray-100); }

/* === Modal (Feature 5) === */
.modal-backdrop { position: fixed; inset: 0; background: rgba(0,0,0,0.4); z-index: 200; display: flex; align-items: center; justify-content: center; }
.modal-box { background: white; border-radius: 12px; padding: 32px; max-width: 560px; width: 90%; max-height: 80vh; overflow-y: auto; position: relative; }
.modal-box h2 { font-size: 18px; font-weight: 700; color: var(--kk-charcoal); margin-bottom: 16px; }
.modal-box p, .modal-box li { font-size: 13px; line-height: 1.6; color: var(--kk-charcoal); }
.modal-box ul { padding-left: 20px; margin: 8px 0; }
.modal-box section { margin-bottom: 20px; }
.modal-box h3 { font-size: 14px; font-weight: 600; color: var(--kk-charcoal); margin-bottom: 6px; }
.modal-close { position: absolute; top: 16px; right: 16px; background: none; border: none; cursor: pointer; color: var(--kk-gray-500); font-size: 20px; }
.modal-close:hover { color: var(--kk-charcoal); }

/* === Välkomstdialog ===
   Signaturen är likhetstecknet: ombyggnation väger ungefär lika tungt som
   nybyggnation. Approximationen sitter i glyfen, inte i påhittade staplar. */
#welcomeModal { background: rgba(68,68,68,0.45); padding: 20px; }
.welcome-box { background: white; border-radius: 12px; width: 100%; max-width: 620px; max-height: calc(100vh - 40px); overflow-y: auto; position: relative; box-shadow: 0 18px 48px rgba(68,68,68,0.22); }
.welcome-close { position: absolute; top: 14px; right: 14px; background: none; border: none; cursor: pointer; color: var(--kk-gray-500); font-size: 18px; line-height: 1; padding: 6px; border-radius: 50%; }
.welcome-close:hover { color: var(--kk-charcoal); background: rgba(0,0,0,0.04); }

.welcome-intro { padding: 34px 40px 26px; }
.welcome-eyebrow { font-size: 10.5px; font-weight: 700; letter-spacing: 1.6px; text-transform: uppercase; color: var(--kk-gray-500); margin-bottom: 14px; padding-right: 34px; }
.welcome-box h2 { font-size: 27px; font-weight: 700; letter-spacing: -0.5px; line-height: 1.15; color: var(--kk-charcoal); margin: 0 0 12px; }
.welcome-lead { font-size: 16.5px; font-weight: 300; line-height: 1.55; color: var(--kk-charcoal); margin: 0; }

/* Bakgrund och syfte: en lugn inramad mening, inte en figur att avkoda */
.welcome-why { margin: 4px 40px 0; padding: 15px 20px; background: var(--kk-cream); border-left: 3px solid var(--kk-gold); border-radius: 0 4px 4px 0; }
.welcome-why p { font-size: 13.5px; line-height: 1.65; color: #6f6144; margin: 0; }
.welcome-why strong { color: var(--kk-charcoal); font-weight: 700; }

.welcome-body { padding: 24px 40px 8px; }
.welcome-label { font-size: 10.5px; font-weight: 700; letter-spacing: 1.3px; text-transform: uppercase; color: var(--kk-gray-500); margin-bottom: 10px; }
.welcome-body p { font-size: 13.5px; line-height: 1.65; color: var(--kk-charcoal); margin: 0 0 24px; }

.welcome-tips { list-style: none; margin: 0 0 4px; padding: 0; }
.welcome-tips li { padding: 13px 0; border-top: 1px solid var(--kk-gray-100); }
.welcome-tips li:first-child { border-top: none; padding-top: 4px; }
.welcome-tips strong { display: block; font-size: 14px; font-weight: 700; color: var(--kk-charcoal); margin-bottom: 3px; }
.welcome-tips span { display: block; font-size: 13px; line-height: 1.6; color: var(--kk-gray-500); }

.welcome-note { padding: 4px 40px 0; font-size: 12px; line-height: 1.6; color: var(--kk-gray-500); }
.welcome-note a { color: var(--kk-dark-red); }
/* Sticky: kryssrutan och knappen ska nås utan att rulla, även på låga fönster. */
.welcome-foot { position: sticky; bottom: 0; background: white; border-top: 1px solid var(--kk-gray-100); display: flex; align-items: center; justify-content: space-between; gap: 20px; flex-wrap: wrap; padding: 16px 40px 18px; margin-top: 16px; }
.welcome-checkbox { display: flex; align-items: center; gap: 9px; font-size: 13px; color: var(--kk-gray-500); cursor: pointer; }
.welcome-checkbox input { width: 16px; height: 16px; accent-color: var(--kk-charcoal); cursor: pointer; margin: 0; }
.welcome-checkbox:hover { color: var(--kk-charcoal); }
.welcome-start { padding: 11px 26px; background: var(--kk-charcoal); color: white; border: none; border-radius: 8px; font-size: 14px; font-weight: 600; font-family: inherit; cursor: pointer; transition: background 0.2s; }
.welcome-start:hover { background: var(--kk-dark-red); }
.welcome-box :focus-visible { outline: 2px solid var(--kk-red-orange); outline-offset: 2px; }

@keyframes welcomeRise { from { opacity: 0; transform: translateY(16px) scale(0.985); } to { opacity: 1; transform: none; } }
.welcome-box { animation: welcomeRise 0.34s cubic-bezier(0.2, 0.8, 0.3, 1) both; }

@media (max-width: 560px) {
  #welcomeModal { padding: 0; align-items: stretch; }
  .welcome-box { max-width: none; border-radius: 0; max-height: 100vh; }
  .welcome-intro { padding: 30px 22px 22px; }
  .welcome-box h2 { font-size: 23px; }
  .welcome-lead { font-size: 15.5px; }
  .welcome-why { margin: 4px 22px 0; padding: 14px 16px; }
  .welcome-body { padding: 22px 22px 8px; }
  .welcome-note { padding: 4px 22px 0; }
  .welcome-foot { padding: 14px 22px 18px; gap: 12px; }
  .welcome-start { width: 100%; }
}
@media (prefers-reduced-motion: reduce) {
  .welcome-box { animation: none; }
}

/* === Step-back navigation (Feature 10) === */
.step-item.done { cursor: pointer; }
.step-item.done:hover { background: #E3EDE6; }
.step-item:focus-visible { outline: 2px solid var(--kk-dark-red); outline-offset: 2px; }

/* === Responsive (Feature 9) === */
@media (max-width: 768px) {
  .main { flex-direction: column; overflow-y: auto; overflow-x: hidden; padding: 8px 12px 0; gap: 12px; }
  .chat-panel { width: 100%; min-height: 300px; max-height: 50vh; }
  .results-panel { width: 100%; }
  /* The pills keep their names here. Six of them do not fit, so the track
     scrolls sideways, which is why it owns the overflow rather than the bar:
     the mode switch has to stay put at the right while the steps slide. */
  .progress-bar { padding: 8px 12px; gap: 8px; }
  .step-item { padding: 4px 9px; }
  .topbar { padding: 0 12px; }
  .topbar-center { max-width: 120px; overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }
  .summary { grid-template-columns: repeat(2, 1fr); }
  .card .value { font-size: 18px; }
  .comp-table { display: block; overflow-x: auto; -webkit-overflow-scrolling: touch; }
  body { height: auto; min-height: 100vh; }
  .results-content { max-height: none; }
  .footer { padding: 8px 12px; }
}
@media (max-width: 480px) {
  .summary { grid-template-columns: 1fr; }
  .chat-panel { max-height: 45vh; }
  .results-tabs { overflow-x: auto; }
  .tab { padding: 10px 12px; font-size: 12px; white-space: nowrap; }
  .topbar-center { display: none; }
  /* The even split between logo and right cluster exists to centre
     .topbar-center, which is hidden at this width. Left on, it hands the logo
     half the bar and squeezes "Pling på" and "Om verktyget" into three-line
     wraps. */
  .topbar-logo { flex: 0 0 auto; }
  .topbar-new-btn { padding: 0 9px; }
  .topbar-new-btn span { display: none; }
}
</style>
</head>
<body>

{% if has_supabase %}
<!-- Auth overlay -->
<div id="authOverlay">
  <div class="login-box">
    <h1>Aida</h1>
    <p>Klimatkalkyl och beslutsstöd för ombyggnationer</p>
    <input type="email" id="authEmail" placeholder="E-post" autofocus>
    <input type="password" id="authPassword" placeholder="Lösenord" onkeydown="if(event.key==='Enter')handleAuth()">
    <div id="authError" class="error" style="display:none"></div>
    <button onclick="handleAuth()" id="authSubmitBtn">Logga in</button>
    <div style="text-align:center;margin-top:12px;font-size:13px;color:var(--kk-gray-500)">
      <span id="authToggleText">Inget konto?</span>
      <a href="#" onclick="toggleAuthMode(event)" id="authToggleLink" style="color:var(--kk-dark-red)">Skapa konto</a>
    </div>
  </div>
</div>
<div id="appContainer">
{% endif %}

<!-- Top bar -->
<div class="topbar">
  <div class="topbar-logo">
    <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" stroke-linecap="round"><circle cx="12" cy="12" r="5"/><path d="M12 1v3M12 20v3M1 12h3M20 12h3M4.2 4.2l2.1 2.1M17.7 17.7l2.1 2.1M4.2 19.8l2.1-2.1M17.7 6.3l2.1-2.1"/></svg>
    <span>Aida</span>
  </div>
  {% if has_supabase %}
  <div class="topbar-center" id="projectDropdown" style="position:relative">
    <button class="project-btn" onclick="toggleProjectMenu()" id="projectBtn">
      <span id="projectName">Nytt projekt</span>
      <svg width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><polyline points="6 9 12 15 18 9"/></svg>
    </button>
    <div class="dropdown-menu" id="projectMenu" style="display:none;left:50%;transform:translateX(-50%)">
      <div class="dropdown-header">Senaste projekt</div>
      <div id="projectList"></div>
      <div class="dropdown-divider"></div>
      <button class="dropdown-item" onclick="startRenameProject()">Byt namn på projektet</button>
      <button class="dropdown-item" onclick="openProjectMeta()">Fastighet och tidpunkt</button>
    </div>
  </div>
  <div class="topbar-right" id="userDropdown" style="position:relative;display:flex;align-items:center;gap:12px">
    <button class="topbar-new-btn" onclick="createNewProject()" title="Skapa nytt projekt" aria-label="Skapa nytt projekt"><svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5" stroke-linecap="round"><path d="M12 5v14M5 12h14"/></svg><span>Nytt projekt</span></button>
    <a href="#" class="btn-pill" id="soundToggle" onclick="toggleSound();return false" title="Pling när ett steg är klart, även om du är i en annan flik">🔔 Pling på</a>
    <a href="#" class="btn-pill" onclick="openAbout();return false">Om verktyget</a>
    <span id="saveIndicator" style="font-size:11px;color:var(--kk-gray-500);display:none"></span>
    <button class="user-btn" onclick="toggleUserMenu()">
      <svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5"><path d="M20 21v-2a4 4 0 0 0-4-4H8a4 4 0 0 0-4 4v2"/><circle cx="12" cy="7" r="4"/></svg>
    </button>
    <div class="dropdown-menu dropdown-right" id="userMenu" style="display:none">
      <div class="dropdown-header" id="userEmail"></div>
      <div class="dropdown-divider"></div>
      <button class="dropdown-item" onclick="handleLogout()">Logga ut</button>
    </div>
  </div>
  {% else %}
  <div class="topbar-center"></div>
  <div class="topbar-right" style="display:flex;align-items:center;gap:12px"><button class="topbar-new-btn" onclick="createNewProject()" title="Skapa nytt projekt" aria-label="Skapa nytt projekt"><svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.5" stroke-linecap="round"><path d="M12 5v14M5 12h14"/></svg><span>Nytt projekt</span></button><a href="#" class="btn-pill" id="soundToggle" onclick="toggleSound();return false" title="Pling när ett steg är klart, även om du är i en annan flik">🔔 Pling på</a><a href="#" class="btn-pill" onclick="openAbout();return false">Om verktyget</a><span id="saveIndicator" style="font-size:11px;color:var(--kk-gray-500);display:none"></span><span style="font-size:12px;color:var(--kk-gray-500)">Prototyp</span></div>
  {% endif %}
</div>

<!-- Progress tracker -->
<div class="progress-bar">
  <div class="progress-track">
    <button type="button" class="step-item" id="st-planering" disabled><span class="step-dot"></span>Projektbeskrivning</button>
    <span class="step-sep" aria-hidden="true">&#x203A;</span>
    <button type="button" class="step-item" id="st-baslinje" disabled><span class="step-dot"></span>Baslinje</button>
    <span class="step-sep" aria-hidden="true">&#x203A;</span>
    <button type="button" class="step-item" id="st-aterbruk" disabled><span class="step-dot"></span>&#xC5;terbruk</button>
    <span class="step-sep" aria-hidden="true">&#x203A;</span>
    <button type="button" class="step-item" id="st-nyproduktion" disabled><span class="step-dot"></span>Nyproduktion</button>
    <span class="step-sep" aria-hidden="true">&#x203A;</span>
    <button type="button" class="step-item" id="st-sammanstallning" disabled><span class="step-dot"></span>Sammanst&#xE4;llning</button>
    <span class="step-sep" aria-hidden="true">&#x203A;</span>
    <button type="button" class="step-item" id="st-uppfoljning" disabled><span class="step-dot"></span>Uppf&#xF6;ljning</button>
  </div>
  <div class="mode-switch" id="modeSwitch" role="group" aria-label="Vy">
    <button class="mode-btn active" id="mode-stepwise" onclick="setMode('stepwise')" title="Sex steg med bekr&#xE4;ftelse mellan varje">Stegvis</button>
    <button class="mode-btn" id="mode-document" onclick="setMode('document')" title="Allt i ett ark, ingen best&#xE4;md ordning">Arbetsblad</button>
    <button class="mode-btn" id="mode-followup" onclick="setMode('followup')" title="Vad som faktiskt installerades, mot baslinje och plan">Uppf&#xF6;ljning</button>
  </div>
</div>

<!-- Main content -->
<div class="main">
  <!-- Chat panel -->
  <div class="chat-panel">
    <div class="chat-container">
      <div class="chat-header">
        <h2>Aida</h2>
      </div>
      <div class="messages" id="messages">
        <div class="msg bot">Hej! Beskriv ditt projekt. Berätta vad byggnaden används till, byggnadsår, ungefärlig yta och vilka behoven är.</div>
      </div>
      <div class="confirm-bar" id="confirmBar">
        <span class="confirm-bar-text" id="confirmBarText"></span>
        <button class="btn-confirm-sticky" id="confirmBarBtn" onclick="confirmStep()"></button>
      </div>
      <div class="chat-input">
        <input id="userInput" type="text" placeholder="Skriv ditt meddelande..." onkeydown="if(event.key==='Enter')sendMessage()">
        <button id="sendBtn" onclick="sendMessage()" aria-label="Skicka">
          <svg width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><line x1="22" y1="2" x2="11" y2="13"/><polygon points="22 2 15 22 11 13 2 9 22 2"/></svg>
        </button>
      </div>
    </div>
    <div class="chat-disclaimer"></div>
  </div>

  <!-- Results panel -->
  <div class="results-panel" id="results">
    <div class="results-tabs" id="resultTabs" style="display:none">
      <button class="tab" id="tab-projekt" onclick="switchTab('projekt')" disabled>Projekt</button>
      <button class="tab" id="tab-baslinje" onclick="switchTab('baslinje')" disabled>Baslinje</button>
      <button class="tab" id="tab-alternativ" onclick="switchTab('alternativ')" disabled>Alternativ</button>
      <button class="tab" id="tab-rapport" onclick="switchTab('rapport')" disabled>Rapport</button>
    </div>
    <div class="results-content" id="resultContent">
      <div class="empty-state">
        <p>Beskriv ditt projekt i chatten till vänster för att börja.</p>
      </div>
    </div>
  </div>
</div>

<!-- Footer -->
<div class="footer" style="display:flex;justify-content:center;align-items:center;gap:8px">
  <span style="font-size:11px;color:var(--kk-gray-500)">Aida kan g&#xF6;ra misstag. Kontrollera viktig information.</span>
</div>

<!-- Välkomstdialog: visas första besöket, går att kalla tillbaka från Om Aida -->
<div id="welcomeModal" class="modal-backdrop" style="display:none" onclick="if(event.target===this)closeWelcome()" role="dialog" aria-modal="true" aria-labelledby="welcomeTitle">
  <div class="welcome-box">
    <button class="welcome-close" onclick="closeWelcome()" aria-label="Stäng">&#x2715;</button>

    <div class="welcome-intro">
      <div class="welcome-eyebrow">Karlstads kommun &middot; Klimatneutrala Karlstad 2030</div>
      <h2 id="welcomeTitle">Välkommen till Aida</h2>
      <p class="welcome-lead">Ett beslutsstöd som räknar fram klimatpåverkan och kostnad för ett ombyggnadsprojekt medan valen fortfarande går att ändra.</p>
    </div>

    <div class="welcome-why">
      <p>Ombyggnationer och renoveringar orsakar <strong>ungefär lika stora utsläpp som nybyggnationer</strong>, men det är nybyggnationerna som räknas på. Aida finns för att göra beräkningen snabb nog att hinnas med även i ett vanligt ombyggnadsprojekt.</p>
    </div>

    <div class="welcome-body">
      <div class="welcome-label">Vad Aida gör</div>
      <p>Du beskriver projektet med egna ord i chatten. Aida tar fram en klimatbaslinje, letar återbruk, jämför klimatoptimerade nyinköp och rankar alternativen på både utsläpp och kronor. Resultatet blir ett underlag du kan ta med dig i genomförandet av projektet.</p>

      <div class="welcome-label">Så får du bäst resultat</div>
      <ul class="welcome-tips">
        <li>
          <strong>Dela upp stora projekt</strong>
          <span>Kör en analys per byggnadsdel eller etapp. En hel fastighet i en enda beskrivning ger grövre siffror än tre avgränsade analyser.</span>
        </li>
        <li>
          <strong>Beskriv behoven framåt</strong>
          <span>Ju tydligare du berättar vad lokalen ska användas till och vad som ska åtgärdas, desto mer precist blir resultatet. Befintliga material och skick hjälper också. Du behöver inte vara noggrann, Aida visar vilka antaganden den gjort.</span>
        </li>
        <li>
          <strong>Rätta Aida i chatten</strong>
          <span>Stämmer inte ett antagande, skriv det. Du kan justera mitt i flödet och be om en ny beräkning.</span>
        </li>
      </ul>
    </div>

    <div class="welcome-note">Aida är en prototyp. Resultaten är ett underlag för beslut, inte ett klimatbokslut. Metod, datakällor och begränsningar finns under <a href="#" onclick="closeWelcome();openAbout();return false">Om verktyget</a>.</div>

    <div class="welcome-foot">
      <label class="welcome-checkbox" for="welcomeDontShow">
        <input type="checkbox" id="welcomeDontShow">
        <span>Visa inte det här igen</span>
      </label>
      <button class="welcome-start" id="welcomeStartBtn" onclick="closeWelcome()">Kom igång</button>
    </div>
  </div>
</div>

<!-- About modal (Feature 5) -->
<div id="projectMetaModal" class="modal-backdrop" style="display:none" onclick="if(event.target===this)closeProjectMeta()" role="dialog" aria-modal="true" aria-labelledby="projectMetaTitle">
  <div class="modal-box" style="max-width:440px">
    <button class="modal-close" onclick="closeProjectMeta()" aria-label="St&#xE4;ng">&#x2715;</button>
    <h2 id="projectMetaTitle">Fastighet och tidpunkt</h2>
    <p style="margin-bottom:22px">Kopplar analysen till en byggnad och en ungef&#xE4;rlig tidpunkt, s&#xE5; att flera analyser p&#xE5; samma fastighet g&#xE5;r att l&#xE4;sa ihop och planera efter. B&#xE5;da f&#xE4;lten &#xE4;r frivilliga.</p>
    <label class="meta-label" for="metaPropertyRef">Fastighet eller objektsnummer</label>
    <input class="meta-input" id="metaPropertyRef" type="text" maxlength="120" placeholder="T.ex. Stadsbiblioteket">
    <p class="meta-hint">Samma ben&#xE4;mning som i Pythagoras, om du har den.</p>
    <label class="meta-label" for="metaPlannedStart">Planerad utf&#xF6;randetidpunkt</label>
    <input class="meta-input" id="metaPlannedStart" type="month">
    <p class="meta-hint">Ungef&#xE4;rlig m&#xE5;nad r&#xE4;cker. L&#xE4;mna tomt om det inte &#xE4;r best&#xE4;mt &#xE4;n.</p>
    <div class="meta-actions">
      <button class="btn btn-secondary" onclick="closeProjectMeta()">Avbryt</button>
      <button class="btn" onclick="saveProjectMeta()">Spara</button>
    </div>
  </div>
</div>

<div id="aboutModal" class="modal-backdrop" style="display:none" onclick="if(event.target===this)closeAbout()">
  <div class="modal-box">
    <button class="modal-close" onclick="closeAbout()" aria-label="St&#xE4;ng">&#x2715;</button>
    <h2>Om Aida</h2>
    <section>
      <h3>Vad &#xE4;r Aida?</h3>
      <p>Aida &#xE4;r ett AI-drivet beslutsst&#xF6;d f&#xF6;r klimatber&#xE4;kning vid ombyggnation av kommunala fastigheter. Verktyget utvecklas inom Klimatneutrala Karlstad 2030, finansierat av Vinnova, Energimyndigheten och Formas inom ramen f&#xF6;r strategiska innovationsprogrammet Viable Cities.</p>
    </section>
    <section>
      <h3>Datak&#xE4;llor</h3>
      <ul>
        <li><strong>Klimatdata:</strong> Boverkets klimatdatabas, NollCO2-metoden</li>
        <li><strong>Alternativ:</strong> Environdec EPD-databas (verifierade produktdeklarationer)</li>
        <li><strong>Priser:</strong> AI-driven webbs&#xF6;kning mot svenska bygghandlare</li>
        <li><strong>&#xC5;terbruk:</strong> Palats (Karlstads kommuns &#xE5;terbruksplattform)</li>
      </ul>
    </section>
    <section>
      <h3>Metod</h3>
      <p>Aida j&#xE4;mf&#xF6;r konventionella materialval (baslinje) mot klimatoptimerade alternativ med hj&#xE4;lp av verifierade EPD:er. Ber&#xE4;kningarna avser produktskedet (A1-A3) om inget annat anges.</p>
    </section>
    <section>
      <h3>Begr&#xE4;nsningar</h3>
      <ul>
        <li>Resultaten &#xE4;r ett underlag f&#xF6;r beslut, inte ett slutgiltigt klimatbokslut</li>
        <li>Kostnadsuppskattningar baseras p&#xE5; webbs&#xF6;kning &#x2014; inh&#xE4;mta offerter f&#xF6;r exakta v&#xE4;rden</li>
        <li>AI kan g&#xF6;ra fel &#x2014; kontrollera k&#xE4;llh&#xE4;nvisningar vid viktiga beslut</li>
      </ul>
    </section>
    <section>
      <h3>Kontakt</h3>
      <p>Henric Barkman, <a href="mailto:henric.barkman@karlstad.se" style="color:var(--kk-blue)">henric.barkman@karlstad.se</a></p>
    </section>
    <section style="margin-bottom:0">
      <p><a href="#" onclick="closeAbout();openWelcome(true);return false" style="color:var(--kk-dark-red)">Visa introduktionen igen</a></p>
    </section>
  </div>
</div>

{% if has_supabase %}</div><!-- /appContainer -->{% endif %}

<script>
// Configure marked
if (typeof marked !== 'undefined') {
  marked.setOptions({ breaks: true, gfm: true });
}

let _step = 'idle';
let state = {
  project: null, baseline: null, alternatives: null,
  selections: {}, pendingDesc: null, reportMarkdown: null,
  // Orchestration increment 4: the full chat, displayed and remembered. Durable
  // per analysis (analyses.conversation_data). Entries tagged with a `role` are
  // the model's context; see llmHistory().
  conversation: [],
  // Orchestration increment 2: standing user directives that shape alternatives.
  // global = applies to every component ("bara svenska tillverkare"); byComponent
  // = scoped to one component ("tänk bredare på golv"). Persisted per analysis and
  // REPLAYED into every alternatives rerun so refinements stop vanishing.
  directives: {global: [], byComponent: {}},
  // Orchestration increment 3: which alternative the user chose per component,
  // kept apart from the concrete binding in `selections`. Durable across reruns
  // so a choice is only lost when the alternative genuinely stops being offered.
  selectionIntent: {},
  // Which building the analysis is about, and roughly when the work is planned.
  // Both optional, both free of the analysis flow: nothing downstream reads them
  // yet. They are stored so a year of test runs adds up to a usable set instead
  // of a pile of unrelated reports. propertyRef is free text (a Pythagoras
  // designation if the user has one), plannedStart is 'YYYY-MM' in the UI and
  // 'YYYY-MM-01' in the database.
  propertyRef: '',
  plannedStart: '',
  // Which shape the analysis is shown in (orchestration-redesign §12). Not a
  // separate code path: the same state, the same computations, the same chat,
  // rendered by the same section functions in a different arrangement. Rides in
  // project_data like directives do, so it is null until intake succeeds; a mode
  // chosen before the first description lives in memory for that session only.
  mode: 'stepwise',
  // Orchestration §12.5: figures the user has a source for, laid on top of the
  // computed ones. Keyed by component id, then by field:
  //   overrides.c1.baseline_co2e = {value, note, at}
  // Never written into baseline/alternatives/selections. Reruns compute their
  // own figure as usual and this is applied afterwards at read time, which is
  // why an override survives every rerun and why lifting one shows Aida's number
  // again without recomputing anything. Rides in project_data beside `mode`.
  overrides: {},
  // Orchestration §12.6: what was actually installed, keyed by component id.
  //   as_built.c1 = {installed_name, quantity, unit, epd, match_quality, ...}
  // Its own Supabase column rather than a ride in project_data: an analysis
  // started straight in follow-up mode has no project_data until intake runs,
  // and following up a project that was never calculated in Aida is normal.
  as_built: {},
  // The computed outcome table. Derived, never saved, and never computed here:
  // the server owns that arithmetic (see /api/followup), so there is one
  // implementation of the figure that ends up in a klimatredovisning instead of
  // one per language that have to be kept honest about each other.
  followup: null,
  get step() { return _step; },
  set step(v) { _step = v; updatePlaceholder(); },
};
let activeTab = null;

// === Pling när Aida är klar (Johanna feedback punkt 1) ===
//
// A short two-note pling via Web Audio (no asset, works offline). On by default,
// toggleable, persisted. Browsers require a user gesture before audio plays, so
// we unlock the context on the first interaction and reuse it thereafter.
//
// The chime alone only reaches her while she is sitting in front of the tab
// with the sound up. A baseline takes minutes, and minutes is exactly when she
// goes and does something else, so "klart" also has to reach her outside the
// tab: a system notification when the page is away (hidden, or visible on a
// second screen she is not looking at), plus a title badge for the tab she can
// see in her tab strip.
//
// Failures ping too. Walking away and coming back to a stopped spinner with no
// explanation is the worst case, not the one to stay quiet about.
//
// What this does NOT cover: closing the tab. The analysis is still the
// browser's own request, so it dies with the page. That needs the job model
// (design §5b), where the work lives on the server.
let _audioCtx = null;
let soundEnabled = true;
try { soundEnabled = (localStorage.getItem('aida_sound') !== 'off'); } catch (e) {}

function _unlockAudio() {
  try {
    if (!_audioCtx) _audioCtx = new (window.AudioContext || window.webkitAudioContext)();
    if (_audioCtx.state === 'suspended') _audioCtx.resume();
  } catch (e) { /* audio unavailable — stays silent */ }
  document.removeEventListener('pointerdown', _unlockAudio);
  document.removeEventListener('keydown', _unlockAudio);
}
document.addEventListener('pointerdown', _unlockAudio);
document.addEventListener('keydown', _unlockAudio);

function playStepDone(ok) {
  if (!soundEnabled) return;
  try {
    if (!_audioCtx) _audioCtx = new (window.AudioContext || window.webkitAudioContext)();
    const ctx = _audioCtx;
    if (ctx.state === 'suspended') ctx.resume();
    const now = ctx.currentTime;
    // Soft ascending C6 -> E6 when it worked, descending when it did not: a
    // success chime for a failed run tells her the opposite of what happened.
    const notes = (ok === false) ? [[1318.5, 0], [1046.5, 0.12]] : [[1046.5, 0], [1318.5, 0.12]];
    notes.forEach(([freq, t]) => {
      const osc = ctx.createOscillator();
      const gain = ctx.createGain();
      osc.type = 'sine';
      osc.frequency.value = freq;
      gain.gain.setValueAtTime(0.0001, now + t);
      gain.gain.linearRampToValueAtTime(0.16, now + t + 0.02);
      gain.gain.exponentialRampToValueAtTime(0.0001, now + t + 0.25);
      osc.connect(gain); gain.connect(ctx.destination);
      osc.start(now + t); osc.stop(now + t + 0.26);
    });
  } catch (e) { /* audio unavailable — silent */ }
}

function toggleSound() {
  soundEnabled = !soundEnabled;
  try { localStorage.setItem('aida_sound', soundEnabled ? 'on' : 'off'); } catch (e) {}
  updateSoundToggle();
  if (soundEnabled) { playStepDone(true); maybeAskForNotifications(); }  // confirm audibly, offer notices
}
function updateSoundToggle() {
  const el = document.getElementById('soundToggle');
  // One switch for "säg till när du är klar": it gates both the chime and the
  // system notification. Two separate toggles for one wish is a worse model.
  if (el) el.textContent = soundEnabled ? '🔔 Pling på' : '🔕 Pling av';
}

// --- What to say, and whether to say it out loud -----------------------------
// Pure, so the decisions are testable without a browser.

const STEP_NOTICES = {
  intake:             {ok: 'Projektet är inläst',      fail: 'Inläsningen misslyckades'},
  baseline:           {ok: 'Baslinjen är klar',        fail: 'Baslinjeberäkningen misslyckades'},
  alternatives:       {ok: 'Alternativen är klara',    fail: 'Alternativsökningen misslyckades'},
  report:             {ok: 'Rapporten är klar',        fail: 'Rapportgenereringen misslyckades'},
  baseline_rerun:     {ok: 'Baslinjen är omräknad',    fail: 'Omräkningen av baslinjen misslyckades'},
  alternatives_rerun: {ok: 'Alternativen är uppdaterade', fail: 'Uppdateringen av alternativen misslyckades'},
};

function stepNotice(stepKey, ok, detail) {
  const entry = STEP_NOTICES[stepKey];
  const title = entry ? (ok === false ? entry.fail : entry.ok)
                      : (ok === false ? 'Något gick fel' : 'Aida är klar');
  // The body carries the project, because she may have several analyses going
  // and a bare "Baslinjen är klar" does not say which one.
  const parts = [];
  if (detail) parts.push(detail);
  const projectName = (state.project && (state.project.name || state.project.building_type)) || null;
  if (projectName) parts.push(projectName);
  return {title: title, body: parts.join(' · ')};
}

// Hidden covers another tab; unfocused covers a second screen or another app on
// top. Both mean she is not watching, which is the whole point of the pling.
function pageIsAway() {
  const hidden = (typeof document.hidden === 'boolean') ? document.hidden : false;
  const focused = (typeof document.hasFocus === 'function') ? document.hasFocus() : true;
  return hidden || !focused;
}

function shouldRaiseSystemNotification(away, permission, enabled) {
  return !!away && enabled === true && permission === 'granted';
}

// The tab strip is the one place a notification cannot reach but she can still
// glance at. Prefix rather than replace, so the tab is still recognisable.
function badgeTitle(baseTitle, kind) {
  if (kind === 'done') return '✅ ' + baseTitle;
  if (kind === 'error') return '⚠️ ' + baseTitle;
  return baseTitle;
}

const BASE_TITLE = (typeof document !== 'undefined' && document.title) ? document.title : 'Aida';

function _setBadge(kind) {
  try { document.title = badgeTitle(BASE_TITLE, kind); } catch (e) {}
}

// Clear the badge the moment she comes back, not on the next action: a stale
// checkmark on a tab she is already reading is noise.
if (typeof document !== 'undefined' && document.addEventListener) {
  document.addEventListener('visibilitychange', () => { if (!document.hidden) _setBadge(null); });
  window.addEventListener('focus', () => _setBadge(null));
}

function _notificationPermission() {
  try { return (typeof Notification === 'undefined') ? 'unsupported' : Notification.permission; }
  catch (e) { return 'unsupported'; }
}

// Ask from a user gesture (confirmStep / sendMessage call this before their
// first await), once per browser. Asking on page load is the pattern everyone
// has learned to dismiss, and a dismissed prompt is harder to recover from than
// one never shown.
function maybeAskForNotifications() {
  if (_notificationPermission() !== 'default') return;
  if (!soundEnabled) return;
  try {
    if (localStorage.getItem('aida_notify_asked') === '1') return;
    localStorage.setItem('aida_notify_asked', '1');
  } catch (e) { /* private mode: ask this once, do not persist */ }
  addMsg('En analys tar ofta ett par minuter. Säg ja till notiser så säger jag till när den är klar, även om du håller på med något annat.', 'system');
  try { Notification.requestPermission(); } catch (e) { /* older API shape */ }
}

// The one call every finished (or failed) step makes.
function notifyStepDone(stepKey, ok, detail) {
  const succeeded = ok !== false;
  playStepDone(succeeded);
  const away = pageIsAway();
  if (!away) return;
  _setBadge(succeeded ? 'done' : 'error');
  if (!shouldRaiseSystemNotification(away, _notificationPermission(), soundEnabled)) return;
  const notice = stepNotice(stepKey, succeeded, detail);
  try {
    // Same tag for every step: a run that finishes several steps should leave
    // one current notice, not a stack she has to dismiss one by one.
    const n = new Notification(notice.title, {body: notice.body, tag: 'aida-step'});
    n.onclick = () => { try { window.focus(); } catch (e) {} n.close(); };
  } catch (e) { /* notifications unavailable — chime and badge already fired */ }
}

// Dynamic placeholder (Feature 4)
const STEP_PLACEHOLDERS = {
  idle: 'Beskriv ditt ombyggnadsprojekt...',
  intake_done: 'Korrigera eller bekr\u00e4fta...',
  baseline_done: 'Diskutera, korrigera eller bekr\u00e4fta...',
  alternatives_done: 'Diskutera, korrigera eller generera rapport...',
  report_done: 'Diskutera eller korrigera analysen...',
};
function updatePlaceholder() {
  document.getElementById('userInput').placeholder = STEP_PLACEHOLDERS[state.step] || 'Skriv ditt meddelande...';
  updateConfirmBar();
}

// Sticky confirm bar
const CONFIRM_BAR_CONFIG = {
  intake_done: { text: 'Projektbeskrivning klar.', btn: 'Ber\u00e4kna baslinje \u2192' },
  baseline_done: { text: 'Baslinjen \u00e4r klar.', btn: 'S\u00f6k alternativ \u2192' },
  alternatives_done: { text: 'Alternativ redo.', btn: 'Generera rapport \u2192' },
};
function updateConfirmBar() {
  const bar = document.getElementById('confirmBar');
  // The sheet has no gates (§12.1): runs start from the section that is missing,
  // so a bar saying "confirm to continue" would point at a step that isn't there.
  const cfg = isSheet() ? null : CONFIRM_BAR_CONFIG[state.step];
  if (cfg) {
    document.getElementById('confirmBarText').textContent = cfg.text;
    const btn = document.getElementById('confirmBarBtn');
    btn.textContent = cfg.btn;
    btn.disabled = false;
    btn.style.opacity = '';
    bar.classList.add('visible');
  } else {
    bar.classList.remove('visible');
  }
}

function esc(s) { return String(s==null?'':s).replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;').replace(/"/g,'&quot;'); }

function renderMd(text) {
  text = text.replace(/^(\d+)\)\s/gm, '$1. ');
  let html;
  if (typeof marked !== 'undefined') html = marked.parse(text);
  else html = text.replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;')
    .replace(/\*\*(.+?)\*\*/g,'<strong>$1</strong>').replace(/\n/g,'<br>');
  return typeof DOMPurify !== 'undefined' ? DOMPurify.sanitize(html) : html;
}

// === Orchestration increment 4: one durable conversation per analysis ===
//
// Until now the screen and the model remembered different things. What was
// displayed rode in localStorage; what the model had seen lived only in memory
// and died on every reload. So a user could reload, read the whole discussion
// on screen, refer back to it ("gör som jag sa om golvet") and be answered by a
// model that had seen none of it. The transcript lied about what the system
// remembered.
//
// One array now carries both: every entry renders, and the entries tagged with
// a `role` are the model's context (see llmHistory). It persists server-side in
// analyses.conversation_data, so it also survives switching device.
//
// Entry: {text, cls, confirm?{btnLabel,hint}, role?:'user'|'assistant'}

// Bound the persisted row. Trims whole entries from the front (oldest first);
// a report markdown blob in the log can be tens of thousands of characters.
const CONVERSATION_MAX_CHARS = 120000;

function _trimConversation(entries) {
  let total = 0, cut = 0;
  for (let i = entries.length - 1; i >= 0; i--) {
    total += (entries[i].text || '').length + 40;  // + rough per-entry overhead
    if (total > CONVERSATION_MAX_CHARS) { cut = i + 1; break; }
  }
  // A single message over the cap (a pasted report, say) must not take the whole
  // conversation with it: keep the newest entry whole even when it busts the
  // budget alone. An oversized row beats an empty transcript.
  if (cut >= entries.length) cut = entries.length - 1;
  return cut > 0 ? entries.slice(cut) : entries;
}

function _chatStorageKey() { return 'aida_chat_' + (currentAnalysisId || 'new'); }

function _saveConversation() {
  state.conversation = _trimConversation(state.conversation);
  try { localStorage.setItem(_chatStorageKey(), JSON.stringify(state.conversation)); } catch(e) {}
  // Ride to the server only once a row exists. Creating one for a bare greeting
  // would litter the project list with empty analyses; the turns that happened
  // before the row existed are included wholesale when it is first created.
  if (currentAnalysisId) scheduleAutoSave();
}

// The model's view: the entries tagged with a role, in order. Untagged entries
// (system notices, step summaries, and everything saved before increment 4)
// render but stay out of the prompt — exactly what the old in-memory
// chatHistory held.
function llmHistory() {
  return state.conversation
    .filter(m => m.role === 'user' || m.role === 'assistant')
    .map(m => ({role: m.role, content: m.text}));
}

// What to show for the analysis being opened: the server's copy when it has
// one, otherwise this browser's local copy. Local is the fallback for analyses
// saved before increment 4 and for the build without Supabase, never a merge —
// two copies of the same conversation cannot be interleaved without inventing
// an order, and the server's is the one that followed the user here.
// DOM-free, so it can move into the orchestrator with the rest of §3.
function conversationToRestore(fromServer, fromLocal) {
  if (Array.isArray(fromServer) && fromServer.length) return fromServer;
  if (Array.isArray(fromLocal) && fromLocal.length) return fromLocal;
  return [];
}

// A fresh intake reassigns component ids, so earlier turns describe a project
// shape that no longer exists. Drop the model's context without erasing what
// the user can read: the transcript stays, the prompt starts clean.
function forgetLlmContext() {
  state.conversation.forEach(m => { delete m.role; });
  _saveConversation();
}

function addMsg(text, cls, role) {
  const entry = role ? {text, cls, role} : {text, cls};
  state.conversation.push(entry);
  _saveConversation();
  const d = document.createElement('div');
  d.className = 'msg ' + cls;
  if (cls === 'bot' || cls === 'system') { d.innerHTML = renderMd(text); }
  else { d.textContent = text; }
  document.getElementById('messages').appendChild(d);
  d.scrollIntoView({behavior:'smooth'});
  return entry;
}

// The confirm message after intake (Henric 2026-09-08). Intake already writes
// the questions Aida would have liked answered into needs_analysis.would_clarify,
// but they were only rendered in the results panel, where nobody answers them.
// Putting them in the chat turns "vad Aida hade velat veta" into an invitation:
// answer what you know, or go straight on. Capped so it reads as a nudge and
// not as a form, and the button keeps 'baslinje' in its label because the
// reload path matches on that word to decide which confirm is still live.
const INTAKE_CLARIFY_MAX = 3;
function intakeSummary(project) {
  const compList = (project.components || []).map(c => '- ' + c.name + ' (' + c.quantity + ' ' + c.unit + ')').join('\n');
  let text = '**' + project.building_type + '**, ' + project.area_bta + ' m²\n\n**Komponenter:**\n' + compList;
  const na = project.needs_analysis || {};
  const asks = (Array.isArray(na.would_clarify) ? na.would_clarify : [])
    .map(q => String(q || '').trim()).filter(Boolean).slice(0, INTAKE_CLARIFY_MAX);
  let hint;
  if (asks.length) {
    text += '\n\n**Bra att veta innan jag räknar:**\n' + asks.map(q => '- ' + q).join('\n');
    hint = 'Svara i chatten på det du vet, eller gå vidare direkt. Skriv också om något inte stämmer.';
  } else {
    hint = 'Skriv i chatten om något inte stämmer.';
  }
  return {text, btnLabel: 'Bekräfta och beräkna baslinje →', hint};
}

// The confirm block is rendered, not stored, per mode. The conversation entry
// keeps its `confirm` either way, so switching back to Stegvis restores the
// buttons: the mode changes the view, never the data.
function confirmBlockHtml(btnLabel, hint) {
  if (isSheet()) return '';
  return '<div class="confirm-actions"><button class="btn-confirm" onclick="confirmStep()">' + btnLabel + '</button></div>' +
         '<div class="confirm-hint">' + hint + '</div>';
}

function addConfirmMsg(text, btnLabel, hint) {
  state.conversation.push({text, cls: 'bot', confirm: {btnLabel, hint}});
  _saveConversation();
  const d = document.createElement('div');
  d.className = 'msg bot';
  d.innerHTML = renderMd(text) + confirmBlockHtml(btnLabel, hint);
  document.getElementById('messages').appendChild(d);
  d.scrollIntoView({behavior:'smooth'});
}

function removeConfirmButtons() {
  document.querySelectorAll('.confirm-actions').forEach(el => {
    const btn = el.querySelector('.btn-confirm');
    if (btn) {
      btn.disabled = true;
      btn.style.cssText = 'background:var(--kk-gray-200);color:var(--kk-gray-400);cursor:default;pointer-events:none';
      btn.textContent = 'Bekr\u00e4ftad \u2713';
    }
    const hint = el.closest('.msg')?.querySelector('.confirm-hint');
    if (hint) hint.remove();
  });
}

function setProgressStep(name) {
  const order = ['planering','baslinje','aterbruk','nyproduktion','sammanstallning','uppfoljning'];
  const STEP_TAB = {planering:'projekt',baslinje:'baslinje',aterbruk:'alternativ',nyproduktion:'alternativ',sammanstallning:'alternativ',uppfoljning:'rapport'};
  const ni = order.indexOf(name);
  order.forEach((s, i) => {
    const pill = document.getElementById('st-' + s);
    if (!pill) return;
    const isDone = i < ni;
    pill.className = 'step-item' + (isDone ? ' done' : i === ni ? ' active' : '');
    // Only a finished step is somewhere you can go back to, and a disabled
    // button drops out of the tab order, so the rail never offers a keyboard
    // stop that does nothing.
    const tab = STEP_TAB[s];
    pill.disabled = !(isDone && tab);
    pill.onclick = isDone && tab ? () => switchTab(tab) : null;
    pill.title = isDone ? 'G\u00e5 till ' + pill.textContent.trim() : '';
  });
  // Below roughly 900px the six pills are wider than the rail and the track
  // scrolls, with its scrollbar hidden. Nothing would bring the step you are
  // standing on back into view, so on a phone the rail could sit on
  // "Projektbeskrivning" while the run is three steps further on.
  const here = document.getElementById('st-' + name);
  if (here) here.scrollIntoView({block: 'nearest', inline: 'nearest'});
}

// One place resets the rail, called from both the new-project path and the
// analysis-load path. They used to carry their own copies of this loop.
function resetProgressRail() {
  document.querySelectorAll('.step-item').forEach(pill => {
    pill.className = 'step-item';
    pill.disabled = true;
    pill.onclick = null;
    pill.title = '';
  });
}

let _loadingTimer = null;
let _loadingStart = null;
// Whether a pipeline step is running. Editable cells read this: a cell edited
// during a rerun would be overwritten by that rerun's answer without a word
// (orchestration-redesign §12.4, regel 3).
let _runInFlight = false;
function setLoading(on) {
  _runInFlight = on;
  if (typeof setCellsDisabled === 'function') setCellsDisabled(on);
  document.getElementById('sendBtn').disabled = on;
  document.getElementById('userInput').disabled = on;
  // Always clean up previous loading state first (prevents duplicates)
  if (_loadingTimer) { clearInterval(_loadingTimer); _loadingTimer = null; }
  const prev = document.getElementById('typingBubble'); if (prev) prev.remove();
  if (on) {
    _loadingStart = Date.now();
    const el = document.createElement('div');
    el.className = 'msg bot'; el.id = 'typingBubble';
    el.innerHTML = '<div class="typing-indicator"><div class="typing-dot"></div><div class="typing-dot"></div><div class="typing-dot"></div><span class="typing-text" id="typingText">Aida jobbar...</span><span class="elapsed-time" id="elapsedTime"></span></div>';
    document.getElementById('messages').appendChild(el);
    el.scrollIntoView({behavior:'smooth'});
    _loadingTimer = setInterval(() => {
      const s = Math.floor((Date.now() - _loadingStart) / 1000);
      const t = document.getElementById('elapsedTime');
      if (t && s >= 3) t.textContent = s + 's';
      const tx = document.getElementById('typingText');
      if (tx && s >= 5 && !tx.dataset.long) {
        tx.textContent = 'Aida jobbar. Detta kan ta cirka 1-3 minuter.';
        tx.dataset.long = '1';
      }
    }, 1000);
  } else {
    updatePlaceholder();
  }
}

// === Modes (orchestration-redesign §12) ===
//
// A mode is a configuration of the view, not a branch of the pipeline. Labels
// live here and nowhere else, so renaming one is a one-line change; the keys are
// what the rest of the code and the database see. 'followup' arrives in step 5.
const MODES = ['stepwise', 'document', 'followup'];
const MODE_LABELS = {stepwise: 'Stegvis', document: 'Arbetsblad', followup: 'Uppföljning'};
function isDoc() { return state.mode === 'document'; }
function isFollowup() { return state.mode === 'followup'; }
// Both non-stepwise modes are a stacked sheet rather than a tab strip. Almost
// every place that used to ask isDoc() was really asking this, and the two
// questions only stayed the same answer while there were two modes.
function isSheet() { return isDoc() || isFollowup(); }

let _anyTabEnabled = false;

// The strip holds tabs and nothing else, so it can leave the screen entirely
// when there are none to show: before any section exists, and in the sheet
// modes where every section is already on one page. Three call sites used to
// set its display directly and could drift apart; the decision lives here so
// each of them only has to say "re-decide".
function applyTabStripChrome() {
  const strip = document.getElementById('resultTabs');
  if (strip) strip.style.display = (!isSheet() && _anyTabEnabled) ? 'flex' : 'none';
}

// The furthest section that exists, in the order restoreUI already uses. Null
// means the analysis has not started.
function defaultTab() {
  if (state.reportMarkdown) return 'rapport';
  if (state.alternatives) return 'alternativ';
  if (state.baseline) return 'baslinje';
  if (state.project) return 'projekt';
  return null;
}

const STEPWISE_EMPTY = '<div class="empty-state"><p>Beskriv ditt projekt i chatten till vänster för att börja.</p></div>';

function setMode(m) {
  if (MODES.indexOf(m) === -1 || m === state.mode) return;
  state.mode = m;
  applyModeChrome();
  // Leaving the sheet has to go through switchTab, not refreshResults. Two
  // reasons. With no tab current the tab view draws nothing, and the sheet would
  // sit there under Stegvis chrome. And in the sheet switchTab returns before it
  // touches the tab strip, so whichever tab was marked before is still marked -
  // the pipeline may have moved on twice since. Only switchTab sets both the
  // content and the mark, so the strip cannot claim a different section than the
  // one on screen.
  if (isSheet()) {
    refreshResults();
    // Entering follow-up with a project already loaded: the table is empty until
    // the server has computed it, so ask straight away rather than leaving three
    // sections looking broken until the first edit.
    if (isFollowup()) refreshFollowup();
  } else {
    const t = activeTab || defaultTab();
    if (t) switchTab(t);
    else document.getElementById('resultContent').innerHTML = STEPWISE_EMPTY;
  }
  updateConfirmBar();
  scheduleAutoSave();
}

// The chrome that only Stegvis owns: the six-step rail, the tab strip and the
// sticky gate. In the sheet they are not hidden decoration but claims that would
// be false - there is no current step, no current tab, and nothing to confirm.
function applyModeChrome() {
  MODES.forEach(m => {
    const b = document.getElementById('mode-' + m);
    if (b) b.classList.toggle('active', m === state.mode);
  });
  applyTabStripChrome();
  // The rail stays on screen in every mode because it carries the mode switch.
  // Only the steps go: a sheet has no step you are standing on. Same shape as
  // the tab strip's rule, one row up.
  const rail = document.querySelector('.progress-bar');
  if (rail) rail.classList.toggle('steps-hidden', isSheet());
}

// Ask the server for the outcome table and redraw. Called when the mode is
// entered, after every mutation while it is on, and after an analysis loads.
// Not on a timer and not on every render: every input to it changes only
// through a request, so anything more often would be asking the same question
// twice.
let _followupInFlight = false;
let _followupAgain = false;
async function refreshFollowup() {
  if (!state.project) return;
  // A request that arrives mid-flight is remembered, not dropped. Returning here
  // was wrong in the way that is hardest to see: the call succeeds, nothing
  // errors, and the table simply keeps the previous edit's numbers. Binding a
  // declaration and marking the row as reuse in the same second left the outcome
  // reading "Ingen EPD bunden" for a row that had one. Same lesson autoSave
  // already carries: the second caller has newer state, so it must not be the
  // one thrown away.
  if (_followupInFlight) { _followupAgain = true; return; }
  _followupInFlight = true;
  try {
    do {
      _followupAgain = false;
      const r = await authFetch('/api/followup', {method: 'POST', headers: {'Content-Type': 'application/json'},
        body: JSON.stringify({project: state.project, baseline: state.baseline,
                              selections: state.selections, as_built: state.as_built,
                              overrides: state.overrides})});
      const d = await r.json();
      if (d && !d.error) {
        state.followup = d;
        if (isFollowup()) renderSheet();
      }
    } while (_followupAgain);
  } catch (e) {
    // A failed outcome fetch leaves the previous table on screen rather than
    // blanking it. The numbers are stale, not wrong, and the sections around
    // them still work; a thrown error here used to take the whole redraw with it.
    console.warn('uppföljningen kunde inte hämtas', e);
  } finally {
    _followupInFlight = false;
    _followupAgain = false;
  }
}

function sectionExists(tab) {
  if (tab === 'projekt') return !!state.project;
  if (tab === 'baslinje') return !!state.baseline;
  if (tab === 'alternativ') return !!state.alternatives;
  if (tab === 'rapport') return !!state.reportMarkdown;
  return false;
}

// One place that re-renders whatever is currently on screen. Callers that used
// to test activeTab themselves now ask for this instead, so adding a mode does
// not mean finding every re-render site again.
function refreshResults() {
  if (isSheet()) { renderSheet(); return; }
  // A rerun can invalidate the very section being looked at: a full baseline
  // rerun nulls alternatives, the report, and every selection. Rendering nothing
  // would leave the old table sitting there with its radio buttons ticked,
  // describing choices that no longer exist. Fall back to the furthest section
  // that still does. (The old code called the renderer unguarded here, which
  // threw on the null and left the same stale table, minus the fallback.)
  if (!sectionExists(activeTab)) {
    const t = defaultTab();
    if (t) switchTab(t);
    else document.getElementById('resultContent').innerHTML = STEPWISE_EMPTY;
    return;
  }
  if (activeTab === 'projekt') renderProjektContent();
  else if (activeTab === 'baslinje') renderBaslinjeContent();
  else if (activeTab === 'alternativ') renderAlternativContent();
  else if (activeTab === 'rapport') renderRapportContent();
}

// === Tab system ===
function enableTab(name) {
  const tab = document.getElementById('tab-' + name);
  if (tab) tab.disabled = false;
  _anyTabEnabled = true;
  applyTabStripChrome();
}

function switchTab(name) {
  activeTab = name;
  // In the sheet every section is already on screen, so "switch to baslinje"
  // means "take me there". The pipeline calls this when a run finishes, which is
  // exactly when the user wants to be moved to the result.
  if (isSheet()) { renderSheet(); scrollToSection(name); return; }
  document.querySelectorAll('.results-tabs .tab').forEach(t => t.classList.remove('active'));
  const tab = document.getElementById('tab-' + name);
  if (tab) tab.classList.add('active');
  refreshResults();
}

function scrollToSection(key) {
  const el = document.getElementById('sheet-' + key);
  if (!el) return;
  // The move is worth animating: it shows the reader where they were taken. But
  // it is motion nobody asked for, so it obeys the system setting.
  const reduce = window.matchMedia && window.matchMedia('(prefers-reduced-motion: reduce)').matches;
  el.scrollIntoView({behavior: reduce ? 'auto' : 'smooth', block: 'start'});
}

// === Chat input ===
// Note: "kör" tas medvetet bort — det matchade "Kör omräkningen" och triggade
// oavsiktligt nästa steg. "kör vidare" fungerar fortfarande via "vidare".
const ADVANCE_RE = /\b(vidare|nästa|fortsätt|gå vidare|next|confirm|bekräfta)\b/i;
const ADVANCE_EXACT_RE = /^(ok|okej|ja)$/i;
const CORRECTION_RE = /\b(ändra|nej|fel|byt|korrigera|gör om|uppdatera|ta bort|lägg till|ändring|rätta|fixa|nytt? antal|inte \d|ska vara|stämmer inte|borde vara)\b/i;

// Build the context string we feed back into intake when the user makes a
// correction. Includes everything intake might otherwise re-ask about:
// projektnamn, byggnadstyp, area, komponenter, plus a tail of chat history so
// previously-answered clarifications (byggnadsår, krav, omfattning) are not
// re-asked. Reason: intake.py runs stateless on a single string description,
// so anything the model needs must travel in that string.
// Strip delimiter strings from message content so user or model text cannot
// impersonate the structural headers in the prompt below. Without this, a chat
// turn containing "Korrigering från användaren:" would inject a fake
// correction section into the next intake call.
function _scrubCtxDelimiters(s) {
  if (typeof s !== 'string') return '';
  return s
    .replace(/Korrigering från användaren:/gi, '[korrigering]')
    .replace(/Tidigare diskussion i sessionen:/gi, '[diskussion]')
    .replace(/^(Användare|Aida):/gim, '$1​:');  // zero-width space prevents role-line spoofing
}

function buildCorrectionContext(text) {
  // Guard: without core project fields we cannot build a useful ctx; pass the
  // raw correction text so intake at least sees the user's intent without
  // "undefined, undefined m2" garbage.
  if (!state.project || !state.project.building_type || state.project.area_bta == null) {
    return text;
  }
  const compSummary = (state.project.components || []).map(c => c.name + ' (' + c.quantity + ' ' + c.unit + ')').join(', ');
  let ctx = '';
  if (state.project.name) ctx += 'Projektnamn: ' + state.project.name + '. ';
  ctx += state.project.building_type + ', ' + state.project.area_bta + ' m2. Komponenter: ' + compSummary + '.';
  const hist = llmHistory();
  if (hist.length > 0) {
    const tail = hist.slice(-8).map(m => {
      const role = (m.role === 'user') ? 'Användare' : 'Aida';
      return role + ': ' + _scrubCtxDelimiters(m.content || '');
    }).join('\n');
    ctx += '\n\nTidigare diskussion i sessionen:\n' + tail;
  }
  ctx += '\n\nKorrigering från användaren: ' + _scrubCtxDelimiters(text);
  return ctx;
}

async function sendMessage() {
  const input = document.getElementById('userInput');
  const text = input.value.trim();
  if (!text) return;
  input.value = '';
  // Same gesture window as confirmStep: a chat message can kick off a rerun
  // that takes just as long as a first run.
  maybeAskForNotifications();
  // Keep the entry: whether this message becomes a turn the model remembers
  // depends on which branch below handles it (advisory and chat do, intake and
  // "kör vidare" do not — same split the in-memory chatHistory had).
  const userEntry = addMsg(text, 'user');
  setLoading(true);

  // Detect "advance to next step" intent at confirmation gates
  const wantsAdvance = ADVANCE_RE.test(text) || ADVANCE_EXACT_RE.test(text.trim());
  const wantsCorrection = CORRECTION_RE.test(text);

  // Orchestration increment 1: route advisory questions to a server-side answer
  // before the flow switch. This is the missing "question" branch — without it a
  // question in `idle` falls through to runIntake and crashes the intake parser.
  // Skip the router for messages the regexes already classify reliably: advance
  // commands ("kör vidare") and corrections ("byt golvet"). Corrections have a
  // working regex-gated mutation path; letting the LLM router reclassify one as
  // advisory would silently drop the edit (advisory mutates nothing). So the
  // router only fires on messages neither regex catches — exactly the pure
  // questions we want it for. Any routing error falls through (fail-safe).
  if (!wantsAdvance && !wantsCorrection) {
    try {
      const routed = await routeMessage(text);
      if (routed && routed.intent === 'advisory_question') {
        userEntry.role = 'user';
        addMsg(routed.reply, 'bot', 'assistant');  // saves both entries
        setLoading(false);
        return;
      }
    } catch (e) {
      // Router unavailable — proceed with the existing flow rather than block.
    }
  }

  switch (state.step) {
    case 'idle':
      if (state.pendingDesc) {
        await runIntake(state.pendingDesc + '\n\nFörtydligande: ' + text);
      } else {
        await runIntake(text);
      }
      break;
    case 'intake_done':
      if (wantsAdvance) {
        setLoading(false);
        confirmStep();
      } else {
        addMsg('Uppdaterar projektbeskrivning...', 'system');
        await runIntake(buildCorrectionContext(text));
      }
      break;
    case 'baseline_done':
      if (wantsAdvance) {
        setLoading(false);
        confirmStep();
      } else if (wantsCorrection) {
        // Re-run intake with correction, then auto-trigger baseline
        addMsg('Uppdaterar projektet och r\u00e4knar om baslinjen...', 'system');
        await runIntake(buildCorrectionContext(text));
        if (state.step === 'intake_done') {
          await runBaseline();
        }
      } else {
        await runChat(text, userEntry);
      }
      break;
    case 'alternatives_done':
      if (wantsAdvance) {
        const allSel = state.alternatives && state.alternatives.components.every(c => state.selections[c.component_id]);
        if (!allSel) {
          const missing = state.alternatives.components.filter(c => !state.selections[c.component_id]).map(c => c.component_name);
          addMsg('Välj alternativ för alla komponenter först: ' + missing.join(', '), 'system');
          setLoading(false);
          break;
        }
        setLoading(false);
        generateReport();
      } else if (wantsCorrection) {
        // Re-run alternatives with user feedback
        addMsg('G\u00f6r om alternativs\u00f6kningen med dina kommentarer...', 'system');
        await runAlternatives(text);
      } else {
        await runChat(text, userEntry);
      }
      break;
    case 'report_done':
      if (wantsCorrection) {
        // Re-run from intake with correction
        addMsg('G\u00f6r om analysen med dina kommentarer...', 'system');
        await runIntake(buildCorrectionContext(text));
        if (state.step === 'intake_done') {
          await runBaseline();
          if (state.step === 'baseline_done') {
            await runAlternatives();
          }
        }
      } else {
        await runChat(text, userEntry);
      }
      break;
    default:
      setLoading(false);
  }
}

// === Confirm step ===
function confirmStep() {
  // Synchronously, before anything async: browsers only honour a permission
  // request inside the user gesture that triggered it, and an await ends that
  // window. This is the click that starts the minutes-long work, so it is also
  // the moment where "säg till när du är klar" makes sense to her.
  maybeAskForNotifications();
  document.getElementById('confirmBarBtn').disabled = true;
  document.getElementById('confirmBarBtn').style.opacity = '0.5';
  if (state.step === 'intake_done') runBaseline();
  else if (state.step === 'baseline_done') runAlternatives();
  else if (state.step === 'alternatives_done') {
    const allSel = state.alternatives && state.alternatives.components.every(c => state.selections[c.component_id]);
    if (!allSel) {
      const missing = state.alternatives.components.filter(c => !state.selections[c.component_id]).map(c => c.component_name);
      addMsg('Välj alternativ för alla komponenter först: ' + missing.join(', '), 'system');
      document.getElementById('confirmBarBtn').disabled = false;
      document.getElementById('confirmBarBtn').style.opacity = '';
      return;
    }
    generateReport();
  }
}

// === Pipeline: Intake ===
async function runIntake(desc) {
  addMsg('Analyserar projektbeskrivning...', 'system');
  setProgressStep('planering');
  try {
    const r = await authFetch('/api/intake', {method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({description: desc})});
    const d = await r.json();
    if (d.error) { addMsg('Fel: ' + d.error, 'system'); setLoading(false); return; }

    if (d.clarification_needed) {
      state.pendingDesc = desc;
      state.project = null;
      state.step = 'idle';
      if (d.components && d.components.length) {
        const list = d.components.map(c => '- ' + c.name).join('\n');
        addMsg('Hittade hittills:\n' + list, 'bot');
      }
      addMsg(d.clarification_needed, 'bot');
      setLoading(false);
      return;
    }

    state.pendingDesc = null;
    state.project = d;
    state.baseline = null;
    state.alternatives = null;
    state.selections = {};
    state.reportMarkdown = null;
    forgetLlmContext();
    // Intake assigns fresh component IDs, so component-scoped directives keyed on
    // the old IDs are now orphans — drop them. Global directives are not ID-bound
    // and survive (e.g. "bara svenska tillverkare").
    _ensureDirectives();
    state.directives.byComponent = {};
    // Same reason, same fix for selection intent: it is keyed on component id.
    state.selectionIntent = {};
    // And for overrides, which are keyed the same way. A manual figure that
    // survived a re-intake would attach itself to whatever component happened to
    // get that id next, which is the one failure this feature must not have.
    state.overrides = {};
    // And what was registered as installed, keyed the same way. Re-running intake
    // on a project someone had already followed up is rare, but the failure it
    // would leave behind is a delivered quantity filed under the wrong component.
    state.as_built = {};
    state.followup = null;
    state.step = 'intake_done';
    notifyStepDone('intake', true);
    if (HAS_SUPABASE) { document.getElementById('projectName').textContent = d.name || d.building_type || 'Nytt projekt'; }
    scheduleAutoSave();

    enableTab('projekt');
    switchTab('projekt');
    // Disable later tabs if re-running
    ['baslinje','alternativ','rapport'].forEach(t => { const el = document.getElementById('tab-'+t); if(el) el.disabled = true; });

    const summary = intakeSummary(d);
    addConfirmMsg(summary.text, summary.btnLabel, summary.hint);
    setLoading(false);
  } catch(e) { addMsg('Fel: ' + e.message, 'system'); setLoading(false); }
}

// === Pipeline: Baseline ===
// Remove any lingering recompute action-row buttons. Used when we start a
// fresh baseline/alternatives run so an orphaned "S\u00f6k nya alternativ" button
// (rendered from a stale state earlier) does not stay clickable after we have
// already cleared state.alternatives.
function clearActionRows() {
  document.querySelectorAll('.msg.action-row').forEach(el => el.remove());
}

async function runBaseline() {
  addMsg('Ber\u00e4knar baslinje (NollCO2)...', 'system');
  setProgressStep('baslinje');
  setLoading(true);
  clearActionRows();
  try {
    const r = await authFetch('/api/baseline', {method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({project: state.project})});
    const d = await r.json();
    if (d.error) {
      addMsg('Fel: ' + d.error, 'system');
      addConfirmMsg('Baslinjeberäkning misslyckades.', 'Försök igen \u2192', '');
      notifyStepDone('baseline', false);
      setLoading(false); return;
    }
    removeConfirmButtons();
    // Same as the scoped path: the binding cannot outlive an emptied
    // alternatives bag, but the intent carries to the next alternatives run
    // (increment 3). Recomputing a baseline should not cost you your choices.
    backfillIntent();
    state.baseline = d;
    state.alternatives = null;
    state.selections = {};
    state.reportMarkdown = null;
    state.step = 'baseline_done';
    notifyStepDone('baseline', true);
    scheduleAutoSave();

    enableTab('baslinje');
    switchTab('baslinje');
    ['alternativ','rapport'].forEach(t => { const el = document.getElementById('tab-'+t); if(el) el.disabled = true; });

    addConfirmMsg(baselineDoneMsg(), 'Bekr\u00e4fta och s\u00f6k alternativ \u2192',
                  'Skriv i chatten om du vill korrigera n\u00e5got.');
    setLoading(false);
  } catch(e) {
    addMsg('Fel: ' + e.message, 'system');
    addConfirmMsg('Baslinjeberäkning misslyckades.', 'Försök igen \u2192', '');
    notifyStepDone('baseline', false);
    setLoading(false);
  }
}

// === Orchestration increment 2: standing directives ===
// Directives are user instructions that shape alternatives ("tänk bredare",
// "bara svenska tillverkare"). They are stored durably and replayed into EVERY
// alternatives rerun for their scope, so a later rerun (triggered by anything)
// does not silently revert to defaults — the matsal bug.
function _ensureDirectives() {
  if (!state.directives) state.directives = {global: [], byComponent: {}};
  if (!state.directives.global) state.directives.global = [];
  if (!state.directives.byComponent) state.directives.byComponent = {};
}
// Record a new directive. Empty scopeCids = global; otherwise scoped per component.
// Returns true if it was newly added (vs already standing).
function addDirective(scopeCids, text) {
  text = (text || '').trim();
  if (!text) return false;
  _ensureDirectives();
  let added = false;
  if (!scopeCids || scopeCids.length === 0) {
    if (!state.directives.global.includes(text)) { state.directives.global.push(text); added = true; }
  } else {
    scopeCids.forEach(cid => {
      if (!state.directives.byComponent[cid]) state.directives.byComponent[cid] = [];
      if (!state.directives.byComponent[cid].includes(text)) { state.directives.byComponent[cid].push(text); added = true; }
    });
  }
  return added;
}
// Standing directives applicable to a rerun scope. Scoped rerun = global + that
// component's directives. Full rerun (empty scope) = GLOBAL ONLY: a component
// wish ("tänk bredare på golv") must never leak onto other components' prompts.
// Per-component directives still live in state and replay on that component's
// own next scoped rerun.
function directivesForScope(scopeCids) {
  _ensureDirectives();
  const out = state.directives.global.slice();
  if (scopeCids && scopeCids.length) {
    scopeCids.forEach(cid => (state.directives.byComponent[cid] || []).forEach(d => { if (!out.includes(d)) out.push(d); }));
  }
  return out;
}
// Persist newFeedback as a standing directive, then return the FULL standing set
// for this scope as a single user_feedback string. This is the replay: the alt
// search always sees every accumulated instruction, not just the latest one.
function effectiveFeedback(scopeCids, newFeedback) {
  if (newFeedback) addDirective(scopeCids, newFeedback);
  let all = directivesForScope(scopeCids);
  // Cap so a long session cannot bloat the alternatives prompt (server's chat
  // path caps at 500; this direct path had none). Always keep global directives;
  // fill remaining budget with the most recent scoped ones.
  const CAP = 800;
  if (all.join('; ').length > CAP) {
    const globals = state.directives.global.filter(d => all.includes(d));
    const rest = all.filter(d => !globals.includes(d));
    const kept = [];
    for (let i = rest.length - 1; i >= 0; i--) {
      if ([...globals, rest[i], ...kept].join('; ').length > CAP) break;
      kept.unshift(rest[i]);
    }
    all = [...globals, ...kept];
  }
  return all.length ? all.join('; ') : null;
}
// Drop directives for removed components so they cannot resurrect a deleted scope.
function pruneDirectives(removedCids) {
  _ensureDirectives();
  (removedCids || []).forEach(cid => { delete state.directives.byComponent[cid]; });
}

// === Orchestration increment 3: selection intent + status-based merge ===
// Same split the directives got: what the user chose is durable, what it points
// at right now is derived. A rerun replaces the alternatives bag for a
// component, and the old code dropped the selection whenever that happened —
// even when the rerun was a superset and the chosen alternative was still in
// the list. Intent lives in state.selectionIntent and survives the gaps
// (baseline rerun, empty alternatives); the binding in state.selections is
// recomputed after every merge by reconcileSelections().
function _ensureIntent() {
  if (!state.selectionIntent) state.selectionIntent = {};
}

// The model rewrites product names between runs — an en dash was enough to
// break matching in PR #557. Normalise dashes, symbols, spacing and case.
// Cosmetics only: see reconcileSelections for why this stops short of fuzzy.
function normAltName(name) {
  return String(name === undefined || name === null ? '' : name)
    .replace(/[‐-―]/g, '-')
    .replace(/[®™©]/g, '')
    .replace(/\s+/g, ' ')
    .trim()
    .toLowerCase();
}

// Record a choice as intent. Called wherever a selection is written, so the two
// never drift apart.
function rememberIntent(cid, componentName, altName) {
  _ensureIntent();
  state.selectionIntent[cid] = {componentName: componentName, altName: altName};
}

// Adopt intent for any selection that has none — loaded analyses saved before
// increment 3, and paths that write state.selections wholesale (the chat agent).
function backfillIntent() {
  _ensureIntent();
  Object.keys(state.selections || {}).forEach(cid => {
    if (state.selectionIntent[cid]) return;
    const sel = state.selections[cid];
    if (sel && sel.selected_alternative) {
      rememberIntent(cid, sel.name, sel.selected_alternative.name);
    }
  });
}

// Intent dies with its component, nothing else.
function dropIntent(removedCids) {
  _ensureIntent();
  (removedCids || []).forEach(cid => { delete state.selectionIntent[cid]; });
}

// Build a fresh selection object for an intent, or null when the alternative is
// no longer on offer. Always reads the CURRENT numbers: a rerun can change kg
// and price, and carrying the old object forward would leave a stale figure in
// the summary and in the report — a quieter bug than the one being fixed.
function bindIntent(comp, altName) {
  if (!comp) return null;
  if (altName === 'Baslinje') {
    return {id: comp.component_id, name: comp.component_name,
      selected_alternative: {name: 'Baslinje', co2e_kg: comp.baseline_co2e_kg,
        cost_sek: comp.baseline_cost_sek, source: 'NollCO2'},
      baseline_co2e_kg: comp.baseline_co2e_kg, baseline_cost_sek: comp.baseline_cost_sek};
  }
  const want = normAltName(altName);
  const alt = (comp.alternatives || []).find(a => normAltName(a.name) === want);
  if (!alt) return null;
  return {id: comp.component_id, name: comp.component_name,
    selected_alternative: {name: alt.name, co2e_kg: alt.co2e_kg, cost_sek: alt.cost_sek,
      source: alt.source,
      available_quantity: (alt.available_quantity === undefined ? null : alt.available_quantity),
      price_basis: alt.price_basis || '', gwp_basis: alt.gwp_basis || ''},
    baseline_co2e_kg: comp.baseline_co2e_kg, baseline_cost_sek: comp.baseline_cost_sek};
}

// The one owner of "does this selection survive a rerun". Runs AFTER the
// alternatives merge, never before — the old order deleted the selection in the
// baseline step and left nothing for the alternatives step to pick back up.
// `cids` scopes it; empty/null means every component carrying intent.
// Matching is normalised-exact and deliberately not fuzzy: binding a choice to
// the WRONG product is worse than dropping it, because a dropped binding shows
// up in the UI and costs a click, while a wrong one travels into the report
// unseen. Same direction as #558 — rigged against false matches, not for cover.
// Returns {kept, lost} for the caller to report; mutates no DOM, so it can move
// into the server-side orchestrator later unchanged.
function reconcileSelections(cids) {
  _ensureIntent();
  if (!state.selections) state.selections = {};
  const comps = (state.alternatives && Array.isArray(state.alternatives.components))
    ? state.alternatives.components : [];
  const byId = new Map();
  const byName = new Map();
  comps.forEach(c => {
    byId.set(c.component_id, c);
    if (!byName.has(c.component_name)) byName.set(c.component_name, c);
  });
  const scope = (cids && cids.length) ? cids.slice() : Object.keys(state.selectionIntent);
  const kept = [], lost = [];
  scope.forEach(cid => {
    const intent = state.selectionIntent[cid];
    if (!intent) return;
    // A full rerun can hand the same component back under a fresh id, which is
    // why the original restore matched on name. Follow the name before
    // concluding the component is gone.
    const comp = byId.get(cid) || byName.get(intent.componentName);
    if (!comp) {
      // No alternatives for it yet. Drop the binding, keep the intent waiting.
      delete state.selections[cid];
      return;
    }
    let key = cid;
    if (comp.component_id !== cid) {
      key = comp.component_id;
      delete state.selections[cid];
      delete state.selectionIntent[cid];
      state.selectionIntent[key] = intent;
    }
    const bound = bindIntent(comp, intent.altName);
    if (bound) {
      state.selections[key] = bound;
      kept.push(comp.component_name);
    } else {
      // Intent stays: if a later, broader rerun surfaces the same product, the
      // choice rebinds itself. That is what durable intent means.
      delete state.selections[key];
      lost.push({component: comp.component_name, alt: intent.altName});
    }
  });
  return {kept: kept, lost: lost};
}

// Say something only when a choice could NOT be carried over. Silence means it
// worked. Johanna's complaint was that the loss happened invisibly, so the loss
// is the part that has to speak.
function reportReconcile(res) {
  if (!res || !res.lost || !res.lost.length) return;
  res.lost.forEach(l => {
    addMsg('Ditt val "' + l.alt + '" för ' + l.component
      + ' finns inte kvar efter omkörningen. Välj ett nytt alternativ i resultatpanelen.', 'system');
  });
}

// === Pipeline: Alternatives ===
async function runAlternatives(userFeedback) {
  // Selections made before intent tracking existed become intent now, so the
  // reconcile below has something to work from (increment 3).
  backfillIntent();

  addMsg('S\u00f6ker alternativ...', 'system');
  setProgressStep('aterbruk');
  setLoading(true);
  clearActionRows();
  const subStepTimer = setTimeout(() => {
    setProgressStep('nyproduktion');
  }, 2000);
  try {
    const body = {project: state.project, baseline: state.baseline};
    // Full rerun: replay all standing directives (global + every component's).
    const fb = effectiveFeedback([], userFeedback);
    if (fb) body.user_feedback = fb;
    const r = await authFetch('/api/alternatives', {method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify(body)});
    const d = await r.json();
    clearTimeout(subStepTimer);
    if (d.error) {
      clearTimeout(subStepTimer);
      addMsg('Fel: ' + d.error, 'system');
      addConfirmMsg('S\u00f6kning av alternativ misslyckades.', 'F\u00f6rs\u00f6k igen \u2192', '');
      notifyStepDone('alternatives', false);
      setLoading(false); return;
    }
    removeConfirmButtons();
    state.alternatives = d;
    state.reportMarkdown = null;
    state.step = 'alternatives_done';
    notifyStepDone('alternatives', true);

    // Rebind choices against the new bag (increment 3). Replaces the bespoke
    // name-match restore this function used to own alone — the chat-driven
    // reruns below now go through the same rules, which is where the old
    // asymmetry lived: buttons kept your choice, chat threw it away.
    reportReconcile(reconcileSelections(null));
    scheduleAutoSave();
    setProgressStep('sammanstallning');

    enableTab('alternativ');
    switchTab('alternativ');
    document.getElementById('tab-rapport').disabled = true;

    const commentary = d.commentary || '';
    if (commentary) {
      addMsg(commentary, 'bot');
      addMsg('V\u00e4lj alternativ per komponent i resultatpanelen. Skriv i chatten om du vill ha fler f\u00f6rslag.', 'bot');
    } else {
      addMsg('Alternativ klara! V\u00e4lj per komponent i resultatpanelen.\n\nSkriv i chatten om du vill ha fler alternativ.', 'bot');
    }
    setLoading(false);
  } catch(e) {
    clearTimeout(subStepTimer);
    addMsg('Fel: ' + e.message, 'system');
    addConfirmMsg('S\u00f6kning av alternativ misslyckades.', 'F\u00f6rs\u00f6k igen \u2192', '');
    notifyStepDone('alternatives', false);
    setLoading(false);
  }
}

// === Pipeline: Report ===
async function generateReport() {
  setProgressStep('uppfoljning');
  addMsg('Genererar rapport...', 'system');
  // #reportBtn only exists while the Alternativ tab is rendered, but the sticky
  // confirm bar can trigger this from any tab — null-guard to avoid a throw
  // that leaves the UI stuck (the error paths below already guard the same way).
  const rb0 = document.getElementById('reportBtn'); if (rb0) rb0.disabled = true;
  setLoading(true);
  try {
    const sels = {components: Object.values(state.selections)};
    // The raw selections plus the overrides, not the overridden ones: the
    // server lays them on with the same function the report's appendix is
    // built from, so the totals and the list of manual figures cannot
    // disagree about which numbers were replaced.
    const r = await authFetch('/api/report', {method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({project: state.project, selections: sels, overrides: state.overrides})});
    const d = await r.json();
    if (d.error) {
      addMsg('Fel: ' + d.error, 'system');
      addConfirmMsg('Rapportgenerering misslyckades.', 'Försök igen →', '');
      notifyStepDone('report', false);
      const rb = document.getElementById('reportBtn'); if (rb) rb.disabled = false;
      setLoading(false);
      return;
    }
    state.reportMarkdown = d.markdown;
    state.step = 'report_done';
    notifyStepDone('report', true);
    scheduleAutoSave();
    addMsg('Rapport klar!', 'bot');
    enableTab('rapport');
    switchTab('rapport');
    setLoading(false);
  } catch(e) {
    addMsg('Fel: ' + e.message, 'system');
    addConfirmMsg('Rapportgenerering misslyckades.', 'Försök igen →', '');
    notifyStepDone('report', false);
    const rb = document.getElementById('reportBtn'); if (rb) rb.disabled = false;
    setLoading(false);
  }
}

// The follow-up's own report. Written from the outcome table rather than from
// the plan, so it does not go through generateReport: that one refuses when no
// alternative has been chosen, which is the normal state of a project that was
// followed up without ever being planned in Aida.
async function generateFollowupReport() {
  addMsg('Skriver klimatredovisningen...', 'system');
  setLoading(true);
  try {
    const r = await authFetch('/api/followup/report', {
      method: 'POST', headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({project: state.project, baseline: state.baseline,
                            selections: state.selections, as_built: state.as_built,
                            overrides: state.overrides, property_ref: state.propertyRef || ''})});
    const d = await r.json();
    if (d.error) { addMsg('Fel: ' + d.error, 'system'); setLoading(false); return; }
    state.reportMarkdown = d.markdown;
    scheduleAutoSave();
    // Deliberately not state.step = 'report_done'. The step machine describes
    // how far the six-step analysis has run, and a follow-up on a project that
    // was never analysed here has not run any of it. Writing the redovisning
    // would otherwise light up a progress bar for work nobody did.
    addMsg('Klimatredovisningen är klar. Den ligger under Klimatredovisning här till höger.', 'bot');
    refreshResults();
    recordFollowupFacts();
    setLoading(false);
  } catch(e) {
    addMsg('Fel: ' + e.message, 'system');
    setLoading(false);
  }
}

// What the follow-up learned: estimated against actual, per component and field.
// Written when the redovisning is, because that is the point at which the numbers
// stop moving.
//
// Deliberately silent, in both directions. It is bookkeeping for a decision two
// years out (§12.6: nothing here feeds back into pricing yet), so a failure must
// not land on top of a document the user just successfully produced, and a
// success is not news either.
async function recordFollowupFacts() {
  if (!HAS_SUPABASE || !currentUser || !currentAnalysisId) return;
  try {
    await authFetch('/api/followup/facts', {
      method: 'POST', headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({analysis_id: currentAnalysisId, project: state.project,
                            baseline: state.baseline, selections: state.selections,
                            as_built: state.as_built, overrides: state.overrides})});
  } catch (e) {
    console.warn('uppföljningsfakta kunde inte sparas', e);
  }
}

// Orchestration increment 1: classify a message server-side, getting an advisory
// answer back in the same call when applicable. Mirrors runChat's body shape.
async function routeMessage(text) {
  const body = {
    message: text,
    history: llmHistory().slice(-10),
    project: state.project || null,
    baseline: state.baseline || null,
    alternatives: state.alternatives || null,
    selections: (state.selections && Object.keys(state.selections).length) ? state.selections : null,
  };
  const r = await authFetch('/api/route', {method:'POST', headers:{'Content-Type':'application/json'},
    body: JSON.stringify(body)});
  const d = await r.json();
  if (d.error) throw new Error(d.error);
  return d;
}

// === Conversational chat (agent with tool-use) ===
async function runChat(text, userEntry) {
  setLoading(true);
  // Tag before building the body, so the current message rides in `history` the
  // way it always has. run_chat_agent drops that trailing user turn before it
  // appends `message`, so the model still sees it once.
  // Every caller displays the message first and hands us its entry. Warn rather
  // than push a second one: a duplicate would render twice after a reload.
  if (userEntry) userEntry.role = 'user';
  else console.warn('runChat utan conversation-entry: turen sparas inte i modellens kontext');
  try {
    const body = {
      message: text,
      history: llmHistory().slice(-10),
      project: state.project || null,
      baseline: state.baseline || null,
      alternatives: state.alternatives || null,
      selections: (state.selections && Object.keys(state.selections).length) ? state.selections : null,
      // So Aida talks about the figures the user is looking at, and so a
      // material change made in the chat drops a stale manual figure exactly
      // as the same change made in a cell does.
      overrides: state.overrides,
    };
    const r = await authFetch('/api/chat', {method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify(body)});
    const d = await r.json();
    if (d.error) { addMsg('Fel: ' + d.error, 'system'); setLoading(false); return; }

    applyAgentStateUpdates(d.state_updates);

    addMsg(d.reply, 'bot', 'assistant');

    // Chat agent may have requested baseline/alternatives reruns. Execute them
    // sequentially so each action sees state from the previous one's merge.
    // Keep loading true throughout so the user cannot send a second message
    // mid-rerun and race the state.baseline / state.alternatives merge.
    const pendingActions = d.state_updates && d.state_updates.pending_actions;
    if (Array.isArray(pendingActions) && pendingActions.length > 0) {
      await processPendingActions(pendingActions);
    }
    setLoading(false);
  } catch(e) { addMsg('Fel: ' + e.message, 'system'); setLoading(false); }
}

// Execute reruns requested by the chat agent. Full reruns (empty component_ids)
// require an extra click; partial reruns run immediately since the user already
// initiated them via the chat correction. We pass orchestrated=true so the
// inner rerun functions do not toggle setLoading per action — runChat already
// holds setLoading(true) for the whole sequence to prevent the user from
// sending a second chat message that would race the state merges.
//
// Defense in depth against the prompt-only confirmation gate: if the agent
// emits explicit component_ids that cover every component in the project, we
// re-classify the action as full and route it to the confirmation button. The
// LLM cannot bypass the confirm gate by spelling out every id.
async function processPendingActions(actions) {
  const totalComponents = (state.project && Array.isArray(state.project.components))
    ? state.project.components.length : 0;
  const knownIds = new Set(
    (state.project && Array.isArray(state.project.components))
      ? state.project.components.map(c => c.id) : []
  );

  // Sort: rerun_baseline before rerun_alternatives for the same scope, so the
  // alternatives call sees the freshly recomputed baseline. The LLM can emit
  // them in any order and nothing in the schema enforces it.
  const ordered = [...actions].sort((a, b) => {
    if (a.type === b.type) return 0;
    if (a.type === 'rerun_baseline') return -1;
    if (b.type === 'rerun_baseline') return 1;
    return 0;
  });

  for (const action of ordered) {
    try {
      let cids = Array.isArray(action.component_ids) ? action.component_ids.filter(c => knownIds.has(c)) : [];
      // Empty after filter (unknown ids) and not originally full means we have
      // nothing actionable. Skip rather than treat as "all".
      const originallyEmpty = !Array.isArray(action.component_ids) || action.component_ids.length === 0;
      if (!originallyEmpty && cids.length === 0) {
        addMsg('Hoppar över ' + action.type + ': inga giltiga komponent-id.', 'system');
        continue;
      }
      // Re-classify as full when explicit list covers every component.
      const coversAll = totalComponents > 0 && cids.length === totalComponents;
      const isFull = originallyEmpty || coversAll;

      if (action.type === 'rerun_baseline') {
        if (isFull) {
          renderActionRow(
            'Bekräfta: räkna om hela baslinjen' + (action.reason ? ' (' + action.reason + ')' : ''),
            async () => { await runBaselineForComponents([], action.reason, false); },
          );
        } else {
          await runBaselineForComponents(cids, action.reason, true);
        }
      } else if (action.type === 'rerun_alternatives') {
        if (isFull) {
          renderActionRow(
            'Bekräfta: kör om alla alternativ' + (action.reason ? ' (' + action.reason + ')' : ''),
            async () => { await runAlternativesForComponents([], action.user_feedback, action.reason, false); },
          );
        } else {
          await runAlternativesForComponents(cids, action.user_feedback, action.reason, true);
        }
      } else {
        addMsg('Okänd åtgärd från chatten: ' + action.type, 'system');
      }
    } catch (e) {
      addMsg('Fel vid ' + action.type + ': ' + e.message, 'system');
    }
  }
}

// Partial baseline rerun. `componentIds` empty = full rerun (same outcome as
// runBaseline but without the pipeline side effects like tab switching).
// `orchestrated=true` means runChat already holds setLoading(true) for the
// whole pending_actions sequence — do not toggle it per call, otherwise the
// user can fire a second chat message in the gap between two reruns.
async function runBaselineForComponents(componentIds, reason, orchestrated) {
  if (!state.project) {
    addMsg('Inget projekt att räkna baslinje på.', 'system');
    return;
  }
  const isFull = !componentIds || componentIds.length === 0;
  const scope = isFull
    ? 'hela baslinjen'
    : 'komponent(er) ' + componentIds.join(', ');
  const reasonNote = reason ? ' (' + reason + ')' : '';
  addMsg('Aida räknar om ' + scope + reasonNote + '...', 'system');
  if (!orchestrated) setLoading(true);
  try {
    const body = {project: state.project};
    if (!isFull) body.component_ids = componentIds;
    const r = await authFetch('/api/baseline', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify(body),
    });
    const d = await r.json();
    if (d.error) {
      addMsg('Fel vid baslinjebberäkning: ' + d.error, 'system');
      notifyStepDone('baseline_rerun', false);
      // Respect orchestration: runChat holds the lock for the whole pending
      // sequence. Unlocking here would let a user message race the merges.
      if (!orchestrated) setLoading(false);
      return;
    }
    backfillIntent();
    if (isFull) {
      state.baseline = d;
      state.alternatives = null;
      // Bindings cannot survive an emptied alternatives bag, but the intent
      // does: the alternatives rerun that follows rebinds whatever is still
      // on offer (increment 3).
      state.selections = {};
      state.reportMarkdown = null;
    } else {
      mergeBaselineDelta(d, new Set(componentIds));
      invalidateDownstreamFor(new Set(componentIds));
    }
    scheduleAutoSave();
    refreshResults();
    addMsg('Baslinje uppdaterad' + (isFull ? '' : ' för ' + componentIds.join(', ')) + '.', 'system');
    notifyStepDone('baseline_rerun', true, isFull ? null : componentIds.join(', '));
  } catch (e) {
    addMsg('Fel vid baslinjebberäkning: ' + e.message, 'system');
    notifyStepDone('baseline_rerun', false);
  } finally {
    if (!orchestrated) setLoading(false);
  }
}

async function runAlternativesForComponents(componentIds, userFeedback, reason, orchestrated) {
  if (!state.project || !state.baseline) {
    addMsg('Saknar projekt eller baslinje för alternativ.', 'system');
    return;
  }
  const isFull = !componentIds || componentIds.length === 0;
  const scope = isFull
    ? 'alla alternativ'
    : 'alternativ för komponent(er) ' + componentIds.join(', ');
  const reasonNote = reason ? ' (' + reason + ')' : '';
  const feedbackNote = userFeedback ? ' Önskemål (sparas som stående): ' + userFeedback + '.' : '';
  addMsg('Aida kör om ' + scope + reasonNote + feedbackNote + '...', 'system');
  if (!orchestrated) setLoading(true);
  try {
    const body = {project: state.project, baseline: state.baseline};
    if (!isFull) body.component_ids = componentIds;
    // Replay standing directives for this scope (and persist any new feedback as
    // one). So a rerun triggered by anything — even a material change with no new
    // feedback — still carries "tänk bredare", "bara svenska" etc.
    const fb = effectiveFeedback(isFull ? [] : componentIds, userFeedback);
    if (fb) body.user_feedback = fb;
    const r = await authFetch('/api/alternatives', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify(body),
    });
    const d = await r.json();
    if (d.error) {
      addMsg('Fel vid alternativsökning: ' + d.error, 'system');
      notifyStepDone('alternatives_rerun', false);
      // Respect orchestration (see runBaselineForComponents).
      if (!orchestrated) setLoading(false);
      return;
    }
    backfillIntent();
    if (isFull) {
      state.alternatives = d;
    } else {
      mergeAlternativesDelta(d, new Set(componentIds));
    }
    // Rebind rather than invalidate (increment 3). A "tänk bredare"-rerun is a
    // superset, so the choice is usually still right there; dropping it blindly
    // is Johanna's punkt 5. Runs after the merge so it sees the final bag.
    reportReconcile(reconcileSelections(isFull ? null : componentIds));
    state.reportMarkdown = null;
    scheduleAutoSave();
    refreshResults();
    addMsg('Alternativ uppdaterade' + (isFull ? '' : ' för ' + componentIds.join(', ')) + '.', 'system');
    notifyStepDone('alternatives_rerun', true, isFull ? null : componentIds.join(', '));
  } catch (e) {
    addMsg('Fel vid alternativsökning: ' + e.message, 'system');
    notifyStepDone('alternatives_rerun', false);
  } finally {
    if (!orchestrated) setLoading(false);
  }
}

// Replace baseline entries in place, filtered to the components actually
// requested. cidSet acts as a defensive whitelist: a server returning more
// components than requested (regression or future code path) cannot silently
// overwrite unrelated state.
function mergeBaselineDelta(delta, cidSet) {
  if (!delta || !Array.isArray(delta.components)) return;
  const allowed = (delta.components || []).filter(c => cidSet.has(c.component_id));
  if (!state.baseline || !Array.isArray(state.baseline.components)) {
    state.baseline = {components: allowed};
    return;
  }
  const newById = new Map(allowed.map(c => [c.component_id, c]));
  state.baseline.components = state.baseline.components.map(c => newById.get(c.component_id) || c);
  for (const c of allowed) {
    if (!state.baseline.components.find(x => x.component_id === c.component_id)) {
      state.baseline.components.push(c);
    }
  }
}

function mergeAlternativesDelta(delta, cidSet) {
  if (!delta || !Array.isArray(delta.components)) return;
  const allowed = (delta.components || []).filter(c => cidSet.has(c.component_id));
  if (!state.alternatives || !Array.isArray(state.alternatives.components)) {
    state.alternatives = {components: allowed, commentary: delta.commentary || ''};
    return;
  }
  const newById = new Map(allowed.map(c => [c.component_id, c]));
  state.alternatives.components = state.alternatives.components.map(c => newById.get(c.component_id) || c);
  for (const c of allowed) {
    if (!state.alternatives.components.find(x => x.component_id === c.component_id)) {
      state.alternatives.components.push(c);
    }
  }
  if (delta.commentary) state.alternatives.commentary = delta.commentary;
}

// When baseline is rerun for a subset of components, downstream (alternatives,
// selections, report) for those same components is stale by definition.
function invalidateDownstreamFor(cidSet) {
  // Binding goes, intent stays (increment 3). The chat agent emits
  // rerun_baseline before rerun_alternatives for the same scope, and this is
  // the step that used to destroy the choice the next step could have restored.
  backfillIntent();
  if (state.alternatives && Array.isArray(state.alternatives.components)) {
    state.alternatives.components = state.alternatives.components.filter(c => !cidSet.has(c.component_id));
    if (state.alternatives.components.length === 0) state.alternatives = null;
  }
  if (state.selections) {
    Object.keys(state.selections).forEach(cid => { if (cidSet.has(cid)) delete state.selections[cid]; });
  }
  state.reportMarkdown = null;
}

// Apply state deltas returned by the chat agent.
function applyAgentStateUpdates(updates) {
  if (!updates || typeof updates !== 'object') return;
  let touched = false;
  // If the agent returns a fresh baseline/alternatives bag, skaling was applied
  // and no stale warning is needed. Otherwise we flag stale on project mutation.
  let baselineStale = false;
  let altsStale = false;
  let materialChanged = false;

  if (updates.project) {
    const prevIds = new Set((state.project?.components || []).map(c => c.id));
    const newIds = new Set((updates.project.components || []).map(c => c.id));
    const sameIds = prevIds.size === newIds.size && [...prevIds].every(id => newIds.has(id));
    // Detect material/category change vs pure quantity change. If category or
    // name changed on any matching component, the baseline value for it is
    // stale and only a recompute restores correctness.
    if (state.project && sameIds) {
      const prevById = new Map(state.project.components.map(c => [c.id, c]));
      for (const c of updates.project.components) {
        const p = prevById.get(c.id);
        if (p && (p.category !== c.category || p.name !== c.name)) {
          materialChanged = true;
          break;
        }
      }
    } else if (state.project && !sameIds) {
      // Added or removed component => baseline coverage changed.
      materialChanged = true;
    }
    state.project = updates.project;
    touched = true;
    if (!('baseline' in updates)) baselineStale = true;
    if (!('alternatives' in updates)) altsStale = true;
    // If components were removed, mirror in baseline/alternatives/selections.
    if (!sameIds) {
      if (state.baseline && state.baseline.components) {
        state.baseline.components = state.baseline.components.filter(c => newIds.has(c.component_id));
      }
      if (state.alternatives && state.alternatives.components) {
        state.alternatives.components = state.alternatives.components.filter(c => newIds.has(c.component_id));
      }
      if (state.selections) {
        Object.keys(state.selections).forEach(cid => { if (!newIds.has(cid)) delete state.selections[cid]; });
      }
      // Drop directives and selection intent belonging to removed components.
      // Removal is the only thing that ends intent (increment 3).
      pruneDirectives([...prevIds].filter(id => !newIds.has(id)));
      dropIntent([...prevIds].filter(id => !newIds.has(id)));
    }
  }

  if (updates.baseline) {
    state.baseline = updates.baseline;
    touched = true;
  }
  if (updates.alternatives) {
    state.alternatives = updates.alternatives;
    touched = true;
  }
  if (updates.selections) {
    state.selections = updates.selections;
    // The agent writes the bag wholesale; adopt it as intent so a later rerun
    // can carry it (increment 3).
    backfillIntent();
    touched = true;
  }
  // Presence, not truthiness. The guard that actually carries the emptiness sits
  // on the server, in build_state_updates: an empty dict is falsy in PYTHON, so
  // `if overrides:` there would drop the key and never tell the client that the
  // last override went. Here the same shape is defence in depth rather than the
  // fix, because {} is truthy in JS. It earns its place on one input: an explicit
  // null, which `if (updates.overrides)` would skip while leaving the stale bag
  // on screen. The four bags above are never legitimately empty, so the trap is
  // not there.
  if (Object.prototype.hasOwnProperty.call(updates, 'overrides')) {
    state.overrides = updates.overrides || {};
    touched = true;
  }
  // Same test, same reasoning: unbinding the last EPD or clearing the last
  // installed row leaves the bag empty, and that emptiness is the edit.
  if (Object.prototype.hasOwnProperty.call(updates, 'as_built')) {
    state.as_built = updates.as_built || {};
    touched = true;
    // The outcome table is derived server-side, so it does not follow from the
    // new bag by itself. Without this the numbers on screen belong to the
    // previous edit.
    if (isFollowup()) refreshFollowup();
  }

  if (touched) {
    refreshResults();
    scheduleAutoSave();
  }

  // Surface staleness only when the agent couldn't scale — i.e. project changed but baseline/alternatives didn't come back.
  // For material/category changes we surface a concrete action button so the user does not have to guess
  // (chat agent's prompt promises a "Räkna om baslinjen"-button after such changes).
  if (baselineStale && state.baseline) {
    if (materialChanged) {
      renderRecomputeBaselineAction();
    } else {
      addMsg('⚠️ Baslinjen är nu inaktuell efter ändringen. Kör om den för att få nya värden.', 'system');
    }
  }
  if (altsStale && state.alternatives) {
    if (materialChanged) {
      renderRecomputeAlternativesAction();
    } else {
      addMsg('⚠️ Alternativen är nu inaktuella efter ändringen. Kör om dem för aktuella förslag.', 'system');
    }
  }
}

// Action-button row inside the chat. Used after material/category change so the user can recompute
// without typing or hunting for a hidden control.
function renderActionRow(label, handler) {
  document.querySelectorAll('.msg.action-row[data-label="' + label + '"]').forEach(el => el.remove());
  const el = document.createElement('div');
  el.className = 'msg bot action-row';
  el.dataset.label = label;
  const btn = document.createElement('button');
  btn.type = 'button';
  btn.className = 'action-btn';
  btn.textContent = label;
  btn.onclick = async () => {
    btn.disabled = true;
    try { await handler(); } finally { el.remove(); }
  };
  el.appendChild(btn);
  document.getElementById('messages').appendChild(el);
  el.scrollIntoView({behavior:'smooth'});
}

function renderRecomputeBaselineAction() {
  renderActionRow('Räkna om baslinjen', async () => { await runBaseline(); });
}

function renderRecomputeAlternativesAction() {
  renderActionRow('Sök nya alternativ', async () => { await runAlternatives(); });
}

// === Helpers ===
// About modal (Feature 5)
function openAbout() { document.getElementById('aboutModal').style.display = 'flex'; }
function closeAbout() { document.getElementById('aboutModal').style.display = 'none'; }
document.addEventListener('keydown', e => { if (e.key === 'Escape') { closeAbout(); closeWelcome(); closeProjectMeta(); } });

// === Fastighet och tidpunkt ===
// Vilken byggnad analysen gäller och ungefär när arbetet ska ske. Båda frivilliga,
// och medvetet utanför analysflödet: inget steg läser dem idag. De finns för att en
// samling analyser ska gå att läsa som ett bestånd över tid i stället för som lösa
// rapporter. Efterfrågat av Anna Florqvist 2026-09-07.
let _projectMetaReturnFocus = null;

function openProjectMeta() {
  document.getElementById('projectMenu').style.display = 'none';
  _projectMetaReturnFocus = document.activeElement;
  document.getElementById('metaPropertyRef').value = state.propertyRef || '';
  document.getElementById('metaPlannedStart').value = state.plannedStart || '';
  document.getElementById('projectMetaModal').style.display = 'flex';
  document.getElementById('metaPropertyRef').focus();
}

function closeProjectMeta() {
  const m = document.getElementById('projectMetaModal');
  // Guarded because Escape fires this for every dialog. Without the check, a
  // closed dialog would still yank focus back on every Escape press.
  if (!m || m.style.display === 'none') return;
  m.style.display = 'none';
  if (_projectMetaReturnFocus && _projectMetaReturnFocus.focus) _projectMetaReturnFocus.focus();
  _projectMetaReturnFocus = null;
}

function saveProjectMeta() {
  state.propertyRef = document.getElementById('metaPropertyRef').value.trim();
  // A browser without <input type="month"> renders a plain text box, so the value
  // can be anything. Only a well-formed month is kept; the column is a DATE and a
  // bad string would fail the save rather than this one field.
  const month = document.getElementById('metaPlannedStart').value.trim();
  state.plannedStart = /^\d{4}-(0[1-9]|1[0-2])$/.test(month) ? month : '';
  closeProjectMeta();
  if (HAS_SUPABASE && currentUser) scheduleAutoSave();
}

// === Välkomstdialog ===
// Kryssrutan är en preferens, inte en engångshandling: den speglar sparat läge
// när dialogen öppnas och skrivs när den stängs, oavsett hur den stängs.
const WELCOME_KEY = 'aida_welcome_seen';
let _welcomeReturnFocus = null;

function welcomeSuppressed() {
  try { return localStorage.getItem(WELCOME_KEY) === '1'; } catch (e) { return false; }
}

function openWelcome(force) {
  const m = document.getElementById('welcomeModal');
  if (!m || m.style.display === 'flex') return;
  if (!force && welcomeSuppressed()) return;
  const chk = document.getElementById('welcomeDontShow');
  if (chk) chk.checked = welcomeSuppressed();
  _welcomeReturnFocus = document.activeElement;
  m.style.display = 'flex';
  // preventScroll: annars rullar fokuseringen rutan förbi sin egen rubrik när
  // innehållet är högre än fönstret.
  const btn = document.getElementById('welcomeStartBtn');
  if (btn) btn.focus({preventScroll: true});
}

function closeWelcome() {
  const m = document.getElementById('welcomeModal');
  if (!m || m.style.display !== 'flex') return;
  const chk = document.getElementById('welcomeDontShow');
  try {
    if (chk && chk.checked) localStorage.setItem(WELCOME_KEY, '1');
    else localStorage.removeItem(WELCOME_KEY);
  } catch (e) {}
  m.style.display = 'none';
  if (_welcomeReturnFocus && _welcomeReturnFocus.focus) _welcomeReturnFocus.focus();
  _welcomeReturnFocus = null;
}

// Håll tabbningen inne i dialogen så länge den är öppen.
document.addEventListener('keydown', e => {
  if (e.key !== 'Tab') return;
  const m = document.getElementById('welcomeModal');
  if (!m || m.style.display !== 'flex') return;
  const focusable = m.querySelectorAll('button, input, a[href]');
  if (!focusable.length) return;
  const first = focusable[0], last = focusable[focusable.length - 1];
  if (e.shiftKey && document.activeElement === first) { e.preventDefault(); last.focus(); }
  else if (!e.shiftKey && document.activeElement === last) { e.preventDefault(); first.focus(); }
});

// Reasoning toggle (Feature 2)
function toggleReasoning(id, e) {
  e.stopPropagation();
  const row = document.getElementById('reasoning-' + id);
  if (!row) return;
  const isHidden = row.style.display === 'none';
  row.style.display = isHidden ? '' : 'none';
  e.target.textContent = isHidden ? 'D\u00f6lj' : 'Visa mer';
}

// A figure taken from GWP-GHG rather than GWP-fossil sits on a different basis
// than the baseline and the rest of the table. It is used only where an EPD's
// own components did not add up (Henric, 2026-08-17: use the fallback, but mark
// it). The marking is the condition for allowing it at all, so the badge goes
// beside the source rather than into a tooltip nobody opens.
function gwpBasisBadge(alt) {
  if (!alt || alt.gwp_basis !== 'ghg') return '';
  return '<span class="source-badge source-aggregate" title="Bygger på GWP-GHG '
    + '(totalt exklusive biogent kol), inte GWP-fossil som resten av tabellen. '
    + 'Produktens deklaration går inte ihop, så fossilvärdet gick inte att '
    + 'använda.">GWP-GHG</span> ';
}

// On the search hit, not on the row after binding. The choice between two
// declarations is made here, and a marking that only appears once the row is
// bound arrives after the decision it was meant to inform.
function estimatedBadge(c) {
  if (!c || c.gwp_source !== 'estimated') return '';
  return '<span class="source-badge source-aggregate" title="Deklarationen är '
    + 'riktig och går att slå upp, men siffran är uppskattad ur den i stället '
    + 'för avläst ur den. Det står också i klimatredovisningen.">Uppskattad</span>';
}

function formatSource(source) {
  if (!source) return '';
  if (source.startsWith('[EPD]')) return '<span class="source-badge source-verified">EPD</span>' + esc(source.replace('[EPD] ', ''));
  if (source.startsWith('[Palats]')) return '<span class="source-badge source-verified">Palats</span>' + esc(source.replace('[Palats] ', ''));
  if (source.includes('Boverket')) return '<span class="source-badge source-verified">BVK</span>' + esc(source);
  // EPD-typvärde: median of upper-half EPDs by GWP per category — better
  // than Uppskattning, less precise than a single verified EPD. Approximates
  // NollCO2 'Typical' for categories Boverket lacks.
  if (source.includes('EPD-typvärde') || source.includes('EPD-medel') || source.includes('EPD-median')) return '<span class="source-badge source-aggregate">EPD-typvärde</span>' + esc(source);
  if (source.includes('EPD') || source.includes('Environdec')) return '<span class="source-badge source-verified">EPD</span>' + esc(source);
  if (source.startsWith('[Uppskattning]')) return '<span class="source-badge source-estimate">Est.</span>' + esc(source.replace('[Uppskattning] ', ''));
  if (source.includes('Uppskattning')) return '<span class="source-badge source-estimate">Est.</span>' + esc(source);
  return esc(source);
}

function getTypeBadge(alt) {
  if (alt.alternative_type === 'info') return '<span class="type-badge" style="background:var(--kk-gray-100);color:var(--kk-gray-500)">Info</span>';
  if (alt.alternative_type === 'reuse') return '<span class="type-badge type-reuse">\u00c5terbruk</span>';
  if (alt.alternative_type === 'climate_optimized') return '<span class="type-badge type-optimized">Klimatopt.</span>';
  return '<span class="type-badge type-baseline">Baslinje</span>';
}

// === Tab renderers ===
function quantitySourceBadge(src) {
  if (src === 'user_specified') return '<span class="source-badge source-verified" title="Antalet kommer fr\u00e5n din projektbeskrivning">Du angav</span>';
  return '<span class="source-badge source-estimate" title="Aida uppskattade antalet utifr\u00e5n area och byggnadstyp \u2014 granska om n\u00e5got verkar fel">Aida uppskattat</span>';
}

// Snapshot of the inferred text at the moment edit was opened. Restores
// on Cancel even if state.project was swapped by a concurrent chat-agent
// update mid-edit.
let _naEditSnapshot = null;

function renderNeedsAnalysis(na) {
  // Editorial pairing: user voice (gray) \u2194 Aida voice (red), with vertical
  // accent stripes and a clear visual transition between them.
  const hasAny = na && (
    (na.from_user || '').trim() ||
    (na.inferred || '').trim() ||
    (Array.isArray(na.assumptions) && na.assumptions.length) ||
    (Array.isArray(na.would_clarify) && na.would_clarify.length)
  );
  if (!hasAny) {
    return '<div class="needs-card">' +
      '<div class="needs-card-head"><div class="needs-card-title">Aidas behovsanalys</div></div>' +
      '<div class="needs-empty">Ingen behovsanalys finns f\u00f6r det h\u00e4r projektet \u2014 analysen tillkom efter att projektet skapades. K\u00f6ra om intake i chatten ger en analys som styr alternativvalet.</div></div>';
  }
  const fromUser = na.from_user || '';
  const inferred = na.inferred || '';
  const assumptions = Array.isArray(na.assumptions) ? na.assumptions : [];
  const clarify = Array.isArray(na.would_clarify) ? na.would_clarify : [];

  let html = '<div class="needs-card">';
  html += '<div class="needs-card-head">';
  html += '<div class="needs-card-title">Aidas behovsanalys</div>';
  html += '<div class="needs-card-sub">Granska f\u00f6re baslinje \u00b7 korrigera om Aidas l\u00e4sning \u00e4r fel</div>';
  html += '</div>';
  html += '<div class="needs-body">';

  // User voice block
  html += '<div class="voice-block voice-user">';
  html += '<div class="voice-label"><span class="dot"></span>Du sa</div>';
  html += '<div class="voice-text">' + (fromUser ? esc(fromUser) : '<em class="empty">(inget direkt fr\u00e5n din beskrivning)</em>') + '</div>';
  html += '</div>';

  // Transition
  html += '<div class="voice-transition">\u2193 Aidas l\u00e4sning av detta</div>';

  // Aida voice block with edit affordance
  html += '<div class="voice-block voice-aida" id="aidaVoiceBlock">';
  html += '<div class="voice-aida-actions">';
  html += '<button type="button" class="voice-aida-edit" id="naEditBtn" onclick="toggleNeedsEdit()">';
  html += '<svg viewBox="0 0 16 16" fill="none" stroke="currentColor" stroke-width="1.8"><path d="M11 2l3 3-9 9H2v-3l9-9z"/></svg>';
  html += 'Justera Aidas l\u00e4sning</button>';
  html += '</div>';
  html += '<div class="voice-label"><span class="dot"></span>Aida tolkar</div>';
  html += '<div class="voice-text" id="naInferredView">' + (inferred ? esc(inferred) : '<em class="empty">(ingen inferens)</em>') + '</div>';
  // Empty textarea \u2014 value populated via .value to preserve quotes/ampersands
  html += '<textarea class="voice-aida-textarea" id="naInferredEdit"></textarea>';
  html += '<div class="voice-aida-edit-actions" id="naEditActions">';
  html += '<button type="button" class="btn-na-cancel" onclick="cancelNeedsEdit()">Avbryt</button>';
  html += '<button type="button" class="btn-na-save" onclick="saveNeedsEdit()">Spara</button>';
  html += '</div>';
  html += '</div>';

  // Meta row \u2014 assumptions + would_clarify
  if (assumptions.length || clarify.length) {
    html += '<div class="needs-meta-row">';
    html += '<div>';
    if (assumptions.length) {
      html += '<div class="needs-meta-label">Antaganden Aida gjort</div>';
      html += '<ul class="needs-meta-list needs-meta-assumptions">';
      assumptions.forEach(a => { html += '<li>' + esc(a) + '</li>'; });
      html += '</ul>';
    }
    html += '</div><div>';
    if (clarify.length) {
      html += '<div class="needs-meta-label">Aida hade g\u00e4rna vetat</div>';
      html += '<ul class="needs-meta-list needs-meta-clarify">';
      clarify.forEach(q => { html += '<li>' + esc(q) + '</li>'; });
      html += '</ul>';
    }
    html += '</div></div>';
  }

  html += '</div></div>';
  return html;
}

function _populateNeedsTextarea() {
  // Populate the textarea via .value (not innerHTML) so the raw inferred
  // text \u2014 including quotes and ampersands \u2014 is preserved verbatim.
  const ta = document.getElementById('naInferredEdit');
  if (!ta) return;
  const na = state.project && state.project.needs_analysis;
  ta.value = (na && na.inferred) || '';
}

function toggleNeedsEdit() {
  const ta = document.getElementById('naInferredEdit');
  const block = document.getElementById('aidaVoiceBlock');
  if (!ta || !block) return;
  _naEditSnapshot = ta.value;  // remember pre-edit value
  block.classList.add('is-editing');
  ta.focus();
}

function cancelNeedsEdit() {
  const ta = document.getElementById('naInferredEdit');
  const block = document.getElementById('aidaVoiceBlock');
  if (!ta || !block) return;
  // Restore from the pre-edit snapshot \u2014 not from state.project, which may
  // have been mutated by a concurrent chat-agent update during the edit.
  ta.value = _naEditSnapshot != null ? _naEditSnapshot : ta.value;
  _naEditSnapshot = null;
  block.classList.remove('is-editing');
}

function saveNeedsEdit() {
  const ta = document.getElementById('naInferredEdit');
  if (!ta) return;
  const newVal = ta.value.trim();
  if (!state.project.needs_analysis) state.project.needs_analysis = {from_user:'',inferred:'',assumptions:[],would_clarify:[]};
  state.project.needs_analysis.inferred = newVal;
  _naEditSnapshot = null;
  if (typeof scheduleAutoSave === 'function') scheduleAutoSave();
  renderProjektContent();
}

// Section renderers, split in two (orchestration-redesign §12.2, step 1).
//
// <name>Html(st) is pure: state in, HTML string out, no DOM. render<Name>Content()
// keeps the old name and every call site, writes into the tab and binds events.
// The split is what lets the same section render into a tab (Stegvis) or into
// the sheet (Chatt) without a second copy of the markup. Nothing about the
// output changes here; scripts/test_render_parity.js holds that to the letter.
//
// The pure half takes state as an argument rather than reading the global, so a
// caller can render a snapshot. Event handlers stay in the DOM half and keep
// reading the global `state`, because they fire later, when it has moved on.
// `cfg` arrives in step 2, when there is something for it to vary.
// cfg.hideTitle: the sheet prints the section name in its own header, so the
// renderer must not print it a second time. Absent cfg means the tab view, which
// is what every existing caller wants and what the parity test pins.
// An editable cell carries the two things the commit needs, and nothing else:
// which component, which field. The old value rides along so a blur that changed
// nothing can return without a round trip.
function cellInput(cid, field, value, label, type) {
  const v = value == null ? '' : String(value);
  return '<input class="cell-input" data-cid="' + esc(cid) + '" data-field="' + field + '"' +
         ' data-was="' + esc(v) + '" value="' + esc(v) + '" aria-label="' + esc(label) + '"' +
         (type === 'number' ? ' type="number" min="0" step="any"' : ' type="text"') + '>';
}

// Unit is a closed set: add_component rejects anything outside m2/st/lm, so a
// free-text cell could only produce a rejection the user has to guess their way
// out of. An unrecognised stored unit is kept as an option so an older analysis
// does not silently have its unit rewritten by the act of opening it.
const CELL_UNITS = ['m2', 'st', 'lm'];
function cellUnit(cid, unit) {
  const cur = unit == null ? '' : String(unit);
  const opts = CELL_UNITS.concat(CELL_UNITS.indexOf(cur) === -1 && cur ? [cur] : []);
  return '<select class="cell-input" data-cid="' + esc(cid) + '" data-field="unit"' +
         ' data-was="' + esc(cur) + '" aria-label="Enhet">' +
         opts.map(u => '<option value="' + esc(u) + '"' + (u === cur ? ' selected' : '') + '>' + esc(u) + '</option>').join('') +
         '</select>';
}

function projektHtml(st, cfg) {
  const d = st.project;
  let html = (cfg && cfg.hideTitle) ? '' : '<div class="section-title">Projektinformation</div>';
  html += renderNeedsAnalysis(d.needs_analysis);
  html += '<div class="comp-card"><div class="comp-card-header"><h3>' + esc(d.building_type) + ', ' + esc(d.area_bta) + ' m\u00b2 BTA' + (d.name ? ' (' + esc(d.name) + ')' : '') + '</h3></div>';
  const edit = !!(cfg && cfg.editable);
  html += '<table class="comp-table"><thead><tr><th>Komponent</th><th>Antal</th><th>Enhet</th><th>Kategori</th><th>K\u00e4lla</th>' + (edit ? '<th></th>' : '') + '</tr></thead><tbody>';
  d.components.forEach(c => {
    const usage = c.usage_context ? '<div class="usage-context"><span class="usage-context-label">Anv\u00e4ndning</span>' + esc(c.usage_context) + '</div>' : '';
    if (edit) {
      html += '<tr>' +
        '<td>' + cellInput(c.id, 'name', c.name, 'Namn') + usage + '</td>' +
        '<td class="cell-num">' + cellInput(c.id, 'quantity', c.quantity, 'Antal', 'number') + '</td>' +
        '<td>' + cellUnit(c.id, c.unit) + '</td>' +
        '<td>' + cellInput(c.id, 'category', c.category, 'Kategori') + '</td>' +
        '<td>' + quantitySourceBadge(c.quantity_source) + '</td>' +
        '<td><button class="cell-remove" data-remove="' + esc(c.id) + '" title="Ta bort ' + esc(c.name) + '" aria-label="Ta bort ' + esc(c.name) + '">&#xD7;</button></td>' +
        '</tr>';
    } else {
      html += '<tr>' +
        '<td>' + '<div style="font-weight:500">' + esc(c.name) + '</div>' + usage + '</td>' +
        '<td>' + esc(c.quantity) + '</td>' +
        '<td>' + esc(c.unit) + '</td>' +
        '<td>' + esc(c.category || '\u2013') + '</td>' +
        '<td>' + quantitySourceBadge(c.quantity_source) + '</td>' +
        '</tr>';
    }
  });
  html += '</tbody></table>';
  if (edit) {
    html += '<div class="cell-add"><button id="addCompBtn" type="button">L\u00e4gg till komponent</button></div>';
  }
  html += '</div>';
  if (d.description) {
    html += '<div class="comp-card" style="margin-top:12px"><div class="comp-card-header"><h3>Beskrivning</h3></div><div style="padding:12px 16px;font-size:13px;color:var(--kk-gray-500);line-height:1.5">' + esc(d.description) + '</div></div>';
  }
  return html;
}

function renderProjektContent() {
  // Henric, 2026-09-08: Stegvis gets editable cells too, gates and all. The two
  // views differ in how much the tool steers, not in what you are allowed to fix.
  document.getElementById('resultContent').innerHTML = projektHtml(state, {editable: true});
  _populateNeedsTextarea();
  bindCells();
}

// === Editable cells (orchestration-redesign §12.4) ===
//
// Rule 1: a cell and the chat are two doors onto the same mutation. The cell does
// not edit state; it posts to /api/mutate, which runs the same function the chat
// agent's tool calls, and the answer comes back through applyAgentStateUpdates
// exactly like a chat turn. So a cell cannot invent its own idea of what goes
// stale, and the scoped reruns happen either way.
//
// Rule 3: a cell is locked while anything is running. Without it an edit made
// during a rerun is overwritten by that rerun's answer, silently.
let _mutationInFlight = false;

function cellsLocked() { return _mutationInFlight || _runInFlight; }

function setCellsDisabled(on) {
  document.querySelectorAll('.cell-input, .cell-remove, #addCompBtn').forEach(el => { el.disabled = on; });
}

function bindCells() {
  document.querySelectorAll('.cell-input').forEach(el => {
    // 'change', not 'input': commit on blur or Enter, not on every keystroke.
    el.onchange = () => commitCell(el);
    el.onkeydown = e => {
      if (e.key === 'Enter') { e.preventDefault(); el.blur(); }
      // Escape puts the cell back the way it was without a round trip.
      else if (e.key === 'Escape') { el.value = el.dataset.was; el.blur(); }
    };
  });
  document.querySelectorAll('.cell-remove').forEach(el => {
    el.onclick = () => removeComponentCell(el.dataset.remove);
  });
  const add = document.getElementById('addCompBtn');
  if (add) add.onclick = () => addComponentCell();
  if (cellsLocked()) setCellsDisabled(true);
}

// An as-built cell is the same kind of cell with a different destination: it
// describes what was installed, not what is planned, so it must not land in
// update_component and quietly rewrite the plan. Same lock, same escape key,
// same round trip; one attribute decides which tool it reaches.
function asBuiltInput(cid, field, value, label, type) {
  const v = value == null ? '' : String(value);
  return '<input class="cell-input" data-ab="1" data-cid="' + esc(cid) + '" data-field="' + field + '"' +
         ' data-was="' + esc(v) + '" value="' + esc(v) + '" aria-label="' + esc(label) + '"' +
         (type === 'number' ? ' type="number" min="0" step="any"' : ' type="text"') + '>';
}

const AS_BUILT_NUMERIC = ['quantity', 'actual_cost', 'transport_km'];

function commitAsBuiltCell(el) {
  const raw = el.value.trim();
  if (raw === (el.dataset.was || '')) return;
  const field = el.dataset.field;
  const input = {component_id: el.dataset.cid};
  if (AS_BUILT_NUMERIC.indexOf(field) !== -1) {
    // Empty clears the field rather than writing a zero. Zero kronor is a
    // statement (donated, already owned); blank is the absence of one, and the
    // outcome table renders the two differently on purpose.
    if (raw === '') { input[field] = null; }
    else {
      const n = Number(raw);
      if (!(n >= 0)) {
        addMsg('Värdet måste vara ett tal, noll eller större.', 'system');
        el.value = el.dataset.was;
        return;
      }
      input[field] = n;
    }
  } else {
    input[field] = raw;
  }
  el.classList.add('saving');
  return sendMutation('set_as_built', input);
}

function commitCell(el) {
  if (el.dataset.ab) return commitAsBuiltCell(el);
  const value = el.value.trim();
  if (value === (el.dataset.was || '')) return;   // blur without an edit
  const field = el.dataset.field;
  if (field === 'name' && !value) {
    addMsg('En komponent måste ha ett namn. Ta bort raden i stället om den inte ska vara med.', 'system');
    el.value = el.dataset.was;
    return;
  }
  const input = {component_id: el.dataset.cid};
  if (field === 'quantity') {
    const n = Number(value);
    if (!(n > 0)) {
      addMsg('Mängden måste vara ett tal större än noll.', 'system');
      el.value = el.dataset.was;
      return;
    }
    input.quantity = n;
  } else {
    input[field] = value;
  }
  el.classList.add('saving');
  return sendMutation('update_component', input);
}

function removeComponentCell(cid) {
  const comp = (state.project && state.project.components || []).find(c => c.id === cid);
  if (!comp) return;
  if (!confirm('Ta bort ' + comp.name + '? Baslinje, alternativ och val för just den komponenten försvinner. Övriga rörs inte.')) return;
  return sendMutation('remove_component', {component_id: cid});
}

function addComponentCell() {
  const name = (prompt('Vad ska läggas till?') || '').trim();
  if (!name) return;
  const qty = Number((prompt('Hur många eller hur mycket? (bara siffran)') || '').trim());
  if (!(qty > 0)) { addMsg('Ingen komponent tillagd: mängden måste vara ett tal större än noll.', 'system'); return; }
  const unit = (prompt('Enhet: m2, st eller lm') || '').trim();
  if (CELL_UNITS.indexOf(unit) === -1) { addMsg('Ingen komponent tillagd: enheten måste vara m2, st eller lm.', 'system'); return; }
  return sendMutation('add_component', {name, quantity: qty, unit, quantity_source: 'user_specified'});
}

async function sendMutation(tool, input) {
  if (cellsLocked()) return;
  _mutationInFlight = true;
  setCellsDisabled(true);
  try {
    const r = await authFetch('/api/mutate', {
      method: 'POST', headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({tool, input, project: state.project, baseline: state.baseline,
                            alternatives: state.alternatives, selections: state.selections,
                            overrides: state.overrides, as_built: state.as_built}),
    });
    const d = await r.json();
    if (d.error) { addMsg('Fel: ' + d.error, 'system'); refreshResults(); return; }
    // A refused edit is not a failure to report as one: the handler explains why
    // ("det finns redan en komponent som heter..."), and the view goes back to
    // what is actually stored so the cell never shows a value the state lacks.
    if (!d.ok) { addMsg(d.message, 'system'); refreshResults(); return; }
    addMsg(d.message, 'system');
    applyAgentStateUpdates(d.state_updates);
    refreshResults();
    const pending = d.state_updates && d.state_updates.pending_actions;
    if (pending && pending.length) {
      // Hand the lock over, do not drop it. These reruns are dispatched with
      // orchestrated=true, and that flag is exactly what stops them calling
      // setLoading, because it was written for runChat which already holds the
      // lock around the whole sequence. So nobody holds it here unless we do,
      // and an unlocked rerun is the worst window there is: applyAgentStateUpdates
      // replaces baseline, alternatives and selections wholesale, so a second
      // edit sent meanwhile answers from the pre-rerun snapshot and its reply
      // overwrites everything the rerun had merged in. setLoading also greys the
      // chat input and shows the working indicator, which is what should be on
      // screen while a run the user's own edit started is still going.
      setLoading(true);
      _mutationInFlight = false;
      try {
        await processPendingActions(pending);
      } finally {
        setLoading(false);
      }
    }
  } catch (e) {
    addMsg('Ändringen gick inte igenom: ' + e.message, 'system');
    refreshResults();
  } finally {
    _mutationInFlight = false;
    setCellsDisabled(cellsLocked());
  }
}

// === Overrides (orchestration-redesign §12.5) ===
//
// Twin of `apply` in src/aida/overrides.py. It exists because the view renders
// without asking the server, and the report is written without asking the
// browser; there is no single runtime both read from. The twin is allowed,
// drifting is not: scripts/test_overrides_agree.js runs the same fixtures
// through both and requires identical output.
//
// The same baseline figure lives in four places at once - the baseline bag, the
// alternatives bag's baseline row, the selection's baseline field, and the
// selected alternative itself when the user chose "Baslinje". One override has
// to reach all four, or the sheet contradicts itself half a screen apart.
const OVERRIDE_FIELDS = ['baseline_co2e', 'baseline_cost'];
const OVERRIDE_NOTE_MAX = 120;
const OVERRIDE_LABELS = {baseline_co2e: 'Baslinje CO₂e', baseline_cost: 'Baslinje kostnad'};

function overrideOf(cid, field) {
  const entry = state.overrides && state.overrides[cid];
  const one = entry && entry[field];
  return (one && one.value !== undefined && one.value !== null) ? one : null;
}

function hasOverrides(ov) {
  if (!ov) return false;
  return Object.keys(ov).some(cid => OVERRIDE_FIELDS.some(f => {
    const one = ov[cid] && ov[cid][f];
    return one && one.value !== undefined && one.value !== null;
  }));
}

// Returns a state-shaped object with the overrides laid on. Never mutates the
// argument: the stored bags stay exactly what the pipeline produced, which is
// what makes "what would Aida have said?" answerable at any point.
function effectiveState(st) {
  const ov = st.overrides;
  if (!hasOverrides(ov)) return st;
  const out = Object.assign({}, st);
  out.baseline = st.baseline ? JSON.parse(JSON.stringify(st.baseline)) : st.baseline;
  out.alternatives = st.alternatives ? JSON.parse(JSON.stringify(st.alternatives)) : st.alternatives;
  out.selections = st.selections ? JSON.parse(JSON.stringify(st.selections)) : st.selections;

  Object.keys(ov).forEach(cid => {
    const co2e = (ov[cid] && ov[cid].baseline_co2e) || null;
    const cost = (ov[cid] && ov[cid].baseline_cost) || null;
    if (!co2e && !cost) return;

    if (out.baseline && out.baseline.components) {
      out.baseline.components.forEach(c => {
        if (c.component_id !== cid) return;
        if (co2e) { c.co2e_kg = co2e.value; c.co2e_override = co2e.note; }
        if (cost) { c.cost_sek = cost.value; c.cost_override = cost.note; }
      });
    }
    if (out.alternatives && out.alternatives.components) {
      out.alternatives.components.forEach(c => {
        if (c.component_id !== cid) return;
        if (co2e) { c.baseline_co2e_kg = co2e.value; c.baseline_co2e_override = co2e.note; }
        if (cost) { c.baseline_cost_sek = cost.value; c.baseline_cost_override = cost.note; }
      });
    }
    // hasOwnProperty, not out.selections[cid]. A key named "__proto__" survives
    // JSON.parse as a real own key, and indexing a plain object with it hands
    // out Object.prototype to write on - so one such key in a saved analysis
    // would corrupt every object in the tab on load. It cannot come through the
    // handlers, which check the id against the project, but it can come through
    // a hand-edited analysis, and then it is re-applied on every load. Python's
    // dict.get is inert here, which is exactly why the JS side has to say so.
    const sel = (out.selections && Object.prototype.hasOwnProperty.call(out.selections, cid))
      ? out.selections[cid] : null;
    if (sel) {
      const chosen = sel.selected_alternative;
      // "Baslinje" as a choice means the selected figure IS the baseline figure.
      // Overriding one and not the other would put two different numbers for the
      // same decision in the same report.
      const pickedBaseline = chosen && chosen.name === 'Baslinje';
      if (co2e) {
        sel.baseline_co2e_kg = co2e.value;
        sel.baseline_co2e_override = co2e.note;
        if (pickedBaseline) { chosen.co2e_kg = co2e.value; chosen.co2e_override = co2e.note; }
      }
      if (cost) {
        sel.baseline_cost_sek = cost.value;
        sel.baseline_cost_override = cost.note;
        if (pickedBaseline) { chosen.cost_sek = cost.value; chosen.cost_override = cost.note; }
      }
    }
  });
  return out;
}

// The line the chat writes when a baseline is done, and the one restoreUI
// reconstructs when an analysis is reopened without its chat log. One function
// because it is one sentence: two copies drifted apart is how the same screen
// ended up able to show two different totals. It sums the effective state for
// the same reason the tab beside it renders the effective state - a manually set
// figure survives a rerun, so the raw response is no longer what anyone sees.
function baselineDoneMsg() {
  const shown = effectiveState(state).baseline.components;
  const total = shown.reduce((s,c) => s + c.co2e_kg, 0);
  return 'Baslinje klar: **' + Math.round(total).toLocaleString('sv')
    + ' kg CO₂e** totalt för ' + shown.length + ' komponenter.';
}

// Same shape as gwpBasisBadge (#550): a number with a different origin than the
// rest must never look like the rest. The note is the whole point, so it is in
// the title where a hover reaches it and in the report where a reader does.
// "Manuellt" is the same word the report's appendix uses, deliberately: one name
// for one thing across the sheet, the document and the chat.
function overrideBadge(note) {
  if (!note) return '';
  return ' <span class="source-badge source-manual" title="' + esc(note) + '">Manuellt</span>';
}

// Which cell has its editor open. One at a time: two open forms in one table
// would let someone start a figure in one and save it from the other.
let _overrideOpen = null;

function openOverride(cid, field) {
  if (cellsLocked()) return;
  _overrideOpen = {cid: cid, field: field};
  refreshResults();
  const el = document.getElementById('ovValue');
  if (el) { el.focus(); el.select(); }
}

function closeOverride() {
  _overrideOpen = null;
  refreshResults();
}

// The number itself is the affordance, so a dense numeric column gains no icon
// per row. It wraps rather than appends: an empty span next to the figure is a
// zero-width target, which is the same as no target at all. Outside an editable
// configuration it returns the html untouched, so the pinned renderer output
// stays byte-identical.
function overrideEdit(html, cid, field, note, cfg) {
  if (!(cfg && cfg.editable) || !cid) return html;
  return '<span class="override-open" onclick="openOverride(\'' + esc(cid) + '\',\'' + field
    + '\')" title="' + (note ? 'Ändra det manuella värdet' : 'Ange ett eget värde')
    + '">' + html + '</span>';
}

// Aida's own figure, read from the stored bag rather than from the row we were
// handed. The row comes from effectiveState, so on an overridden cell it already
// carries the user's number - and the form was promising to restore "Aidas
// beräknade värde (213)" when 213 was the very number they had typed.
function computedOf(cid, field) {
  const src = ((state.baseline || {}).components || []).filter(c => c.component_id === cid)[0];
  if (!src) return 0;
  return (field === 'baseline_co2e' ? src.co2e_kg : src.cost_sek) || 0;
}

// A full-width row under the component's own, because a note field cannot live
// in a right-aligned numeric cell and shrinking it to fit would say the note is
// an afterthought. It is the part that makes the number legitimate.
function overrideFormRow(c, cfg) {
  if (!(cfg && cfg.editable) || !_overrideOpen || _overrideOpen.cid !== c.component_id) return '';
  const field = _overrideOpen.field;
  const current = overrideOf(c.component_id, field);
  const computed = computedOf(c.component_id, field);
  const unit = field === 'baseline_co2e' ? 'kg CO₂e' : 'SEK';
  const value = current ? current.value : (computed > 0 ? Math.round(computed) : '');
  return '<tr class="override-row"><td colspan="5">'
    + '<div class="override-form">'
    + '<div><label for="ovValue">' + OVERRIDE_LABELS[field] + ' (' + unit + ')</label>'
    + '<input class="ov-value" id="ovValue" type="number" min="0" step="any" value="' + esc(String(value)) + '"></div>'
    + '<div class="ov-note"><label for="ovNote">Var kommer siffran ifrån?</label>'
    + '<input id="ovNote" type="text" maxlength="' + OVERRIDE_NOTE_MAX + '" value="'
    + esc(current ? current.note : '') + '" placeholder="Till exempel EPD S-P-01234 eller ramavtal 2025-114"></div>'
    // The buttons live in their own box so a narrow results panel wraps them as
    // a group. Loose in the flex row, Avbryt kept the line and Spara dropped
    // below it, which puts the primary action where nothing else is.
    + '<div class="ov-actions">'
    + '<button type="button" class="btn-na-cancel" onclick="closeOverride()">Avbryt</button>'
    + '<button type="button" class="btn-na-save" onclick="saveOverride()">Spara värdet</button>'
    + (current ? '<button type="button" class="btn-na-cancel ov-clear" onclick="clearOverride()">Använd Aidas värde igen</button>' : '')
    + '</div>'
    + '<div class="override-hint" id="ovHint">Aidas beräknade värde ('
    + (computed > 0 ? fmtNum(computed) + ' ' + unit : 'saknas') + ') sparas kvar och visas igen om du tar bort ditt.</div>'
    + '</div></td></tr>';
}

function saveOverride() {
  if (!_overrideOpen) return;
  const value = (document.getElementById('ovValue') || {}).value;
  const note = ((document.getElementById('ovNote') || {}).value || '').trim();
  const hint = document.getElementById('ovHint');
  // Said as direction, not as a complaint: the note is what lets a reader of the
  // report judge the number, so the form explains what to write rather than
  // reporting that a field is empty.
  if (!note) {
    if (hint) { hint.className = 'override-hint bad'; hint.textContent = 'Skriv var siffran kommer ifrån, till exempel ett EPD-nummer eller ett ramavtal. Det följer med till rapporten.'; }
    const el = document.getElementById('ovNote'); if (el) el.focus();
    return;
  }
  if (!(Number(value) >= 0) || value === '') {
    if (hint) { hint.className = 'override-hint bad'; hint.textContent = 'Värdet måste vara ett tal, noll eller större.'; }
    const el = document.getElementById('ovValue'); if (el) el.focus();
    return;
  }
  const target = _overrideOpen;
  _overrideOpen = null;
  return sendMutation('set_override', {component_id: target.cid, field: target.field,
                                       value: Number(value), note: note});
}

function clearOverride() {
  if (!_overrideOpen) return;
  const target = _overrideOpen;
  _overrideOpen = null;
  return sendMutation('clear_override', {component_id: target.cid, field: target.field});
}

// Baseline costs come from the LLM and can be absent. Same rule as everywhere
// else: an absent price is not zero kronor.
function knownCostRollup(rows) {
  const priced = rows.filter(c => c.cost_sek > 0);
  return {
    known: priced.reduce((s,c) => s + c.cost_sek, 0),
    unpriced: rows.filter(c => !(c.cost_sek > 0)).map(c => c.component_name || c.name || ''),
    pricedCount: priced.length,
    total: rows.length,
  };
}

// Small-caps label above a value, the idiom the Boverket-produkt line already
// used. Factored out so the three transparency lines cannot drift apart.
function subLine(label, valueHtml) {
  return '<div style="font-size:11px;color:var(--kk-gray-500);margin-top:3px;font-style:italic">'
    + '<span style="font-style:normal;font-weight:500;color:var(--kk-gray-500);font-size:9.5px;letter-spacing:0.8px;text-transform:uppercase;display:block;margin-bottom:1px">'
    + label + '</span>' + valueHtml + '</div>';
}

function fmtNum(n) {
  if (n === null || n === undefined || isNaN(n)) return '';
  // Two decimals matter at the low end: paint is 0.55 kg/m², and rounding that
  // to "1" would make the most-questioned number on the page look invented.
  const dec = Math.abs(n) < 10 ? 2 : (Math.abs(n) < 100 ? 1 : 0);
  return Number(n).toLocaleString('sv', { maximumFractionDigits: dec });
}

// What the number rests on, in the user's terms. The spread is the honest part:
// a typvärde is one point in a distribution, and showing the range is what lets
// someone judge whether it is a tight estimate or a wide guess.
function basisLine(c) {
  const b = c.basis;
  if (!b || b.kind !== 'epd_typvärde') return '';
  const n = b.sample_size;
  let txt = n + ' EPD:er';
  if (b.min !== null && b.min !== undefined && b.max !== null && b.max !== undefined) {
    txt += ', spann ' + fmtNum(b.min) + ' till ' + fmtNum(b.max);
  }
  if (b.full_median !== null && b.full_median !== undefined) {
    txt += ', median ' + fmtNum(b.full_median);
  }
  // A subtype request that fell back to the category is the case worth calling
  // out: the user asked about vinyl and got all floors. Saying so is the whole
  // point of carrying `level` through instead of inferring it.
  let scope = '';
  if (b.level === 'category' && b.requested_subtype) {
    scope = '<div style="color:var(--kk-red-orange);margin-top:1px">För få EPD:er för '
      + esc(b.requested_subtype) + ', visar hela kategorin (flera materialtyper)</div>';
  }
  return subLine(esc(b.label || 'EPD-typvärde'), txt + scope);
}

function baslinjeHtml(st, cfg) {
  const d = st.baseline;
  const total = d.components.reduce((s,c) => s + c.co2e_kg, 0);
  const cost = knownCostRollup(d.components);
  let html = (cfg && cfg.hideTitle) ? '' : '<div class="section-title">Baslinje (NollCO2-metoden)</div>';
  html += '<div class="method-label">Klimatmetod: GWP-fossil, livscykelskedena A1-A3 (Boverkets klimatdatabas)</div>';
  html += '<div class="source-legend"><span><span class="source-badge source-verified">EPD</span> Verifierad k\u00e4lla</span><span><span class="source-badge source-aggregate">EPD-typvärde</span> Kategori-typvärde (övre halvan)</span><span><span class="source-badge source-estimate">Est.</span> Uppskattning</span></div>';
  html += '<div class="summary">';
  html += '<div class="card"><div class="card-title">Total CO\u2082e</div><div class="value">' + Math.round(total).toLocaleString('sv') + '</div><div class="sublabel">kg CO\u2082e</div></div>';
  html += '<div class="card"><div class="card-title">' + (cost.unpriced.length ? 'Kostnad (delsumma)' : 'Total kostnad') + '</div><div class="value">' + Math.round(cost.known).toLocaleString('sv') + '</div><div class="sublabel">' + (cost.unpriced.length ? 'SEK för ' + cost.pricedCount + ' av ' + cost.total + ' komponenter' : 'SEK') + '</div></div>';
  html += '<div class="card"><div class="card-title">Komponenter</div><div class="value">' + d.components.length + '</div><div class="sublabel">st</div></div>';
  if (cost.unpriced.length) {
    html += '<div style="grid-column:1/-1;font-size:11px;color:var(--kk-red-orange);margin-top:-4px">Saknar pris: ' + esc(cost.unpriced.join(', ')) + '. Ingen prisuppgift hittades, posten är alltså inte gratis.</div>';
  }
  html += '</div>';
  html += '<div class="comp-card"><div class="comp-card-header"><h3>Per komponent</h3></div>';
  html += '<table class="comp-table"><thead><tr><th>Komponent</th><th style="text-align:right">CO\u2082e (kg)</th><th>Klimatk\u00e4lla</th><th style="text-align:right">Kostnad (SEK)</th><th>Prisk\u00e4lla</th></tr></thead><tbody>';
  d.components.forEach(c => {
    // boverket_product is set only when the component's standard material is
    // genuinely in Boverket (same material, e.g. gips→gips). Cross-material
    // proxies are no longer used — components Boverket lacks get an EPD-typvärde
    // instead, with boverket_product empty.
    const productLine = c.boverket_product ? subLine('Boverket-produkt', esc(c.boverket_product)) : '';
    // "Vilket golv har den räknat på?" A golv category aggregate spans a factor
    // of three, so the total alone does not identify a material. This is the
    // standard material the baseline agent named from building type and the
    // component's function — the assumption the whole number rests on, and the
    // one a förvaltare is most likely to disagree with.
    const materialLine = c.assumed_material ? subLine('Antaget standardmaterial', esc(c.assumed_material)) : '';
    // Per-unit figure with the multiplication spelled out, so the total is
    // checkable by hand rather than something to take on faith. Suppressed when
    // the total is overridden: the stored per-unit figure is Aida's and no
    // longer multiplies out to what the column shows, and a multiplication that
    // does not add up is worse than none.
    const perUnit = (c.co2e_per_unit > 0 && c.unit && !c.co2e_override)
      ? '<div style="font-size:10.5px;color:var(--kk-gray-500);margin-top:2px">' + fmtNum(c.co2e_per_unit) + ' kg CO₂e/' + esc(c.unit) + ' × ' + fmtNum(c.quantity) + ' ' + esc(c.unit) + '</div>'
      : '';
    // An overridden figure is marked in the cell it replaced, not only in the
    // report. Someone reading the sheet has to be able to see which numbers are
    // theirs without opening the document.
    const co2eCell = overrideEdit(Math.round(c.co2e_kg).toLocaleString('sv'),
                                  c.component_id, 'baseline_co2e', c.co2e_override, cfg)
      + overrideBadge(c.co2e_override) + perUnit;
    const costCell = overrideEdit(
        (c.cost_sek > 0
          ? Math.round(c.cost_sek).toLocaleString('sv')
          : '<span style="color:var(--kk-gray-500)">Pris saknas</span>'),
        c.component_id, 'baseline_cost', c.cost_override, cfg)
      + overrideBadge(c.cost_override);
    html += '<tr><td style="font-weight:500">' + esc(c.component_name) + materialLine + '</td><td style="text-align:right">' + co2eCell + '</td><td style="font-size:11px">' + formatSource(c.source) + productLine + basisLine(c) + '</td><td style="text-align:right">' + costCell + '</td><td style="font-size:11px">' + esc(c.cost_source || '') + '</td></tr>';
    html += overrideFormRow(c, cfg);
  });
  html += '</tbody></table></div>';
  return html;
}

function renderBaslinjeContent() {
  document.getElementById('resultContent').innerHTML = baslinjeHtml(effectiveState(state), {editable: true});
}

// Climate benefit per krona. Johanna, juni 2026: "Alternativen ska visas med
// bäst klimatnytta per krona, oavsett återbruk eller nyproducerade alternativ."
// Henric confirmed the sort key on 2026-06-10: kr per saved kg CO2e, lowest
// first. Until now the table simply showed whatever order the agent produced.
function altValue(comp, alt) {
  const saved = comp.baseline_co2e_kg - alt.co2e_kg;
  const priced = alt.cost_sek > 0;
  const extra = alt.cost_sek - comp.baseline_cost_sek;
  if (!priced) return { tier: 2, value: 0, saved: saved, priced: false };
  if (saved <= 0) return { tier: 3, value: -saved, saved: saved, priced: true };
  // Cheaper AND lower emissions is strictly better than anything you pay for,
  // so it gets its own tier rather than competing on a negative ratio.
  return { tier: extra <= 0 ? 0 : 1, value: extra / saved, saved: saved, priced: true };
}

// Sort for display while keeping each alternative's index into
// comp.alternatives — selectAlt() looks rows up by that index.
function rankedAlternatives(comp) {
  const rows = comp.alternatives.map((alt, i) => ({ alt: alt, i: i }));
  const info = rows.filter(r => r.alt.alternative_type === 'info');
  const rest = rows.filter(r => r.alt.alternative_type !== 'info');
  rest.sort((a, b) => {
    const va = altValue(comp, a.alt), vb = altValue(comp, b.alt);
    if (va.tier !== vb.tier) return va.tier - vb.tier;
    if (va.value !== vb.value) return va.value - vb.value;
    return a.i - b.i;
  });
  return rest.concat(info);
}

// Reuse totals cover the whole component quantity even when Palats holds
// fewer. That is the intended behaviour for an early-planning tool (stock turns
// over long before procurement), but until 2026-08-15 the table said nothing
// about it: a row could read "9 600 kr" for 30 windows off a listing with 3 in
// stock. Show the gap instead of hiding it.
function stockShortfall(alt, pc) {
  if (alt.alternative_type !== 'reuse') return null;
  if (alt.available_quantity === null || alt.available_quantity === undefined) return null;
  if (!pc || !(pc.quantity > 0)) return null;
  if (alt.available_quantity >= pc.quantity) return null;
  return { have: alt.available_quantity, need: pc.quantity };
}

function stockNote(alt, pc) {
  const short = stockShortfall(alt, pc);
  if (!short) return '';
  return '<div style="font-size:10px;color:var(--kk-red-orange);margin-top:2px">'
    + esc(String(short.have)) + ' av ' + esc(String(short.need))
    + ' i lager. Siffrorna räknas på hela behovet.</div>';
}

// Three kinds of figure land in the same column and must not read alike.
// "Annonspris" is what a specific second-hand item actually costs.
// "Marknadspris" is a web-searched installed price for that KIND of material.
// "AI-uppskattat pris" is the model's own guess, used when the search found no
// source at all. Better than an empty cell for someone comparing two materials,
// but the weakest of the three, so it is the one that gets a warning colour.
function priceBasisNote(alt) {
  if (alt.price_basis === 'market_estimate')
    return '<div style="font-size:10px;color:var(--kk-gray-500)">marknadspris</div>';
  if (alt.price_basis === 'llm_estimate')
    return '<div style="font-size:10px;color:var(--kk-red-orange)" title="Ingen priskälla hittades vid webbsökning. Siffran är språkmodellens egen uppskattning av ett typiskt installerat pris och behöver kontrolleras innan den används som underlag.">AI-uppskattat pris</div>';
  return '';
}

function formatCost(alt) {
  // An EPD-verified alternative that could not be web-priced survives the B1
  // filter with cost_sek 0. Rendering that as "0 kr" read as free; Johanna
  // reported it as "pris visas inte för nyproducerade material".
  if (!(alt.cost_sek > 0)) return '<span style="color:var(--kk-gray-500)">Pris saknas</span>';
  return Math.round(alt.cost_sek).toLocaleString('sv') + ' kr';
}

function formatValuePerKg(comp, alt) {
  const v = altValue(comp, alt);
  if (!v.priced) return '<span style="color:var(--kk-gray-500)">-</span>';
  if (v.saved <= 0) return '<span style="color:var(--kk-gray-500)">ingen besparing</span>';
  if (v.tier === 0) return '<span style="color:var(--green-saving)">billigare</span>';
  return Math.round(v.value).toLocaleString('sv') + ' kr';
}

function alternativHtml(st, cfg) {
  const data = st.alternatives;
  let html = (cfg && cfg.hideTitle) ? '' : '<div class="section-title">J\u00e4mf\u00f6relse per komponent</div>';
  html += '<div class="method-label">Klimatmetod: GWP-fossil, livscykelskedena A1-A3 (Boverkets klimatdatabas)</div>';
  html += '<div class="source-legend"><span><span class="source-badge source-verified">EPD</span> Verifierad k\u00e4lla</span><span><span class="source-badge source-aggregate">EPD-typvärde</span> Kategori-typvärde (övre halvan)</span><span><span class="source-badge source-estimate">Est.</span> Uppskattning</span></div>';
  const projComps = (st.project && st.project.components) || [];
  data.components.forEach(comp => {
    const pc = projComps.find(p => p.id === comp.component_id);
    const qtyLabel = pc ? esc(pc.quantity) + ' ' + esc(pc.unit) + ' ' + quantitySourceBadge(pc.quantity_source) : '';
    const usageBlock = (pc && pc.usage_context) ? '<div class="usage-context"><span class="usage-context-label">Användning</span>' + esc(pc.usage_context) + '</div>' : '';
    const header = '<h3>' + esc(comp.component_name) + '</h3>' + (qtyLabel ? '<div style="font-size:12px;color:var(--kk-gray-500);margin-top:2px">Antal: ' + qtyLabel + '</div>' : '') + usageBlock;
    html += '<div class="comp-card"><div class="comp-card-header">' + header + '</div>';
    html += '<table class="comp-table"><thead><tr><th style="width:32px"></th><th>Typ</th><th>Material</th><th>K\u00e4lla</th><th style="text-align:right">CO\u2082e (kg)</th><th style="text-align:right">Kostnad</th><th style="text-align:right" title="Merkostnad delat med sparade kilo CO\u2082e. L\u00e4gst v\u00e4rde \u00f6verst.">kr/sparat kg</th><th></th></tr></thead><tbody>';
    const blSel = st.selections[comp.component_id] && st.selections[comp.component_id].selected_alternative.name === 'Baslinje';
    // Look up the Boverket product used for this component's baseline so the
    // baseline row shows the actual material (e.g. "Gipsskiva, standardskiva").
    // Set only for genuine same-material Boverket matches; empty for EPD-typvärde
    // and LLM-uppskattning baselines.
    const blBaselineComp = (st.baseline && st.baseline.components) ? st.baseline.components.find(b => b.component_id === comp.component_id) : null;
    const blProduct = (blBaselineComp && blBaselineComp.boverket_product) ? blBaselineComp.boverket_product : '';
    const blSource = (blBaselineComp && blBaselineComp.source) ? blBaselineComp.source : 'NollCO2';
    const blMaterialCell = blProduct
      ? '<div style="font-weight:500">Konventionellt</div><div style="font-size:11px;color:var(--kk-gray-500);font-style:italic;margin-top:2px"><span style="font-style:normal;font-weight:500;color:var(--kk-gray-500);font-size:9.5px;letter-spacing:0.8px;text-transform:uppercase;display:block;margin-bottom:1px">Boverket-produkt</span>' + esc(blProduct) + '</div>'
      : '<div style="font-weight:500">Konventionellt</div>';
    // component_id is LLM-generated and round-trips through user-editable
    // Supabase rows — escape it before interpolating into HTML attributes / ids.
    const cid = esc(comp.component_id);
    html += '<tr class="alt-row' + (blSel ? ' selected' : '') + '" data-comp="' + cid + '" data-alt="baseline">' +
      '<td><input type="radio" name="' + cid + '"' + (blSel ? ' checked' : '') + '></td>' +
      '<td><span class="type-badge type-baseline">Baslinje</span></td>' +
      '<td>' + blMaterialCell + '</td><td style="font-size:11px">' + (blSource.includes('Boverket') ? '<span class="source-badge source-verified">BVK</span>' : (blSource.includes('typvärde') ? '<span class="source-badge source-aggregate">EPD-typvärde</span>' : '<span class="source-badge source-estimate">Est.</span>')) + ' NollCO2</td>' +
      '<td style="text-align:right">' + Math.round(comp.baseline_co2e_kg) + overrideBadge(comp.baseline_co2e_override) + '</td>' +
      '<td style="text-align:right">' + Math.round(comp.baseline_cost_sek).toLocaleString('sv') + ' kr' + overrideBadge(comp.baseline_cost_override) + '</td>' +
      '<td style="text-align:right;color:var(--kk-gray-500)">referens</td><td></td></tr>';
    rankedAlternatives(comp).forEach(entry => {
      const alt = entry.alt, i = entry.i;
      const rowId = cid + '_' + i;
      if (alt.alternative_type === 'info') {
        html += '<tr style="opacity:0.6">' +
          '<td></td>' +
          '<td>' + getTypeBadge(alt) + '</td>' +
          '<td colspan="5" style="font-size:12px;color:var(--kk-gray-500)">' + esc(alt.name) + '</td>' +
          '<td>' + (alt.reasoning ? '<button class="reasoning-toggle" onclick="toggleReasoning(\'' + rowId + '\',event)">Visa mer</button>' : '') + '</td></tr>';
        if (alt.reasoning) {
          html += '<tr class="reasoning-row" id="reasoning-' + rowId + '" style="display:none"><td colspan="8">' + esc(alt.reasoning) + '</td></tr>';
        }
        return;
      }
      const saving = comp.baseline_co2e_kg > 0 ? Math.round((1 - alt.co2e_kg / comp.baseline_co2e_kg) * 100) : 0;
      const isSel = st.selections[comp.component_id] && st.selections[comp.component_id].selected_alternative.name === alt.name;
      // Decompose total for reuse alternatives where units match (no trailing *).
      // Lets the user see e.g. "45 st \u00d7 320 kr = 14 400 kr" inline rather than
      // hidden in Visa mer \u2014 answers Johanna's "varf\u00f6r 45 lampor" without exposing
      // adjustability yet.
      const showBreakdown = alt.alternative_type === 'reuse' && !alt.name.endsWith('*') && pc && pc.quantity > 0 && alt.cost_sek > 0;
      const perUnit = showBreakdown ? Math.round(alt.cost_sek / pc.quantity) : 0;
      const costCell = !(alt.cost_sek > 0)
        ? formatCost(alt)
        : (alt.name.endsWith('*')
          ? Math.round(alt.cost_sek).toLocaleString('sv') + ' kr/st *'
          : (showBreakdown
            ? '<div style="line-height:1.3">' + Math.round(alt.cost_sek).toLocaleString('sv') + ' kr<div style="font-size:10px;color:var(--kk-gray-500)">' + esc(String(pc.quantity)) + ' \u00d7 ' + perUnit.toLocaleString('sv') + ' kr annonspris</div></div>'
            : '<div style="line-height:1.3">' + Math.round(alt.cost_sek).toLocaleString('sv') + ' kr' + priceBasisNote(alt) + '</div>'));
      html += '<tr class="alt-row' + (isSel ? ' selected' : '') + '" data-comp="' + cid + '" data-alt="' + i + '">' +
        '<td><input type="radio" name="' + cid + '"' + (isSel ? ' checked' : '') + '></td>' +
        '<td>' + getTypeBadge(alt) + '</td>' +
        '<td style="font-weight:500">' + esc(alt.name) + stockNote(alt, pc) + '</td>' +
        '<td style="font-size:11px">' + gwpBasisBadge(alt) + formatSource(alt.source) + '</td>' +
        '<td style="text-align:right">' + Math.round(alt.co2e_kg) + ' <span style="color:' + (saving >= 0 ? 'var(--green-saving)' : 'var(--kk-red-orange)') + ';font-size:11px">' + (saving >= 0 ? '\u2193' : '\u2191') + Math.abs(saving) + '%</span></td>' +
        '<td style="text-align:right">' + costCell + '</td>' +
        '<td style="text-align:right;font-size:12px">' + formatValuePerKg(comp, alt) + '</td>' +
        '<td>' + (alt.reasoning ? '<button class="reasoning-toggle" onclick="toggleReasoning(\'' + rowId + '\',event)">Visa mer</button>' : '') + '</td></tr>';
      if (alt.reasoning) {
        html += '<tr class="reasoning-row" id="reasoning-' + rowId + '" style="display:none"><td colspan="8">' + esc(alt.reasoning) + '</td></tr>';
      }
    });
    html += '</tbody></table></div>';
  });
  // Check if any alternatives have per-article pricing
  const hasPerArticle = data.components.some(c => c.alternatives.some(a => a.name.endsWith('*')));
  if (hasPerArticle) {
    html += '<div style="font-size:12px;color:var(--kk-gray-500);margin:8px 0;font-style:italic">* Pris per artikel (yta per artikel ok\u00e4nd). Se \u201cVisa mer\u201d f\u00f6r detaljer.</div>';
  }
  // One shared explanation when any reuse row is priced beyond current stock,
  // so the assumption is stated once in plain language rather than only as a
  // red line per row.
  const hasShortfall = data.components.some(c => {
    const pc2 = projComps.find(p => p.id === c.component_id);
    return c.alternatives.some(a => stockShortfall(a, pc2));
  });
  if (hasShortfall) {
    html += '<div style="font-size:12px;color:var(--kk-gray-500);margin:8px 0">'
      + 'Kostnad och klimatnytta f\u00f6r \u00e5terbruk r\u00e4knas p\u00e5 hela behovet, \u00e4ven n\u00e4r Palats '
      + 'har f\u00e4rre artiklar just nu. Det \u00e4r avsiktligt i ett tidigt planeringsskede, '
      + 'eftersom lagret oms\u00e4tts innan n\u00e5got handlas upp. St\u00e4m av tillg\u00e5ngen innan '
      + 'siffran g\u00e5r vidare till ett beslutsunderlag.</div>';
  }
  html += '<div id="summaryArea"></div>';
  html += '<button class="btn" id="reportBtn" onclick="generateReport()" disabled title="V\u00e4lj ett alternativ per komponent">Generera rapport</button>';
  html += '<div id="missingHint" style="font-size:12px;color:var(--kk-gray-500);margin-top:6px;font-style:italic"></div>';
  return html;
}

// The bindings live apart from the two callers (tab and sheet) so the markup has
// one owner and the handlers have one owner. A second copy is how the sheet and
// the tab would quietly drift apart, which is the thing §12 exists to prevent.
function bindAltRows() {
  document.querySelectorAll('.alt-row').forEach(row => {
    row.onclick = function() { selectAlt(this.dataset.comp, this.dataset.alt, this); };
  });
  if (state.alternatives && Object.keys(state.selections).length > 0) updateSummary();
}

function renderAlternativContent() {
  document.getElementById('resultContent').innerHTML = alternativHtml(effectiveState(state));
  bindAltRows();
}

function rapportHtml(st) {
  let html = '<div class="report-area">' + renderMd(st.reportMarkdown) + '</div>';
  html += '<div style="margin-top:12px;display:flex;gap:8px">';
  html += '<button class="btn" id="dlDocxBtn"><svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" style="vertical-align:-2px;margin-right:4px"><path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/></svg>Ladda ner Word (.docx)</button>';
  html += '<button class="btn btn-secondary" id="dlBtn"><svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" style="vertical-align:-2px;margin-right:4px"><path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/></svg>Ladda ner (.md)</button>';
  html += '</div>';
  return html;
}

function renderRapportContent() {
  document.getElementById('resultContent').innerHTML = rapportHtml(state);
  bindReportDownloads();
}

function bindReportDownloads() {
  if (!document.getElementById('dlBtn')) return;
  document.getElementById('dlBtn').onclick = () => {
    const blob = new Blob([state.reportMarkdown], {type:'text/markdown'});
    const a = document.createElement('a'); a.href = URL.createObjectURL(blob); a.download = 'aida-rapport.md'; a.click();
  };
  document.getElementById('dlDocxBtn').onclick = async () => {
    const btn = document.getElementById('dlDocxBtn');
    btn.disabled = true; btn.textContent = 'Skapar dokument...';
    try {
      const r = await authFetch('/api/report/docx', {method:'POST', headers:{'Content-Type':'application/json'}, body: JSON.stringify({markdown: state.reportMarkdown})});
      if (!r.ok) { const d = await r.json(); alert('Fel: ' + d.error); return; }
      const blob = await r.blob();
      const today = new Date().toISOString().slice(0,10);
      const a = document.createElement('a'); a.href = URL.createObjectURL(blob); a.download = 'Aida_rapport_' + today + '.docx'; a.click();
    } catch(e) { alert('Fel: ' + e.message); }
    finally { btn.disabled = false; btn.innerHTML = '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" style="vertical-align:-2px;margin-right:4px"><path d="M21 15v4a2 2 0 0 1-2 2H5a2 2 0 0 1-2-2v-4"/><polyline points="7 10 12 15 17 10"/><line x1="12" y1="15" x2="12" y2="3"/></svg>Ladda ner Word (.docx)'; }
  };
}

// === The sheet (Arbetsblad, orchestration-redesign §12.2) ===
//
// Same four section renderers as the tabs, stacked in one column instead of
// hidden behind each other. The section list is data, so a fifth section (or the
// follow-up set in step 5) is a row here rather than a new render function.
//
// `ready` is what has to be true before the run can start, and it is the reason
// the empty state can offer a button at all: an empty section with no way to
// fill it would just be a scold. `action` is null when the prerequisite is
// missing, and then the text alone says what is needed.
// === Uppföljning (§12.6) ===
// The three renderers read st.followup, which the server computed. There is no
// arithmetic in here on purpose: the outcome figure ends up in a klimatredovisning
// and having two implementations of it, one per language, is how the sheet and
// the document end up disagreeing while both look authoritative.

const QUALITY_LABELS = {product: 'Produktspecifik EPD', generic: 'Generisk EPD',
                        typvarde: 'Kategorins typvärde', reuse: 'Återbruk',
                        none: 'Ingen träff'};

// The badge says what the figure rests on, in the same visual grammar as every
// other source badge on the page: filled and warm when it comes from a real
// declaration, outlined when it does not come from our sources at all.
function qualityBadge(quality) {
  const label = QUALITY_LABELS[quality] || QUALITY_LABELS.none;
  if (quality === 'product') return '<span class="source-badge source-verified" title="' + esc(label) + '">EPD</span>';
  if (quality === 'generic') return '<span class="source-badge source-verified" title="' + esc(label) + '">Generisk</span>';
  if (quality === 'reuse') return '<span class="type-badge type-reuse" title="' + esc(label) + '">Återbruk</span>';
  if (quality === 'typvarde') return '<span class="source-badge source-aggregate" title="' + esc(label) + '">Typvärde</span>';
  return '<span class="source-badge source-estimate" title="' + esc(label) + '">Ingen träff</span>';
}

function followupRows(st) {
  return (st.followup && st.followup.rows) || [];
}

function installeratHtml(st, cfg) {
  const rows = followupRows(st);
  const edit = !!(cfg && cfg.editable);
  let html = (cfg && cfg.hideTitle) ? '' : '<div class="section-title">Installerat</div>';
  html += '<div class="method-label">Vad som faktiskt sattes in, bredvid det som planerades</div>';
  html += '<div class="comp-card"><table class="comp-table"><thead><tr>'
        + '<th>Komponent</th><th>Installerad produkt</th><th>Mängd</th>'
        + '<th>Planerat</th><th>Transport (km)</th>'
        + '<th>Kostnad (SEK)</th><th>Prisunderlag</th></tr></thead><tbody>';
  rows.forEach(r => {
    const cid = r.component_id;
    const planned = (r.planned_quantity != null ? fmtNum(r.planned_quantity) + ' ' + esc(r.unit || '') : '—');
    html += '<tr>'
      + '<td><strong>' + esc(r.name) + '</strong></td>'
      + '<td>' + (edit ? asBuiltInput(cid, 'installed_name', r.installed_name, 'Installerad produkt för ' + r.name, 'text')
                       : esc(r.installed_name || '—')) + '</td>'
      + '<td>' + (edit ? asBuiltInput(cid, 'quantity', r.installed_quantity, 'Installerad mängd för ' + r.name, 'number')
                       : (r.installed_quantity != null ? fmtNum(r.installed_quantity) : '—'))
      + ' <span style="color:var(--kk-gray-500)">' + esc(r.installed_unit || '') + '</span></td>'
      + '<td style="color:var(--kk-gray-500)">' + planned + '</td>'
      // Recorded, never converted to emissions: that needs a mass per component
      // Aida does not have. The column exists because the outcome note promises
      // the figure ("Transporten (140 km) är inte omräknad"), and a note naming
      // a number the sheet cannot collect is a note about nothing.
      + '<td>' + (edit ? asBuiltInput(cid, 'transport_km', r.transport_km, 'Transportsträcka för ' + r.name, 'number')
                       : (r.transport_km != null ? fmtNum(r.transport_km) : '—')) + '</td>'
      + '<td>' + (edit ? asBuiltInput(cid, 'actual_cost', r.actual_cost_sek, 'Verklig kostnad för ' + r.name, 'number')
                       : (r.actual_cost_sek != null ? fmtNum(r.actual_cost_sek) : '—')) + '</td>'
      + '<td>' + (edit ? asBuiltInput(cid, 'cost_source', r.cost_source, 'Prisunderlag för ' + r.name, 'text')
                       : esc(r.cost_source || '—')) + '</td>'
      + '</tr>';
  });
  html += '</tbody></table></div>';
  return html;
}

function matchningHtml(st, cfg) {
  const rows = followupRows(st);
  const edit = !!(cfg && cfg.editable);
  let html = (cfg && cfg.hideTitle) ? '' : '<div class="section-title">Matchning</div>';
  html += '<div class="method-label">Vilken miljödeklaration som gäller för det installerade</div>';
  html += '<div class="comp-card"><table class="comp-table"><thead><tr>'
        + '<th>Komponent</th><th>Bunden deklaration</th><th>Underlag</th>'
        + '<th>kg CO₂e per enhet</th>' + (edit ? '<th></th>' : '') + '</tr></thead><tbody>';
  rows.forEach(r => {
    const e = r.epd;
    // An unbound row is not an error and is not styled as one. It goes to the
    // report as "uppskattad", which §12.6 names as a valid outcome.
    const bound = e ? esc(e.name || e.id) + (e.reg_no ? subLine('Reg.nr', esc(e.reg_no)) : '')
                    : '<span style="color:var(--kk-gray-500)">Ingen bunden ännu</span>';
    const per = (e && e.gwp_per_unit != null)
      ? fmtNum(e.gwp_per_unit) + ' / ' + esc(e.unit || '')
        + (e.gwp_basis === 'ghg' ? ' ' + gwpBasisBadge({gwp_basis: 'ghg'}) : '')
      : '<span style="color:var(--kk-gray-500)">—</span>';
    html += '<tr>'
      + '<td><strong>' + esc(r.name) + '</strong>'
      + (r.installed_name ? subLine('Installerat', esc(r.installed_name)) : '') + '</td>'
      + '<td>' + bound + '</td>'
      + '<td>' + qualityBadge(r.match_quality) + '</td>'
      + '<td>' + per + '</td>'
      + (edit ? '<td><button type="button" class="btn-na-cancel" onclick="openMatch(\'' + esc(r.component_id) + '\')">'
                + (e ? 'Byt' : 'Sök') + '</button></td>' : '')
      + '</tr>';
    if (edit) html += matchFormRow(r);
  });
  html += '</tbody></table></div>';
  return html;
}

// Which row has its search open, and what came back. One at a time, for the
// same reason the override editor is: two open searches in one table would let
// someone start in one and bind from the other.
let _matchOpen = null;
let _matchResults = [];
let _matchBusy = false;

function openMatch(cid) {
  if (cellsLocked()) return;
  _matchOpen = cid;
  _matchResults = [];
  refreshResults();
  const el = document.getElementById('matchQuery');
  if (el) { el.focus(); el.select(); }
}

function closeMatch() {
  _matchOpen = null;
  _matchResults = [];
  refreshResults();
}

function matchFormRow(r) {
  if (_matchOpen !== r.component_id) return '';
  const hint = _matchBusy ? 'Söker…'
    : (_matchResults.length ? 'Välj den produkt som faktiskt installerades.'
       : 'Sök på produktnamnet. Hittar du den inte, lämna raden obunden: den redovisas då som uppskattad.');
  let html = '<tr class="override-row"><td colspan="5">'
    + '<div class="override-form">'
    + '<div class="ov-note"><label for="matchQuery">Sök i EPD-registret</label>'
    + '<input id="matchQuery" type="text" value="' + esc(r.installed_name || r.name || '') + '"'
    + ' placeholder="Produktnamn, till exempel iQ Granit" onkeydown="if(event.key===\'Enter\'){event.preventDefault();runMatch();}"></div>'
    + '<div class="ov-actions">'
    + '<button type="button" class="btn-na-cancel" onclick="closeMatch()">Stäng</button>'
    + '<button type="button" class="btn-na-save" onclick="runMatch()">Sök</button>'
    + (r.epd ? '<button type="button" class="btn-na-cancel ov-clear" onclick="unbindEpd(\'' + esc(r.component_id) + '\')">Ta bort bindningen</button>' : '')
    + '</div>'
    + '<div class="override-hint" id="matchHint">' + esc(hint) + '</div>';
  if (_matchResults.length) {
    html += '<div class="match-results">';
    _matchResults.forEach(c => {
      html += '<button type="button" class="match-hit" onclick="bindMatch(\'' + esc(r.component_id)
        + '\',\'' + esc(c.id) + '\',\'' + esc(c.version || '') + '\')">'
        + '<span class="match-hit-name">' + esc(c.name) + ' ' + estimatedBadge(c) + '</span>'
        + '<span class="match-hit-meta">' + esc(c.owner || '') + (c.geo ? ' · ' + esc(c.geo) : '')
        + (c.reg_no ? ' · ' + esc(c.reg_no) : '') + '</span></button>';
    });
    html += '</div>';
  }
  return html + '</div></td></tr>';
}

async function runMatch() {
  const el = document.getElementById('matchQuery');
  const query = (el && el.value || '').trim();
  const hint = document.getElementById('matchHint');
  if (!query) {
    if (hint) { hint.className = 'override-hint bad'; hint.textContent = 'Skriv ett produktnamn att söka på.'; }
    return;
  }
  _matchBusy = true;
  if (hint) { hint.className = 'override-hint'; hint.textContent = 'Söker…'; }
  try {
    const comp = (followupRows(effectiveState(state)).find(r => r.component_id === _matchOpen)) || {};
    const r = await authFetch('/api/match', {method: 'POST', headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({query: query, component_hint: comp.name || '', limit: 12})});
    const d = await r.json();
    _matchBusy = false;
    if (d.error) {
      if (hint) { hint.className = 'override-hint bad'; hint.textContent = d.error; }
      return;
    }
    _matchResults = d.candidates || [];
    refreshResults();
    if (!_matchResults.length) {
      const h2 = document.getElementById('matchHint');
      if (h2) h2.textContent = 'Inga träffar. Lämna raden obunden så redovisas den som uppskattad, '
                             + 'eller sök på ett kortare produktnamn.';
    }
  } catch (e) {
    _matchBusy = false;
    if (hint) { hint.className = 'override-hint bad'; hint.textContent = 'Sökningen misslyckades: ' + e.message; }
  }
}

function bindMatch(cid, epdId, version) {
  _matchOpen = null;
  _matchResults = [];
  return sendMutation('bind_epd', {component_id: cid, epd_id: epdId, version: version,
                                   match_quality: 'product'});
}

function unbindEpd(cid) {
  _matchOpen = null;
  _matchResults = [];
  return sendMutation('bind_epd', {component_id: cid, epd_id: null, match_quality: 'none'});
}

function utfallHtml(st, cfg) {
  const rows = followupRows(st);
  const t = (st.followup && st.followup.totals) || null;
  let html = (cfg && cfg.hideTitle) ? '' : '<div class="section-title">Utfall</div>';
  html += '<div class="method-label">Utfall mot baslinje och mot plan, GWP-fossil A1-A3</div>';
  if (t) {
    html += '<div class="summary">';
    html += '<div class="card"><div class="card-title">Utfall</div><div class="value">'
          + fmtNum(t.outcome_co2e_kg) + '</div><div class="sublabel">kg CO₂e</div></div>';
    html += '<div class="card"><div class="card-title">Mot baslinjen</div><div class="value">'
          + fmtNum(t.avoided_vs_baseline_kg) + '</div><div class="sublabel">kg CO₂e undveks</div></div>';
    html += '<div class="card"><div class="card-title">Mot planen</div><div class="value">'
          + (t.deviation_vs_plan_kg > 0 ? '+' : '') + fmtNum(t.deviation_vs_plan_kg)
          + '</div><div class="sublabel">kg CO₂e avvikelse</div></div>';
    html += '</div>';
    // The sentence that keeps the three cards honest. Without it the totals
    // read as covering the whole project, and a saving computed over three of
    // five components looks exactly like one computed over all five.
    if (t.rows_counted !== t.rows_total) {
      html += '<div class="method-label" style="color:var(--kk-red-orange)">Summorna gäller '
            + t.rows_counted + ' av ' + t.rows_total + ' komponenter. Utanför: '
            + esc((t.uncounted_names || []).join(', '))
            + '. Baslinje och plan är räknade över samma rader, så jämförelsen gäller.</div>';
    }
    if (t.cost_rows_counted) {
      html += '<div class="method-label">Kostnad: ' + fmtNum(t.actual_cost_sek) + ' SEK mot planerade '
            + fmtNum(t.planned_cost_sek) + ' SEK ('
            + (t.cost_difference_sek > 0 ? '+' : '') + fmtNum(t.cost_difference_sek) + ' SEK), över '
            + t.cost_rows_counted + ' komponenter med verkligt pris.</div>';
    }
  }
  html += '<div class="comp-card"><table class="comp-table"><thead><tr>'
        + '<th>Komponent</th><th>Utfall (kg)</th><th>Underlag</th>'
        + '<th>Baslinje (kg)</th><th>Planerat (kg)</th></tr></thead><tbody>';
  rows.forEach(r => {
    const outcome = r.outcome_co2e_kg != null
      ? fmtNum(r.outcome_co2e_kg)
      : '<span style="color:var(--kk-gray-500)">Räknas inte</span>';
    html += '<tr>'
      + '<td><strong>' + esc(r.name) + '</strong></td>'
      + '<td>' + outcome + (r.outcome_note ? subLine('Notering', esc(r.outcome_note)) : '') + '</td>'
      + '<td>' + qualityBadge(r.match_quality) + '</td>'
      + '<td>' + (r.baseline_co2e_kg != null ? fmtNum(r.baseline_co2e_kg) : '—')
      + overrideBadge(r.baseline_co2e_override) + '</td>'
      + '<td>' + (r.planned_co2e_kg != null ? fmtNum(r.planned_co2e_kg) : '—')
      + (r.planned_is_baseline ? subLine('Underlag', 'Inget alternativ valdes, så planen är baslinjen') : '') + '</td>'
      + '</tr>';
  });
  html += '</tbody></table></div>';
  return html;
}

const SHEET_SECTIONS = [
  {key: 'projekt', title: 'Projektinformation', html: projektHtml,
   has: st => !!st.project,
   empty: 'Beskriv projektet i chatten, så förs det in här: byggnad, yta, och vad som ska göras.',
   ready: () => false, action: null},
  {key: 'baslinje', title: 'Baslinje (NollCO2-metoden)', html: baslinjeHtml,
   has: st => !!st.baseline,
   empty: 'Baslinjen räknas fram ur projektets komponenter och visar utsläppen om allt byggs nytt.',
   ready: st => !!st.project, action: {label: 'Räkna baslinje', fn: 'runBaseline'}},
  {key: 'alternativ', title: 'Jämförelse per komponent', html: alternativHtml,
   has: st => !!st.alternatives,
   empty: 'Aida söker återbruk och lägre nyproduktion per komponent, och jämför mot baslinjen.',
   ready: st => !!st.baseline, action: {label: 'Sök alternativ', fn: 'runAlternatives'}},
  {key: 'rapport', title: 'Rapport', html: rapportHtml,
   has: st => !!st.reportMarkdown,
   empty: 'Rapporten skrivs när varje komponent har ett val.',
   ready: st => !!(st.alternatives && st.alternatives.components.length
                   && st.alternatives.components.every(c => st.selections[c.component_id])),
   action: {label: 'Generera rapport', fn: 'generateReport'}},
];

// Follow-up is its own list, not a flag on the one above. The two sheets share
// only 'projekt', and the rest answer different questions: one is about what to
// build, the other about what got built. §12.1 puts baseline and plan in the
// outcome table as COLUMNS rather than as sections, which is why they are not
// repeated here.
const FOLLOWUP_SECTIONS = [
  {key: 'projekt', title: 'Projektinformation', html: projektHtml,
   has: st => !!st.project,
   empty: 'Beskriv projektet i chatten, så förs det in här: byggnad, yta, och vad som gjordes.',
   ready: () => false, action: null},
  {key: 'installerat', title: 'Installerat', html: installeratHtml,
   has: st => !!st.project,
   empty: 'När projektet är inläst listas komponenterna här, så du kan fylla i vad som faktiskt sattes in.',
   ready: () => false, action: null},
  {key: 'matchning', title: 'Matchning', html: matchningHtml,
   has: st => !!st.project,
   empty: 'Här binder du en miljödeklaration till varje installerad produkt.',
   ready: () => false, action: null},
  {key: 'utfall', title: 'Utfall', html: utfallHtml,
   has: st => !!st.project,
   empty: 'Utfallet räknas när något är registrerat som installerat.',
   ready: () => false, action: null},
  {key: 'redovisning', title: 'Klimatredovisning', html: rapportHtml,
   has: st => !!st.reportMarkdown,
   empty: 'Redovisningen skrivs ur utfallet, med källkvalitet per byggdel och osäkerheterna namngivna.',
   ready: st => !!(st.followup && st.followup.totals && st.followup.totals.rows_counted > 0),
   action: {label: 'Skriv klimatredovisning', fn: 'generateFollowupReport'}},
];

function sectionsForMode() {
  return isFollowup() ? FOLLOWUP_SECTIONS : SHEET_SECTIONS;
}

// Stegvis disables its confirm button synchronously, before any await, because
// the click starts minutes of work. The sheet needs the same: it is not redrawn
// until the run returns, so a second click would start a second pipeline against
// the same analysis. The redraw discards this disabled state, which is correct.
function sheetAction(btn, fn) {
  btn.disabled = true;
  btn.style.opacity = '0.5';
  maybeAskForNotifications();
  window[fn]();
}

function sheetHtml(st, sections) {
  let html = '<div class="sheet">';
  (sections || SHEET_SECTIONS).forEach(sec => {
    const filled = sec.has(st);
    html += '<section class="sheet-section" id="sheet-' + sec.key + '">';
    html += '<div class="sheet-head"><h2>' + esc(sec.title) + '</h2>'
          + (filled ? '' : '<span class="sheet-pending">väntar</span>') + '</div>';
    if (filled) {
      html += sec.html(st, {hideTitle: true, editable: true});
    } else {
      const act = sec.ready(st) ? sec.action : null;
      html += '<div class="sheet-empty"><p>' + esc(sec.empty) + '</p>'
            + (act ? '<button class="btn" onclick="sheetAction(this, \'' + act.fn + '\')">' + esc(act.label) + '</button>' : '')
            + '</div>';
    }
    html += '</section>';
  });
  return html + '</div>';
}

function renderSheet() {
  document.getElementById('resultContent').innerHTML =
    sheetHtml(effectiveState(state), sectionsForMode());
  _populateNeedsTextarea();
  bindCells();
  bindAltRows();
  bindReportDownloads();
}

// === Selection handling ===
function selectAlt(compId, altIdx, row) {
  row.closest('table').querySelectorAll('.alt-row').forEach(r => r.classList.remove('selected'));
  row.classList.add('selected');
  row.querySelector('input[type=radio]').checked = true;
  const comp = state.alternatives.components.find(c => c.component_id === compId);
  if (altIdx === 'baseline') {
    state.selections[compId] = { id: compId, name: comp.component_name,
      selected_alternative: {name:'Baslinje', co2e_kg: comp.baseline_co2e_kg, cost_sek: comp.baseline_cost_sek, source:'NollCO2'},
      baseline_co2e_kg: comp.baseline_co2e_kg, baseline_cost_sek: comp.baseline_cost_sek };
  } else {
    const alt = comp.alternatives[parseInt(altIdx)];
    state.selections[compId] = { id: compId, name: comp.component_name,
      // available_quantity and price_basis travel with the selection so the
      // report can state the stock assumption behind a reuse figure. The
      // report is the artifact that leaves the tool, so the caveat has to
      // reach it, not just the on-screen table.
      selected_alternative: {name: alt.name, co2e_kg: alt.co2e_kg, cost_sek: alt.cost_sek, source: alt.source,
        available_quantity: (alt.available_quantity === undefined ? null : alt.available_quantity),
        price_basis: alt.price_basis || '', gwp_basis: alt.gwp_basis || ''},
      baseline_co2e_kg: comp.baseline_co2e_kg, baseline_cost_sek: comp.baseline_cost_sek };
  }
  // Durable intent, so a later rerun can rebind it (increment 3).
  rememberIntent(compId, comp.component_name, state.selections[compId].selected_alternative.name);
  updateSummary();
  scheduleAutoSave();
}

// The only place selection totals get summed. Pure, so the test suite can
// exercise it directly \u2014 updateSummary() below writes to the DOM and was
// therefore never covered, which is how the unpriced-as-zero bug survived the
// row-level fix on 2026-08-14.
//
// A selected alternative with cost_sek <= 0 has NO KNOWN PRICE. It is not free:
// the table renders it "Pris saknas", palats_client documents its price field
// as "0 if free/unknown", and an EPD-verified alternative that could not be
// web-priced survives the B1 filter at 0. Adding it as zero kronor understated
// the basket and produced a saving out of a data gap, so unpriced components are
// named and the comparison runs over the priced subset only.
function summaryTotals(sels) {
  const totalCo2 = sels.reduce((s,c) => s + c.selected_alternative.co2e_kg, 0);
  const blCo2 = sels.reduce((s,c) => s + c.baseline_co2e_kg, 0);
  const priced = sels.filter(c => c.selected_alternative.cost_sek > 0);
  const unpriced = sels.filter(c => !(c.selected_alternative.cost_sek > 0)).map(c => c.name);
  const knownCost = priced.reduce((s,c) => s + c.selected_alternative.cost_sek, 0);
  const comparableBl = priced.reduce((s,c) => s + c.baseline_cost_sek, 0);
  // Same rule on the baseline side: a baseline component the model gave no
  // cost for is not a free component.
  const blPriced = sels.filter(c => c.baseline_cost_sek > 0);
  const blCost = blPriced.reduce((s,c) => s + c.baseline_cost_sek, 0);
  return {
    totalCo2: totalCo2, blCo2: blCo2, blCost: blCost,
    blUnpricedCount: sels.length - blPriced.length,
    knownCost: knownCost, comparableBl: comparableBl,
    unpriced: unpriced, costIsPartial: unpriced.length > 0,
    pricedCount: priced.length, total: sels.length,
  };
}

function updateSummary() {
  // Effective, not stored: the cards sit a screen away from the table and
  // would otherwise show a total that no row on the page adds up to.
  const sels = Object.values(effectiveState(state).selections || {});
  const t = summaryTotals(sels);
  const co2Diff = t.totalCo2 - t.blCo2;
  const co2Pct = t.blCo2 > 0 ? Math.round(Math.abs(co2Diff) / t.blCo2 * 100) : 0;
  const co2Arrow = co2Diff <= 0 ? '\u2193' : '\u2191';
  const costDiff = t.knownCost - t.comparableBl;
  const costPct = t.comparableBl > 0 ? Math.round(Math.abs(costDiff) / t.comparableBl * 100) : 0;
  const costArrow = costDiff <= 0 ? '\u2193' : '\u2191';
  // With a hole in the basket the headline is a part, not a total, and the
  // percentage compares the priced components against their own baseline.
  const costTitle = t.costIsPartial ? 'Kostnad (delsumma)' : 'Kostnad';
  const costSub = t.costIsPartial
    ? 'SEK f\u00f6r ' + t.pricedCount + ' av ' + t.total + ' komponenter (' + costArrow + costPct + '% vs deras baslinje)'
    : 'SEK (' + costArrow + costPct + '% vs baslinje)';
  const costGap = t.costIsPartial
    ? '<div style="grid-column:1/-1;font-size:11px;color:var(--kk-red-orange);margin-top:-4px">Saknar pris: '
      + esc(t.unpriced.join(', '))
      + '. Ingen prisuppgift hittades, posten \u00e4r allts\u00e5 inte gratis \u2014 n\u00e5gon totalkostnad g\u00e5r inte att ange.</div>'
    : '';
  document.getElementById('summaryArea').innerHTML =
    '<div class="summary">' +
    '<div class="card' + (co2Diff <= 0 ? ' saving' : '') + '"><div class="card-title">Klimatp\u00e5verkan</div><div class="value">' + Math.round(t.totalCo2).toLocaleString('sv') + '</div><div class="sublabel">kg CO\u2082e (' + co2Arrow + co2Pct + '% vs baslinje)</div></div>' +
    '<div class="card' + (!t.costIsPartial && costDiff <= 0 ? ' saving' : '') + '"><div class="card-title">' + costTitle + '</div><div class="value">' + Math.round(t.knownCost).toLocaleString('sv') + '</div><div class="sublabel">' + costSub + '</div></div>' +
    '<div class="card"><div class="card-title">Baslinje</div><div class="value">' + Math.round(t.blCo2).toLocaleString('sv') + '</div><div class="sublabel">kg CO\u2082e | ' + Math.round(t.blCost).toLocaleString('sv') + ' SEK' + (t.blUnpricedCount ? ' (' + t.blUnpricedCount + ' utan pris)' : '') + '</div></div>' +
    costGap +
    '</div>';
  const allSelected = state.alternatives.components.every(c => state.selections[c.component_id]);
  document.getElementById('reportBtn').disabled = !allSelected;
  const hintEl = document.getElementById('missingHint');
  if (hintEl) {
    if (!allSelected) {
      const missing = state.alternatives.components.filter(c => !state.selections[c.component_id]).map(c => c.component_name);
      hintEl.textContent = 'V\u00e4lj alternativ f\u00f6r: ' + missing.join(', ');
    } else { hintEl.textContent = ''; }
  }
}

// === Supabase auth + persistence ===
const HAS_SUPABASE = {{ 'true' if has_supabase else 'false' }};
const SUPABASE_URL = {{ supabase_url|tojson }};
const SUPABASE_ANON_KEY = {{ supabase_anon_key|tojson }};
let supabaseClient = null;
let currentUser = null;
let currentAnalysisId = null;
let isSignup = false;
let saveTimeout = null;
let saveInProgress = false;

// Auth-aware fetch wrapper with safe JSON parsing
async function authFetch(url, options) {
  options = options || {};
  options.headers = options.headers || {};
  if (supabaseClient) {
    const sess = await supabaseClient.auth.getSession();
    if (sess.data.session) {
      options.headers['Authorization'] = 'Bearer ' + sess.data.session.access_token;
    }
  }
  const resp = await fetch(url, options);
  // Wrap .json() to catch non-JSON responses (e.g. Vercel timeout HTML pages)
  const origJson = resp.json.bind(resp);
  resp.json = async () => {
    const text = await resp.clone().text();
    try { return JSON.parse(text); }
    catch(e) {
      if (resp.status === 504 || resp.status === 502 || text.includes('FUNCTION_INVOCATION_TIMEOUT'))
        return {error: 'Analysen tog f\u00f6r l\u00e5ng tid. F\u00f6rs\u00f6k igen, eller f\u00f6renkla projektbeskrivningen.'};
      if (resp.status >= 500)
        return {error: 'Serverfel (' + resp.status + '). F\u00f6rs\u00f6k igen om en stund.'};
      return {error: 'Ov\u00e4ntat svar fr\u00e5n servern. F\u00f6rs\u00f6k igen.'};
    }
  };
  return resp;
}

// No-op when Supabase not configured
function scheduleAutoSave() {
  if (!HAS_SUPABASE || !currentUser) return;
  if (saveTimeout) clearTimeout(saveTimeout);
  saveTimeout = setTimeout(autoSave, 2000);
}

async function autoSave() {
  if (!supabaseClient || !currentUser) return;
  // A save already running means this call would drop whatever changed since it
  // started. Come back instead of returning: the debounce timer has already
  // fired, so nothing else is going to retry, and the edit disappears with no
  // error anywhere. Typing in the metadata dialog and closing it is where that
  // shows most, because there is no later step that resaves the field.
  if (saveInProgress) {
    if (saveTimeout) clearTimeout(saveTimeout);
    saveTimeout = setTimeout(autoSave, 500);
    return;
  }
  saveInProgress = true;
  const indicator = document.getElementById('saveIndicator');
  if (indicator) { indicator.textContent = 'Sparar...'; indicator.style.display = 'inline'; indicator.style.color = 'var(--kk-gray-400)'; }
  // Directives persist per analysis. Until a dedicated column exists, they ride
  // inside project_data (the server's Project.from_dict ignores unknown keys).
  // Spread so we never mutate the working state.project object.
  const projectDataToSave = state.project
    ? Object.assign({}, state.project, {directives: state.directives,
                                        selection_intent: state.selectionIntent,
                                        mode: state.mode,
                                        overrides: state.overrides})
    : null;
  const analysisData = {
    // Falls back to the property before 'Nytt projekt': someone who names the
    // building before describing the job has already said something more useful
    // than the placeholder.
    name: state.project ? (state.project.name || state.project.building_type || 'Nytt projekt')
                        : (state.propertyRef || 'Nytt projekt'),
    status: state.step,
    project_data: projectDataToSave,
    baseline_data: state.baseline,
    alternatives_data: state.alternatives,
    selections_data: Object.keys(state.selections).length > 0 ? state.selections : null,
    report_markdown: state.reportMarkdown,
    // Increment 4: the conversation gets its own column, not a project_data
    // piggyback like directives — project_data is null until intake succeeds,
    // and the advisory questions we most want to keep happen before that.
    conversation_data: (state.conversation && state.conversation.length) ? state.conversation : null,
    // Own column, like conversation_data and for the same reason: an analysis
    // opened straight in follow-up mode has no project_data to ride in until
    // intake has run, and following up a job Aida never calculated is a normal
    // case, not an edge one.
    as_built_data: Object.keys(state.as_built || {}).length ? state.as_built : null,
    property_ref: state.propertyRef || null,
    // The month input gives 'YYYY-MM'; the column is a DATE, so anchor it to the
    // first of the month. We only ever show the month back, so the day is a
    // storage detail and never a claim about precision we do not have.
    planned_start: state.plannedStart ? (state.plannedStart + '-01') : null,
  };
  try {
    if (currentAnalysisId) {
      const r = await authFetch('/api/analyses/' + currentAnalysisId, {
        method: 'PUT', headers: {'Content-Type':'application/json'},
        body: JSON.stringify(analysisData),
      });
      await r.json();
    } else {
      const r = await authFetch('/api/analyses', {
        method: 'POST', headers: {'Content-Type':'application/json'},
        body: JSON.stringify(analysisData),
      });
      const result = await r.json();
      if (result && result.id) {
        currentAnalysisId = result.id;
        // The chat written before the row existed sits in the 'new' bucket.
        // Move it to this analysis's key so a reload before the next autosave
        // still finds it locally.
        try { localStorage.removeItem('aida_chat_new'); } catch(e) {}
        _saveConversation();
        await loadAnalysesList();
      }
    }
    if (indicator) { indicator.textContent = 'Sparat'; setTimeout(() => { indicator.style.display = 'none'; }, 2000); }
  } catch (e) {
    console.error('Auto-save failed:', e);
    if (indicator) { indicator.textContent = 'Sparfel'; indicator.style.color = 'var(--kk-dark-red)'; }
  }
  finally { saveInProgress = false; }
}

if (HAS_SUPABASE && SUPABASE_URL && SUPABASE_ANON_KEY) {
  supabaseClient = supabase.createClient(SUPABASE_URL, SUPABASE_ANON_KEY);
  initAuth();
} else {
  // Utan auth finns ingen inloggning att vänta in, så introduktionen visas direkt.
  openWelcome();
}

// Reflect persisted sound preference in the toggle label on load.
updateSoundToggle();

// The mode buttons ship with Stegvis marked active in the markup; this keeps
// that honest if the default ever changes, and hides the step rail when an
// analysis loads straight into the sheet.
applyModeChrome();

async function initAuth() {
  const { data: { session } } = await supabaseClient.auth.getSession();
  if (session) { onLogin(session); }
  else { showAuth(); }
  supabaseClient.auth.onAuthStateChange((event, session) => {
    if (event === 'SIGNED_IN' && session) {
      if (!currentUser) onLogin(session);
    }
    else if (event === 'SIGNED_OUT') { currentUser = null; showAuth(); }
  });
}

function showAuth() {
  document.getElementById('authOverlay').style.display = 'flex';
  document.getElementById('appContainer').style.display = 'none';
}

function showApp() {
  document.getElementById('authOverlay').style.display = 'none';
  document.getElementById('appContainer').style.display = '';
  // Först efter inloggning, annars staplas introduktionen på inloggningsrutan.
  openWelcome();
}

async function onLogin(session) {
  currentUser = session.user;
  document.getElementById('userEmail').textContent = currentUser.email;
  showApp();
  const list = await loadAnalysesList();
  if (list && list.length > 0) { await loadAnalysis(list[0].id); }
}

async function handleAuth() {
  const email = document.getElementById('authEmail').value.trim();
  const password = document.getElementById('authPassword').value;
  const errorEl = document.getElementById('authError');
  errorEl.style.display = 'none';
  if (!email || !password) {
    errorEl.textContent = 'Fyll i e-post och lösenord';
    errorEl.style.display = 'block';
    return;
  }
  document.getElementById('authSubmitBtn').disabled = true;
  try {
    const result = isSignup
      ? await supabaseClient.auth.signUp({ email, password })
      : await supabaseClient.auth.signInWithPassword({ email, password });
    if (result.error) {
      const AUTH_ERRORS = {'Invalid login credentials':'Fel e-post eller l\u00f6senord.','Email not confirmed':'Bekr\u00e4fta din e-post innan du loggar in.','User already registered':'Det finns redan ett konto med den e-postadressen.','Password should be at least 6 characters':'L\u00f6senordet m\u00e5ste vara minst 6 tecken.'};
      errorEl.textContent = AUTH_ERRORS[result.error.message] || result.error.message;
      errorEl.style.display = 'block';
    } else if (isSignup && !result.data.session) {
      errorEl.textContent = 'Kolla din e-post för bekräftelselänk';
      errorEl.style.display = 'block';
      errorEl.style.color = 'var(--green-saving)';
    }
  } catch (e) {
    errorEl.textContent = e.message;
    errorEl.style.display = 'block';
  }
  document.getElementById('authSubmitBtn').disabled = false;
}

function toggleAuthMode(e) {
  e.preventDefault();
  isSignup = !isSignup;
  document.getElementById('authSubmitBtn').textContent = isSignup ? 'Skapa konto' : 'Logga in';
  document.getElementById('authToggleText').textContent = isSignup ? 'Har redan konto?' : 'Inget konto?';
  document.getElementById('authToggleLink').textContent = isSignup ? 'Logga in' : 'Skapa konto';
  document.getElementById('authError').style.display = 'none';
}

async function handleLogout() {
  await supabaseClient.auth.signOut();
  currentUser = null;
  currentAnalysisId = null;
  showAuth();
}

// === Project dropdown ===
function toggleProjectMenu() {
  const m = document.getElementById('projectMenu');
  const u = document.getElementById('userMenu');
  if (u) u.style.display = 'none';
  m.style.display = m.style.display === 'none' ? 'block' : 'none';
}

// === Project rename ===
function startRenameProject() {
  document.getElementById('projectMenu').style.display = 'none';
  const span = document.getElementById('projectName');
  if (!span) return;
  const current = (state.project && state.project.name) ? state.project.name : span.textContent;
  const input = document.createElement('input');
  input.type = 'text';
  input.className = 'project-rename-input';
  input.value = current === 'Nytt projekt' ? '' : current;
  input.placeholder = 'Projektnamn';
  input.maxLength = 80;
  span.style.display = 'none';
  span.parentNode.insertBefore(input, span);
  input.focus();
  input.select();
  const commit = (save) => {
    if (!input.parentNode) return;
    const next = input.value.trim();
    if (save && next) {
      if (!state.project) state.project = {name: next};
      else state.project.name = next;
      span.textContent = next;
      if (HAS_SUPABASE && currentUser) scheduleAutoSave();
    }
    input.remove();
    span.style.display = '';
  };
  input.addEventListener('keydown', (e) => {
    if (e.key === 'Enter') { e.preventDefault(); commit(true); }
    else if (e.key === 'Escape') { e.preventDefault(); commit(false); }
  });
  input.addEventListener('blur', () => commit(true));
}

function toggleUserMenu() {
  const m = document.getElementById('userMenu');
  const p = document.getElementById('projectMenu');
  if (p) p.style.display = 'none';
  m.style.display = m.style.display === 'none' ? 'block' : 'none';
}

document.addEventListener('click', (e) => {
  const pd = document.getElementById('projectDropdown');
  const ud = document.getElementById('userDropdown');
  if (pd && !e.target.closest('#projectDropdown')) document.getElementById('projectMenu').style.display = 'none';
  if (ud && !e.target.closest('#userDropdown')) document.getElementById('userMenu').style.display = 'none';
});

async function loadAnalysesList() {
  if (!supabaseClient || !currentUser) return null;
  try {
    const r = await authFetch('/api/analyses');
    const list = await r.json();
    const container = document.getElementById('projectList');
    if (!container) return list;
    container.innerHTML = '';
    if (list && list.length > 0) {
      list.forEach(a => {
        const item = document.createElement('div');
        item.className = 'dropdown-item' + (a.id === currentAnalysisId ? ' active' : '');
        item.style.cssText = 'display:flex;align-items:center;justify-content:space-between;cursor:pointer';
        // Name on top, building and month underneath when they are set. The second
        // line is what tells the user the metadata actually stuck; without it the
        // dialog writes into a void and nobody fills it in twice.
        const textWrap = document.createElement('div');
        textWrap.style.cssText = 'flex:1;min-width:0';
        const nameSpan = document.createElement('div');
        nameSpan.textContent = a.name || 'Nytt projekt';
        nameSpan.style.cssText = 'overflow:hidden;text-overflow:ellipsis;white-space:nowrap';
        textWrap.appendChild(nameSpan);
        const meta = [a.property_ref, a.planned_start ? String(a.planned_start).slice(0, 7) : '']
          .filter(Boolean).join(', ');
        if (meta) {
          const metaLine = document.createElement('div');
          metaLine.className = 'project-meta-line';
          metaLine.textContent = meta;
          textWrap.appendChild(metaLine);
        }
        textWrap.onclick = () => { loadAnalysis(a.id); toggleProjectMenu(); };
        item.appendChild(textWrap);
        const del = document.createElement('button');
        del.style.cssText = 'background:none;border:none;cursor:pointer;color:var(--kk-gray-500);padding:2px 4px;flex-shrink:0';
        del.title = 'Ta bort';
        del.innerHTML = '\u2715';
        del.onclick = async (e) => {
          e.stopPropagation();
          if (!confirm('Ta bort "' + (a.name || 'Nytt projekt') + '"?')) return;
          try { await authFetch('/api/analyses/' + a.id, {method:'DELETE'}); if (a.id === currentAnalysisId) createNewProject(); await loadAnalysesList(); } catch(ex) { alert('Kunde inte ta bort.'); }
        };
        item.appendChild(del);
        container.appendChild(item);
      });
    } else {
      container.innerHTML = '<div style="padding:8px 16px;font-size:12px;color:var(--kk-gray-500)">Inga projekt ännu</div>';
    }
    return list;
  } catch(e) { console.error('Failed to load list:', e); return null; }
}

async function loadAnalysis(id) {
  if (saveTimeout) { clearTimeout(saveTimeout); saveTimeout = null; }
  try {
    const r = await authFetch('/api/analyses/' + id);
    const data = await r.json();
    if (!data || data.error) return;
    currentAnalysisId = id;
    // Increment 4: the analysis carries its own conversation, so switching
    // project swaps the whole transcript AND the model's context together. The
    // previous project's turns cannot leak in, because they are not in this
    // array. restoreUI renders it (and falls back to localStorage for analyses
    // saved before this shipped).
    state.conversation = Array.isArray(data.conversation_data) ? data.conversation_data : [];
    state.pendingDesc = null;
    state.project = data.project_data;
    state.baseline = data.baseline_data;
    state.alternatives = data.alternatives_data;
    state.selections = data.selections_data || {};
    state.reportMarkdown = data.report_markdown;
    state.step = data.status || 'idle';
    state.propertyRef = data.property_ref || '';
    // DATE comes back as 'YYYY-MM-DD'; the month input wants 'YYYY-MM'.
    state.plannedStart = data.planned_start ? String(data.planned_start).slice(0, 7) : '';
    // Restore standing directives (piggybacked in project_data); keep the working
    // project object clean of the persistence-only key.
    state.directives = (state.project && state.project.directives)
      ? state.project.directives : {global: [], byComponent: {}};
    if (state.project && 'directives' in state.project) delete state.project.directives;
    _ensureDirectives();
    // Same piggyback for selection intent. An analysis saved before increment 3
    // has no key, loads as empty, and backfills from its selections — i.e.
    // today's behaviour, with the choices preserved.
    state.selectionIntent = (state.project && state.project.selection_intent)
      ? state.project.selection_intent : {};
    if (state.project && 'selection_intent' in state.project) delete state.project.selection_intent;
    backfillIntent();
    // Same piggyback for the mode. Anything unrecognised (an older analysis, a
    // key from a mode that no longer exists) falls back to Stegvis rather than
    // leaving the view in a shape nothing can render.
    const savedMode = state.project && state.project.mode;
    state.mode = MODES.indexOf(savedMode) === -1 ? 'stepwise' : savedMode;
    if (state.project && 'mode' in state.project) delete state.project.mode;
    // And the overrides. An analysis saved before §12.5 has none, which loads as
    // {} and renders exactly as it always did.
    state.overrides = (state.project && state.project.overrides) || {};
    if (state.project && 'overrides' in state.project) delete state.project.overrides;
    // What was actually installed has its own column, so it loads straight and
    // needs no delete. state.followup is derived and never stored: leaving a
    // previous project's outcome table up while this one's is fetched would show
    // one project's numbers under another project's name.
    state.as_built = data.as_built_data || {};
    state.followup = null;
    document.getElementById('projectName').textContent = data.name || 'Nytt projekt';
    restoreUI();
    await loadAnalysesList();
  } catch(e) { console.error('Failed to load analysis:', e); }
}

function restoreUI() {
  ['projekt','baslinje','alternativ','rapport'].forEach(t => {
    const el = document.getElementById('tab-' + t); if (el) el.disabled = true;
  });
  resetProgressRail();

  applyModeChrome();

  if (state.project) { enableTab('projekt'); setProgressStep('planering'); }
  if (state.baseline) { enableTab('baslinje'); setProgressStep('baslinje'); }
  if (state.alternatives) { enableTab('alternativ'); setProgressStep('sammanstallning'); }
  if (state.reportMarkdown) { enableTab('rapport'); setProgressStep('uppfoljning'); switchTab('rapport'); }
  else if (state.alternatives) { switchTab('alternativ'); }
  else if (state.baseline) { switchTab('baslinje'); }
  else if (state.project) { switchTab('projekt'); }
  // An empty analysis has nothing to switch to, but the sheet still has four
  // sections to show and one of them says how to start. The tab view keeps its
  // own empty state, which is the same sentence in a different place.
  else if (isSheet()) { renderSheet(); }

  const msgs = document.getElementById('messages');
  msgs.innerHTML = '';

  let localChat = null;
  try { localChat = JSON.parse(localStorage.getItem(_chatStorageKey())); } catch(e) {}
  const savedChat = conversationToRestore(state.conversation, localChat);
  state.conversation = [];

  if (savedChat.length > 0) {
    savedChat.forEach(m => {
      const d = document.createElement('div');
      d.className = 'msg ' + m.cls;
      if (m.cls === 'bot' || m.cls === 'system') { d.innerHTML = renderMd(m.text); }
      else { d.textContent = m.text; }
      // Restore confirm buttons for current step only
      if (m.confirm && (
        (state.step === 'intake_done' && m.confirm.btnLabel.includes('baslinje')) ||
        (state.step === 'baseline_done' && m.confirm.btnLabel.includes('alternativ'))
      )) {
        d.innerHTML += '<div class="confirm-actions"><button class="btn-confirm" onclick="confirmStep()">' + m.confirm.btnLabel + '</button></div><div class="confirm-hint">' + m.confirm.hint + '</div>';
      } else if (m.confirm) {
        // Past confirm — show as completed
        d.innerHTML += '<div class="confirm-actions"><button class="btn-confirm" disabled style="background:var(--kk-gray-200);color:var(--kk-gray-400);cursor:default;pointer-events:none">Bekr\u00e4ftad \u2713</button></div>';
      }
      msgs.appendChild(d);
    });
    state.conversation = savedChat;
    const last = msgs.lastElementChild;
    if (last) last.scrollIntoView({behavior:'smooth'});
  } else {
    // Fallback: reconstruct summary from state
    if (!state.project) {
      addMsg('Hej! Beskriv ditt projekt. Berätta vad byggnaden används till, byggnadsår, ungefärlig yta och vilka behoven är.', 'bot');
    } else {
      addMsg('Projekt laddat: ' + (state.project.building_type || 'Okänt') + ', ' + (state.project.area_bta || '?') + ' m\u00b2.', 'bot');
      if (state.step === 'intake_done') {
        const summary = intakeSummary(state.project);
        addConfirmMsg(summary.text, summary.btnLabel, summary.hint);
      } else if (state.step === 'baseline_done') {
        addConfirmMsg(baselineDoneMsg(), 'Bekr\u00e4fta och s\u00f6k alternativ \u2192',
                      'Skriv i chatten om du vill korrigera n\u00e5got.');
      } else if (state.step === 'alternatives_done') addMsg('V\u00e4lj alternativ per komponent i resultatpanelen.', 'bot');
      else if (state.step === 'report_done') addMsg('Rapporten \u00e4r klar.', 'bot');
    }
  }
  updatePlaceholder();
  // Restoring into follow-up mode: the outcome table is derived server-side, so
  // renderSheet above drew three sections with nothing in them. Ask now.
  if (isFollowup()) refreshFollowup();
}

function createNewProject() {
  // Without an account nothing is saved, so a reset here is the only way to
  // lose the current description. With Supabase the old project stays in the
  // list and this is just a switch.
  if (!HAS_SUPABASE && state.project && !confirm('Börja om med ett nytt projekt? Det du beskrivit försvinner.')) return;
  // Close, never toggle: this runs from the top-bar button too, where the menu
  // is already closed and a toggle would open it on top of the fresh project.
  const menu = document.getElementById('projectMenu');
  if (menu) menu.style.display = 'none';
  // Null the id BEFORE removing, so we clear the "new" chat bucket — not the
  // previous project's saved chat log (which must survive switching back).
  currentAnalysisId = null;
  try { localStorage.removeItem(_chatStorageKey()); } catch(e) {}
  state.conversation = [];
  state.project = null; state.baseline = null; state.alternatives = null;
  state.selections = {}; state.pendingDesc = null; state.reportMarkdown = null;
  state.step = 'idle';
  state.directives = {global: [], byComponent: {}};
  state.selectionIntent = {};
  state.overrides = {};
  state.as_built = {};
  state.followup = null;
  state.propertyRef = '';
  state.plannedStart = '';
  document.getElementById('projectName').textContent = 'Nytt projekt';
  ['projekt','baslinje','alternativ','rapport'].forEach(t => {
    const el = document.getElementById('tab-' + t); if (el) el.disabled = true;
  });
  // Also reset the flag, not just the buttons. It stayed true across a new
  // project before, so a mode round-trip could bring the strip back with four
  // dead tabs in it.
  _anyTabEnabled = false;
  applyTabStripChrome();
  document.getElementById('resultContent').innerHTML = '<div class="empty-state"><p>Beskriv ditt projekt i chatten till vänster för att börja.</p></div>';
  resetProgressRail();
  const msgs = document.getElementById('messages');
  msgs.innerHTML = '';
  addMsg('Hej! Beskriv ditt projekt. Berätta vad byggnaden används till, byggnadsår, ungefärlig yta och vilka behoven är.', 'bot');
  setLoading(false);
}
</script>
</body>
</html>
"""


def main():
    import argparse
    parser = argparse.ArgumentParser(description='Aida Web UI')
    parser.add_argument('--port', type=int, default=5002)
    parser.add_argument('--host', type=str, default='127.0.0.1')
    args = parser.parse_args()

    print(f"Aida web UI: http://{args.host}:{args.port}")
    app.run(host=args.host, port=args.port, debug=False)


if __name__ == '__main__':
    main()
